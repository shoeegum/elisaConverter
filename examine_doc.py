#!/usr/bin/env python3
"""
Examine the populated template document to verify formatting
"""

import docx
from docx.document import Document
from docx.text.paragraph import Paragraph

def examine_document(filename):
    """Examine a document and display key sections for formatting verification"""
    doc = docx.Document(filename)
    
    print(f"\nExamining document: {filename}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    
    # Print first few paragraphs
    print("\n=== FIRST PARAGRAPHS ===")
    for i, para in enumerate(doc.paragraphs[:10]):
        if para.text.strip():
            style = para.style.name if hasattr(para.style, 'name') else "None"
            print(f"Para {i}: [{style}] {para.text[:80]}...")
    
    # Find section headers and their content
    print("\n=== KEY SECTIONS ===")
    sections = ["BACKGROUND", "MATERIALS REQUIRED", "REAGENTS"]
    for section in sections:
        for i, para in enumerate(doc.paragraphs):
            if section in para.text:
                print(f"\n--- {section} ---")
                style = para.style.name if hasattr(para.style, 'name') else "None"
                print(f"Header style: {style}")
                print(f"Header text: {para.text}")
                
                # Print the next few paragraphs to see content
                content_paras = []
                for j in range(1, 6):  # Up to 5 paragraphs after header
                    if i + j < len(doc.paragraphs) and doc.paragraphs[i + j].text.strip():
                        content_style = doc.paragraphs[i + j].style.name if hasattr(doc.paragraphs[i + j].style, 'name') else "None"
                        content_paras.append(f"  [{content_style}] {doc.paragraphs[i + j].text[:150]}...")
                
                print(f"Content paragraphs: {len(content_paras)}")
                for para in content_paras:
                    print(para)
                break

    # Print document tables
    print("\n=== TABLES ===")
    print(f"Total tables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables[:2]):  # Show first 2 tables
        print(f"\nTable {i}: Rows: {len(table.rows)}, Columns: {len(table.rows[0].cells) if table.rows else 0}")
        
        # Show table headers
        if table.rows:
            headers = []
            for cell in table.rows[0].cells:
                headers.append(cell.text)
            print(f"Headers: {headers}")
            
            # Show first data row if available
            if len(table.rows) > 1:
                data = []
                for cell in table.rows[1].cells:
                    data.append(cell.text)
                print(f"First data row: {data}")

if __name__ == "__main__":
    examine_document("output_populated_template.docx")