#!/usr/bin/env python3
"""
Check the Assay Procedure Summary section in the source document.

This script examines the source document to locate and extract the 
ASSAY PROCEDURE SUMMARY section for Red Dot documents.
"""

import logging
from pathlib import Path
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def find_assay_procedure_summary(document_path):
    """
    Find the ASSAY PROCEDURE SUMMARY section in a document.

    Args:
        document_path: Path to the document to inspect
    
    Returns:
        Text content of the ASSAY PROCEDURE SUMMARY section if found, None otherwise
    """
    # Load the document
    doc = Document(document_path)
    
    # Flags to track section boundaries
    in_assay_procedure_summary = False
    start_idx = -1
    end_idx = -1
    
    # First pass: identify section boundaries
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().upper()
        
        # Check for section start - use a more flexible match
        if (("ASSAY PROCEDURE SUMMARY" in text or "ASSAY SUMMARY" in text) and 
            not in_assay_procedure_summary and 
            len(text) < 60):  # Ensure it's a heading, not content
            in_assay_procedure_summary = True
            start_idx = i + 1  # Start after the section heading
            logger.info(f"Found ASSAY PROCEDURE SUMMARY section at paragraph {i}: {para.text}")
        
        # Check for section end (next section starts)
        elif in_assay_procedure_summary and text and any(keyword in text for keyword in [
            "IMPORTANT NOTE", "PRECAUTION", "SENSITIVITY", "DETECTION", "PRECISION", 
            "STABILITY", "DISCLAIMER"
        ]) and len(text) < 60:  # Likely a new section header (short line)
            end_idx = i - 1  # End before the next section
            logger.info(f"Found end of ASSAY PROCEDURE SUMMARY at paragraph {i-1}")
            break
    
    # If we found the start but not the end, assume it continues to the end of the document
    if in_assay_procedure_summary and start_idx > 0 and end_idx == -1:
        end_idx = len(doc.paragraphs) - 1
        logger.info(f"ASSAY PROCEDURE SUMMARY continues to the end of document at paragraph {end_idx}")
    
    # Extract section content
    if start_idx > 0 and end_idx >= start_idx:
        # Join the paragraphs
        content = []
        for i in range(start_idx, end_idx + 1):
            if doc.paragraphs[i].text.strip():  # Skip empty paragraphs
                content.append(doc.paragraphs[i].text.strip())
        
        # Look for any tables between these paragraphs
        table_content = []
        for i, table in enumerate(doc.tables):
            # Locate the table's position (approximate)
            for p_idx in range(len(doc.paragraphs)):
                if p_idx >= start_idx and p_idx <= end_idx:
                    if p_idx < len(doc.paragraphs) - 1 and doc.paragraphs[p_idx].text.strip() == "" and doc.paragraphs[p_idx+1].text.strip() == "":
                        # This could be a table position (empty paragraphs before/after)
                        # We don't have a direct way to locate tables, so this is a heuristic
                        table_content.append(f"[Table {i+1}]")
                        logger.info(f"Possible table {i} in ASSAY PROCEDURE SUMMARY section")
                        break
        
        # Combine paragraph and table content
        if table_content:
            content.extend(table_content)
        
        # Return the section content
        if content:
            full_content = "\n".join(content)
            logger.info(f"Extracted ASSAY PROCEDURE SUMMARY: {full_content}")
            return full_content
        else:
            logger.warning("ASSAY PROCEDURE SUMMARY section found but no content extracted")
            return None
    else:
        # Alternative detection method - look for short steps in the ASSAY PROCEDURE section
        logger.info("ASSAY PROCEDURE SUMMARY section not found, trying to extract from ASSAY PROCEDURE section")
        
        # Find ASSAY PROCEDURE section
        in_assay_procedure = False
        assay_procedure_start = -1
        assay_procedure_end = -1
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().upper()
            
            # Check for ASSAY PROCEDURE section start
            if (("ASSAY PROCEDURE" in text or "ASSAY PROTOCOL" in text) and 
                not in_assay_procedure and 
                len(text) < 60):
                in_assay_procedure = True
                assay_procedure_start = i + 1
                logger.info(f"Found ASSAY PROCEDURE section at paragraph {i}: {para.text}")
                
            # Check for section end
            elif in_assay_procedure and text and any(keyword in text for keyword in [
                "CALCULATION", "RESULT", "TYPICAL", "DETECTION", "SENSITIVITY"
            ]) and len(text) < 60:
                assay_procedure_end = i - 1
                logger.info(f"Found end of ASSAY PROCEDURE at paragraph {i-1}")
                break
        
        # If we found ASSAY PROCEDURE section, extract numbered points or short steps
        if in_assay_procedure and assay_procedure_start > 0:
            if assay_procedure_end == -1:  # If end not found, use some reasonable limit
                assay_procedure_end = min(assay_procedure_start + 20, len(doc.paragraphs) - 1)
                
            # Extract numbered steps or short bullet points
            import re
            steps = []
            
            for i in range(assay_procedure_start, assay_procedure_end + 1):
                text = doc.paragraphs[i].text.strip()
                
                # Skip empty paragraphs
                if not text:
                    continue
                    
                # Look for numbered steps or bullet points
                if re.match(r'^\d+\.', text) or text.startswith('•') or text.startswith('-'):
                    # Only include short steps in the summary (likely action items)
                    if len(text) < 200:
                        steps.append(text)
            
            # If we found steps, format them as a summary
            if steps:
                summary = "\n".join(steps)
                logger.info(f"Created ASSAY PROCEDURE SUMMARY from procedure steps: {summary}")
                return summary
        
        # If all else fails, check for a structured "Assay Summary" section
        # This is a last resort as some documents have this information in a different format
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if ("Assay Summary" in text or "Summary of Procedure" in text or 
                "Protocol Summary" in text or "Procedure at a Glance" in text):
                summary_text = []
                # Collect the next few paragraphs (likely the summary)
                for j in range(i+1, min(i+10, len(doc.paragraphs))):
                    next_text = doc.paragraphs[j].text.strip()
                    if next_text and len(next_text) < 200:  # Only include short paragraphs
                        summary_text.append(next_text)
                    # Stop if we hit another heading
                    if next_text and next_text.isupper() and len(next_text) < 50:
                        break
                
                if summary_text:
                    found_summary = "\n".join(summary_text)
                    logger.info(f"Found Assay Summary alternative section: {found_summary}")
                    return found_summary
                
        logger.warning("Could not find or create ASSAY PROCEDURE SUMMARY")
        return None

if __name__ == "__main__":
    # Search in the Red Dot document
    summary = find_assay_procedure_summary("attached_assets/RDR-LMNB2-Hu.docx")
    if summary:
        print("\nFULL ASSAY PROCEDURE SUMMARY:")
        print("-" * 40)
        print(summary)
        print("-" * 40)