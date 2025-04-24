#!/usr/bin/env python3
"""
Check the actual kit contents in the source document
"""

import re
import logging
from docx import Document
from pathlib import Path

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def check_kit_content_tables(source_path='attached_assets/EK1586_Mouse_KLK1Kallikrein_1_ELISA_Kit_PicoKine_Datasheet.docx'):
    """Check all tables in the source document to find the actual kit contents table."""
    source_path = Path(source_path)
    
    if not source_path.exists():
        logger.error(f"Source file does not exist: {source_path}")
        return
    
    # Parse the ELISA datasheet
    doc = Document(source_path)
    logger.info(f"Examining document: {source_path}")
    
    # Find the kit components section
    kit_section_idx = None
    for i, para in enumerate(doc.paragraphs):
        if any(term in para.text.lower() for term in ['kit components', 'materials provided']):
            logger.info(f"Found kit components section at paragraph {i}: {para.text}")
            kit_section_idx = i
            break
            
    if kit_section_idx is None:
        logger.warning("Kit components section not found")
        return
        
    # Look at all tables in the document
    for i, table in enumerate(doc.tables):
        logger.info(f"Table {i+1}: {len(table.rows)} rows x {len(table.columns)} columns")
        
        # Try to determine what type of table this is
        has_reagent_keywords = False
        has_specs_keywords = False
        
        for row in table.rows:
            row_text = " ".join([cell.text.lower() for cell in row.cells])
            
            # Check for reagent keywords
            if any(keyword in row_text for keyword in ['microplate', 'antibody', 'standard', 'buffer', 'substrate', 'stop solution']):
                has_reagent_keywords = True
                
            # Check for specs keywords
            if any(keyword in row_text for keyword in ['sensitivity', 'specificity', 'storage', 'reactive species']):
                has_specs_keywords = True
        
        # Print the full table content
        print(f"\nTable {i+1} Content:")
        print("-" * 50)
        if has_reagent_keywords:
            print("LIKELY REAGENT TABLE")
        if has_specs_keywords:
            print("LIKELY SPECS TABLE")
            
        for row_idx, row in enumerate(table.rows):
            row_content = [cell.text.strip() for cell in row.cells]
            print(f"Row {row_idx+1}: {row_content}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        check_kit_content_tables(sys.argv[1])
    else:
        check_kit_content_tables()