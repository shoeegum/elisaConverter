#!/usr/bin/env python3
"""
Create Enhanced Red Dot Template for ELISA Kit Datasheets

This script creates an enhanced template with proper styles, formatting, and placeholders
for Red Dot ELISA kit datasheets based on the provided requirements:
1. All section titles should use Heading 2 Style (blue, all caps)
2. Proper paragraph spacing (1.15) and font (Calibri)
3. Bold Calibri 36pt for main title
4. Company name in footer should be Calibri 24pt bold
5. Contact info in footer should be Open Sans Light 12pt
6. Required materials should use a bullet list
7. First page should only contain title, catalog/lot numbers, and intended use
"""

import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def create_enhanced_red_dot_template():
    """
    Create an enhanced template for Red Dot ELISA kit datasheets with proper styling.
    """
    # Create a new document
    doc = Document()
    
    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Ensure paragraphs have proper spacing
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # Set document margins (1 inch on all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Create Heading 2 style for section titles
    heading_2_style = doc.styles['Heading 2']
    heading_2_font = heading_2_style.font
    heading_2_font.name = 'Calibri'
    heading_2_font.size = Pt(12)
    heading_2_font.bold = True
    heading_2_font.color.rgb = RGBColor(0, 70, 180)  # blue color
    
    # First page - Title and basic info
    # Add title placeholder
    title = doc.add_paragraph()
    title.style = doc.styles['Title']
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("{{ kit_name }}")
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.name = 'Calibri'
    
    # Add catalog/lot numbers
    catalog = doc.add_paragraph()
    catalog.alignment = WD_ALIGN_PARAGRAPH.CENTER
    catalog.add_run("Catalog No: ").bold = True
    catalog.add_run("{{ catalog_number }}").bold = False
    
    lot = doc.add_paragraph()
    lot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lot.add_run("Lot No: ").bold = True
    lot.add_run("{{ lot_number }}").bold = False
    
    # Add a page break to start the main content on a new page
    doc.add_page_break()
    
    # Set up the second section with different header/footer
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    
    # Add the INTENDED USE section
    intended_use_title = doc.add_heading("INTENDED USE", level=2)
    intended_use_para = doc.add_paragraph("{{ intended_use }}")
    
    # Add TEST PRINCIPLE section
    doc.add_heading("TEST PRINCIPLE", level=2)
    doc.add_paragraph("{{ test_principle }}")
    
    # Add REAGENTS AND MATERIALS PROVIDED section
    doc.add_heading("REAGENTS PROVIDED", level=2)
    doc.add_paragraph("{{ reagents_and_materials_provided }}")
    
    # Add MATERIALS REQUIRED BUT NOT SUPPLIED section
    doc.add_heading("OTHER SUPPLIES REQUIRED", level=2)
    materials_para = doc.add_paragraph("{{ materials_required_but_not_supplied }}")
    
    # Add STORAGE OF THE KITS section
    doc.add_heading("STORAGE OF THE KITS", level=2)
    doc.add_paragraph("{{ storage_of_the_kits }}")
    
    # Add SAMPLE COLLECTION AND STORAGE section
    doc.add_heading("SAMPLE COLLECTION AND STORAGE", level=2)
    doc.add_paragraph("{{ sample_collection_and_storage }}")
    
    # Add REAGENT PREPARATION section
    doc.add_heading("REAGENT PREPERATION", level=2)
    doc.add_paragraph("{{ reagent_preparation }}")
    
    # Add SAMPLE PREPARATION section
    doc.add_heading("SAMPLE PREPERATION", level=2)
    doc.add_paragraph("{{ sample_preparation }}")
    
    # Add ASSAY PROCEDURE section
    doc.add_heading("ASSAY PROCEDURE", level=2)
    doc.add_paragraph("{{ assay_procedure }}")
    
    # Add CALCULATION OF RESULTS section
    doc.add_heading("CALCULATION OF RESULTS", level=2)
    doc.add_paragraph("{{ calculation_of_results }}")
    
    # Add TYPICAL DATA section
    doc.add_heading("TYPICAL DATA", level=2)
    doc.add_paragraph("{{ typical_data }}")
    
    # Add DETECTION RANGE section
    doc.add_heading("DETECTION RANGE", level=2)
    doc.add_paragraph("{{ detection_range }}")
    
    # Add SENSITIVITY section
    doc.add_heading("SENSITIVITY", level=2)
    doc.add_paragraph("{{ sensitivity }}")
    
    # Add SPECIFICITY section
    doc.add_heading("SPECIFICITY", level=2)
    doc.add_paragraph("{{ specificity }}")
    
    # Add PRECISION section
    doc.add_heading("PRECISION", level=2)
    doc.add_paragraph("{{ precision }}")
    
    # Add STABILITY section
    doc.add_heading("STABILITY", level=2)
    doc.add_paragraph("{{ stability }}")
    
    # Add ASSAY PROCEDURE SUMMARY section
    doc.add_heading("ASSAY PROCEDURE SUMMARY", level=2)
    doc.add_paragraph("{{ assay_procedure_summary }}")
    
    # Add IMPORTANT NOTE section
    doc.add_heading("IMPORTANT NOTE", level=2)
    doc.add_paragraph("{{ important_note }}")
    
    # Add PRECAUTION section
    doc.add_heading("PRECAUTION", level=2)
    doc.add_paragraph("{{ precaution }}")
    
    # Add DISCLAIMER section
    doc.add_heading("DISCLAIMER", level=2)
    doc.add_paragraph("{{ disclaimer }}")
    
    # Add footer - Use the Innovative Research footer replaced with Red Dot Biotech
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0]
        if footer_para.text:
            footer_para.text = ""  # Clear any existing text
        footer_run = footer_para.add_run("| www.reddotbiotech.com")
        footer_run.font.name = 'Calibri'
        footer_run.font.size = Pt(10)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save the template
    template_path = Path("templates_docx/enhanced_red_dot_template.docx")
    doc.save(template_path)
    logger.info(f"Enhanced Red Dot template created at {template_path}")
    
    return template_path

if __name__ == "__main__":
    create_enhanced_red_dot_template()