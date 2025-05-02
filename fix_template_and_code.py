#!/usr/bin/env python3
"""
Fix Template and Code for REAGENTS PROVIDED Table Placement

This script does two things:
1. Updates the enhanced Red Dot template to use a better placeholder for the REAGENTS PROVIDED table
2. Updates the red_dot_template_populator.py file to handle table placement correctly
"""

import logging
import shutil
import re
from pathlib import Path
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def update_enhanced_template(template_path):
    """
    Update the enhanced Red Dot template with better placeholders for the table.
    
    Args:
        template_path: Path to the template to update
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Create a backup of the template
        template_path = Path(template_path)
        backup_path = template_path.with_name(f"{template_path.stem}_before_table_update{template_path.suffix}")
        shutil.copy2(template_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Load the template
        doc = Document(template_path)
        
        # Replace all instances of Reddot with Innovative Research
        replacements = [
            ('Reddot Biotech INC.', 'Innovative Research, Inc.'),
            ('Reddot Biotech', 'Innovative Research'),
        ]
        
        count = 0
        for para in doc.paragraphs:
            original_text = para.text
            new_text = original_text
            for old_text, new_text in replacements:
                if old_text in new_text:
                    new_text = new_text.replace(old_text, new_text)
            
            if new_text != original_text:
                para.text = new_text
                count += 1
        
        logger.info(f"Replaced {count} instances of company names")
        
        # Find the REAGENTS PROVIDED section
        reagents_section_found = False
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == "REAGENTS PROVIDED":
                reagents_section_found = True
                logger.info(f"Found REAGENTS PROVIDED section at paragraph {i}")
                
                # Check if there's a paragraph after this one
                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1]
                    logger.info(f"Paragraph after REAGENTS PROVIDED: '{next_para.text}'")
                    
                    # Replace the placeholder with a clear indicator for a table
                    if "{{" in next_para.text and "}}" in next_para.text:
                        # Clear the paragraph and add our structured placeholder for the table
                        next_para.text = "{{ reagents_table_content }}"
                        logger.info("Updated placeholder for reagents table")
                break
        
        if not reagents_section_found:
            logger.error("REAGENTS PROVIDED section not found in the template")
            return False
            
        # Save the updated template
        doc.save(template_path)
        logger.info(f"Successfully updated template: {template_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error updating template: {e}")
        return False

def update_template_populator_code(file_path):
    """
    Update the red_dot_template_populator.py file to handle table placement correctly.
    
    Args:
        file_path: Path to the red_dot_template_populator.py file
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Create a backup of the file
        file_path = Path(file_path)
        backup_path = file_path.with_name(f"{file_path.stem}_backup{file_path.suffix}")
        shutil.copy2(file_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Read the file content
        with open(file_path, 'r') as f:
            content = f.read()
        
        # Update company names in the code
        content = content.replace('Reddot Biotech INC.', 'Innovative Research, Inc.')
        content = content.replace('Reddot Biotech', 'Innovative Research')
        
        # Find the section where reagents table is processed
        populator_section_pattern = r'# Special handling for REAGENTS PROVIDED section.*?reagents_table_data'
        
        # Add new code for handling table placement
        new_table_handling_code = """
        # Special handling for reagents_table_content to create proper Word table
        if reagents_table:
            # Format a special XML representation of the table for direct insertion
            context['reagents_table_content'] = "TABLE WILL BE INSERTED HERE"
            
            # Replace any company name references
            for key in list(context.keys()):
                if isinstance(context[key], str):
                    context[key] = context[key].replace('Reddot Biotech INC.', 'Innovative Research, Inc.')
                    context[key] = context[key].replace('Reddot Biotech', 'Innovative Research')
        """
        
        # If the pattern exists, replace it; otherwise, add it near other special handling code
        if re.search(populator_section_pattern, content, re.DOTALL):
            content = re.sub(populator_section_pattern, new_table_handling_code, content, flags=re.DOTALL)
        else:
            # Find a suitable insertion point (after other special handling sections)
            insertion_point = content.find("# Special handling for sections that need custom processing")
            if insertion_point > 0:
                insertion_point = content.find("\n", insertion_point) + 1
                content = content[:insertion_point] + new_table_handling_code + content[insertion_point:]
            else:
                # If we can't find a specific insertion point, add it near the end of the populate function
                insertion_point = content.find("# Save populated template")
                if insertion_point > 0:
                    content = content[:insertion_point] + new_table_handling_code + "\n        " + content[insertion_point:]
        
        # Also update the post-processing step for the table
        post_processing_pattern = r'# Apply post-processing to convert text tables to proper Word tables.*?convert_text_to_table\(output_path\)'
        
        # New post-processing code with better error handling and debugging
        new_post_processing_code = """
            # Apply post-processing to convert text tables to proper Word tables
            try:
                from fix_reagents_table_post_processing import convert_text_to_table
                if convert_text_to_table(output_path):
                    logger.info(f"Successfully converted REAGENTS PROVIDED to proper table in: {output_path}")
                else:
                    logger.warning("Could not convert REAGENTS PROVIDED to table, using text format")
                    
                # Also apply fix for company names
                import fix_red_dot_company_and_placement
                fix_red_dot_company_and_placement.fix_company_names(output_path)
                logger.info(f"Applied company name replacements to: {output_path}")
            except Exception as table_error:
                logger.error(f"Error in post-processing: {table_error}")
        """
        
        # If the pattern exists, replace it; otherwise, add it near the footer modification code
        if re.search(post_processing_pattern, content, re.DOTALL):
            content = re.sub(post_processing_pattern, new_post_processing_code, content, flags=re.DOTALL)
        else:
            # Find a suitable insertion point (after footer modification)
            insertion_point = content.find("# Apply the Red Dot footer")
            if insertion_point > 0:
                insertion_point = content.find("\n            except Exception as footer_error:", insertion_point)
                if insertion_point > 0:
                    insertion_point = content.find("\n", insertion_point + 1) + 1
                    content = content[:insertion_point] + new_post_processing_code + content[insertion_point:]
        
        # Write the updated content back to the file
        with open(file_path, 'w') as f:
            f.write(content)
            
        logger.info(f"Successfully updated code in: {file_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error updating code: {e}")
        return False

if __name__ == "__main__":
    # Update the enhanced Red Dot template
    template_path = "templates_docx/enhanced_red_dot_template.docx"
    update_enhanced_template(template_path)
    
    # Update the template populator code
    populator_path = "red_dot_template_populator.py"
    update_template_populator_code(populator_path)