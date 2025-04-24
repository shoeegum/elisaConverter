#!/usr/bin/env python3
"""
Check if the TECHNICAL DETAILS, OVERVIEW, and REPRODUCIBILITY tables are properly populated
in the output document. This verifies that our enhanced template populator is correctly
filling in the tables with extracted data.
"""

import os
from pathlib import Path
from docx import Document
import logging

# Set up logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def check_tables(document_path="output_populated_template.docx"):
    """
    Check if the tables in the document are properly populated.
    
    Args:
        document_path: Path to the document to check
        
    Note:
        Table indices in the document (in order):
        0 - Technical Details Table
        1 - Overview Table 
        2 - Kit Components Table
        3 - Standard Curve Table
        4 - Intra-Assay Table
        5 - Inter-Assay Table
        6 - Lot-to-Lot Table
    """
    if not os.path.exists(document_path):
        logger.error(f"Document not found at {document_path}")
        return
    
    logger.info(f"Checking tables in {document_path}")
    doc = Document(document_path)
    
    # Variables to track our findings
    found_technical_details_table = False
    found_overview_table = False
    found_reproducibility_tables = False
    technical_details_populated = False
    overview_populated = False
    reproducibility_populated = False
    
    # Find sections by heading
    technical_details_section = None
    overview_section = None
    reproducibility_section = None
    
    for i, para in enumerate(doc.paragraphs):
        if "TECHNICAL DETAILS" in para.text.strip().upper():
            technical_details_section = i
            logger.info(f"Found TECHNICAL DETAILS section at paragraph {i}")
        elif "OVERVIEW" in para.text.strip().upper():
            overview_section = i
            logger.info(f"Found OVERVIEW section at paragraph {i}")
        elif "REPRODUCIBILITY" in para.text.strip().upper():
            reproducibility_section = i
            logger.info(f"Found REPRODUCIBILITY section at paragraph {i}")
    
    # Check all tables to identify and validate our target tables
    # Print basic info about all tables first
    print("\n--- Tables in Document ---")
    for i, table in enumerate(doc.tables):
        if not table.rows:
            print(f"Table {i}: Empty table")
            continue
        
        row_sample = " ".join([cell.text for cell in table.rows[0].cells])
        print(f"Table {i}: {len(table.rows)} rows x {len(table.rows[0].cells)} cols - First row: {row_sample[:50]}...")
    
    for i, table in enumerate(doc.tables):
        if not table.rows:
            continue
        
        # Extract table content for analysis
        table_content = ""
        empty_cells = 0
        total_cells = 0
        
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                table_content += cell_text + " "
                total_cells += 1
                if not cell_text or cell_text == "N/A":
                    empty_cells += 1
        
        table_content = table_content.lower()
        
        # Look for key terms in the table to identify it
        contains_capture = 'capture' in table_content or 'antibod' in table_content
        contains_sensitivity = 'sensitivity' in table_content
        contains_detection_range = 'detection range' in table_content or 'range' in table_content
        contains_product = 'product' in table_content and ('name' in table_content)
        contains_species = 'species' in table_content or 'reactive' in table_content
        contains_reproducibility = 'cv' in table_content or 'intra-assay' in table_content or 'inter-assay' in table_content
        
        logger.info(f"Table {i} content keywords: " +
                  f"capture={contains_capture}, " +
                  f"sensitivity={contains_sensitivity}, " +
                  f"detection_range={contains_detection_range}, " +
                  f"product={contains_product}, " +
                  f"species={contains_species}, " +
                  f"reproducibility={contains_reproducibility}")
        
        # Look for the sample type table first (added by our fix_sample_sections function)
        if "sample type" in table_content and "collection and handling" in table_content:
            logger.info(f"Skipping sample type table at index {i}")
            continue
        
        # Check for technical details table
        # - It's usually the first non-sample table
        # - It contains terms like sensitivity, detection, capture, etc.
        if (contains_capture or contains_sensitivity or contains_detection_range) and not found_technical_details_table:
            found_technical_details_table = True
            logger.info(f"Found technical details table at index {i}")
            
            # Check if values are filled in
            empty_value_cells = 0
            for row in table.rows:
                if len(row.cells) >= 2:
                    value_cell = row.cells[1].text.strip()
                    if not value_cell or value_cell == "N/A":
                        empty_value_cells += 1
            
            # Calculate percentage of empty cells
            if len(table.rows) > 0:
                empty_percentage = (empty_value_cells / len(table.rows)) * 100
                logger.info(f"Technical details table has {empty_percentage:.1f}% empty value cells")
                
                if empty_percentage < 50:  # Less than 50% empty is considered populated
                    technical_details_populated = True
                    logger.info("Technical details table is adequately populated")
                else:
                    logger.warning("Technical details table has too many empty cells")
        
        # Check for overview table
        # - It usually contains terms like species, reactivity, etc.
        # - It's often after the technical details table
        elif (contains_product or contains_species) and not found_overview_table:
            found_overview_table = True
            logger.info(f"Found overview table at index {i}")
            
            # Check if values are filled in
            empty_value_cells = 0
            for row in table.rows:
                if len(row.cells) >= 2:
                    value_cell = row.cells[1].text.strip()
                    if not value_cell or value_cell == "N/A":
                        empty_value_cells += 1
            
            # Calculate percentage of empty cells
            if len(table.rows) > 0:
                empty_percentage = (empty_value_cells / len(table.rows)) * 100
                logger.info(f"Overview table has {empty_percentage:.1f}% empty value cells")
                
                if empty_percentage < 50:  # Less than 50% empty is considered populated
                    overview_populated = True
                    logger.info("Overview table is adequately populated")
                else:
                    logger.warning("Overview table has too many empty cells")
        
        # Check for reproducibility tables
        # - These are often near the end of the document
        # - They contain terms like CV, precision, intra-assay, inter-assay, etc.
        elif contains_reproducibility or "sample" in table_content:
            # Identify if this is an intra-assay, inter-assay, or lot-to-lot table
            is_intra_assay = "intra" in table_content.lower()
            is_inter_assay = "inter" in table_content.lower()
            is_lot_to_lot = "lot" in table_content.lower() and len(table.rows[0].cells) > 5
            
            if is_intra_assay or is_inter_assay or is_lot_to_lot:
                found_reproducibility_tables = True
                
                table_type = "intra-assay" if is_intra_assay else "inter-assay" if is_inter_assay else "lot-to-lot"
                logger.info(f"Found {table_type} reproducibility table at index {i}")
                
                # Check if values are filled in for all rows
                empty_cells_in_table = 0
                cell_count = 0
                
                for row_idx, row in enumerate(table.rows):
                    if row_idx == 0:  # Skip header row
                        continue
                        
                    for cell in row.cells:
                        cell_count += 1
                        if not cell.text.strip():
                            empty_cells_in_table += 1
                
                # Calculate percentage of empty cells
                if cell_count > 0:
                    empty_percentage = (empty_cells_in_table / cell_count) * 100
                    logger.info(f"Reproducibility table has {empty_percentage:.1f}% empty cells")
                    
                    if empty_percentage < 20:  # Less than 20% empty is considered populated
                        reproducibility_populated = True
                        logger.info(f"{table_type} table is adequately populated")
                    else:
                        logger.warning(f"{table_type} table has {empty_percentage:.1f}% empty cells")
    
    # Report overall status
    print("\n=== Table Population Status ===")
    print(f"TECHNICAL DETAILS Table: {'Found' if found_technical_details_table else 'Not found'}, {'Populated' if technical_details_populated else 'Not fully populated'}")
    print(f"OVERVIEW Table: {'Found' if found_overview_table else 'Not found'}, {'Populated' if overview_populated else 'Not fully populated'}")
    print(f"REPRODUCIBILITY Tables: {'Found' if found_reproducibility_tables else 'Not found'}, {'Populated' if reproducibility_populated else 'Not fully populated'}")
    
    if found_technical_details_table and found_overview_table and found_reproducibility_tables and technical_details_populated and overview_populated and reproducibility_populated:
        print("\n✅ All tables are found and properly populated!")
    else:
        print("\n❌ Some tables are missing or not fully populated.")
        
        if not found_technical_details_table:
            print("   - TECHNICAL DETAILS table not found")
        elif not technical_details_populated:
            print("   - TECHNICAL DETAILS table has empty cells")
            
        if not found_overview_table:
            print("   - OVERVIEW table not found")
        elif not overview_populated:
            print("   - OVERVIEW table has empty cells")
            
        if not found_reproducibility_tables:
            print("   - REPRODUCIBILITY tables not found")
        elif not reproducibility_populated:
            print("   - REPRODUCIBILITY tables have empty cells")

if __name__ == "__main__":
    check_tables()