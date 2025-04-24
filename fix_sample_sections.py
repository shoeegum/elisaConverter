#!/usr/bin/env python3
"""
Fix the Sample Preparation and Sample Dilution sections in the template.

This script:
1. Restructures the SAMPLE PREPARATION AND STORAGE section to include a proper table
2. Shortens the SAMPLE DILUTION GUIDELINE section
"""

import logging
from pathlib import Path
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def fix_sample_sections(document_path="output_populated_template.docx", 
                       output_path="output_fixed_sample_sections.docx"):
    """
    Fix the sample preparation and dilution sections in the document.
    
    Args:
        document_path: Path to the document to fix
        output_path: Path where the fixed document will be saved
    """
    # Create a new document using the template
    doc = Document(document_path)
    
    # Find the Sample Preparation and Sample Dilution sections
    sample_prep_idx = None
    sample_dilution_idx = None
    assay_procedure_idx = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "SAMPLE PREPARATION AND STORAGE" in text:
            sample_prep_idx = i
            logger.info(f"Found SAMPLE PREPARATION AND STORAGE at paragraph {i}")
        elif "SAMPLE DILUTION GUIDELINE" in text:
            sample_dilution_idx = i
            logger.info(f"Found SAMPLE DILUTION GUIDELINE at paragraph {i}")
        elif "ASSAY PROCEDURE" in text or "ASSAY PROTOCOL" in text:
            assay_procedure_idx = i
            logger.info(f"Found ASSAY PROCEDURE at paragraph {i}")
    
    # Fix Sample Preparation section by creating a new document with the desired structure
    new_doc = Document()
    
    # First, copy all paragraphs up to and including SAMPLE PREPARATION AND STORAGE
    for i in range(sample_prep_idx + 1):
        para = doc.paragraphs[i]
        new_para = new_doc.add_paragraph(para.text)
        new_para.style = para.style
    
    # Add sample preparation content
    new_doc.add_paragraph("These sample collection instructions and storage conditions are intended as a general guideline. Sample stability has not been evaluated.")
    new_doc.add_paragraph("")
    
    # Add SAMPLE COLLECTION NOTES
    sample_notes_para = new_doc.add_paragraph("SAMPLE COLLECTION NOTES")
    sample_notes_para.style = 'Heading 3'
    
    # Add collection notes content
    new_doc.add_paragraph("Innovative Research recommends that samples are used immediately upon preparation.")
    new_doc.add_paragraph("Avoid repeated freeze-thaw cycles for all samples.")
    new_doc.add_paragraph("Samples should be brought to room temperature (18-25°C) before performing the assay.")
    new_doc.add_paragraph("")
    
    # Add a table for sample types
    table = new_doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # Set the table header
    table.cell(0, 0).text = "Sample Type"
    table.cell(0, 1).text = "Collection and Handling"
    
    # Set the table content
    table.cell(1, 0).text = "Cell Culture Supernatant"
    table.cell(1, 1).text = "Centrifuge at 1000 × g for 10 minutes to remove insoluble particulates. Collect supernatant."
    
    table.cell(2, 0).text = "Serum"
    table.cell(2, 1).text = "Use a serum separator tube (SST). Allow samples to clot for 30 minutes before centrifugation for 15 minutes at approximately 1000 × g. Remove serum and assay immediately or store samples at -20°C."
    
    table.cell(3, 0).text = "Plasma"
    table.cell(3, 1).text = "Collect plasma using EDTA or heparin as an anticoagulant. Centrifuge samples for 15 minutes at 1000 × g within 30 minutes of collection. Store samples at -20°C."
    
    table.cell(4, 0).text = "Cell Lysates"
    table.cell(4, 1).text = "Collect cells and rinse with ice-cold PBS. Homogenize at 1×10^7/ml in PBS with a protease inhibitor cocktail. Freeze/thaw 3 times. Centrifuge at 10,000×g for 10 min at 4°C. Aliquot the supernatant for testing and store at -80°C."
    
    # Add Sample Dilution Guideline section
    dilution_para = new_doc.add_paragraph("SAMPLE DILUTION GUIDELINE")
    dilution_para.style = 'Heading 2'
    
    # Add dilution guideline content
    new_doc.add_paragraph("To inspect the validity of experimental operation and the appropriateness of sample dilution proportion, it is recommended to test all plates with the provided samples. Dilute the sample so the expected concentration falls near the middle of the standard curve range.")
    
    # Add all content from the ASSAY PROCEDURE section to the end
    if assay_procedure_idx:
        for i in range(assay_procedure_idx, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            new_para = new_doc.add_paragraph(para.text)
            new_para.style = para.style
    
    # Save the document with the fixed sections
    new_doc.save(output_path)
    logger.info(f"Fixed document saved to {output_path}")
    return True

def main():
    # Fix the sample sections
    success = fix_sample_sections()
    
    if not success:
        logger.error("Failed to fix sample sections")
        return
    
    # Verify the changes
    doc = Document("output_fixed_sample_sections.docx")
    
    # Find the sections and log their content
    sample_prep_idx = None
    sample_dilution_idx = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "SAMPLE PREPARATION AND STORAGE" in text:
            sample_prep_idx = i
            logger.info(f"Verified SAMPLE PREPARATION AND STORAGE at paragraph {i}")
        elif "SAMPLE DILUTION GUIDELINE" in text:
            sample_dilution_idx = i
            logger.info(f"Verified SAMPLE DILUTION GUIDELINE at paragraph {i}")
            
            # Print the dilution section content
            if i + 1 < len(doc.paragraphs):
                logger.info(f"Sample Dilution content: {doc.paragraphs[i+1].text[:100]}...")
    
    # Check for tables near the Sample Preparation section
    tables_found = 0
    for i, table in enumerate(doc.tables):
        # Simple check if this might be the sample preparation table
        if (len(table.rows) >= 5 and len(table.columns) >= 2 and
            'Sample Type' in table.cell(0, 0).text):
            tables_found += 1
            logger.info(f"Found sample preparation table (Table {i}) with {len(table.rows)} rows")
            logger.info(f"  Headers: {table.cell(0, 0).text}, {table.cell(0, 1).text}")
            logger.info(f"  First row: {table.cell(1, 0).text}, {table.cell(1, 1).text[:30]}...")
    
    if tables_found == 0:
        logger.warning("No sample preparation table found in the fixed document")
    else:
        logger.info(f"Found {tables_found} sample preparation tables")

if __name__ == "__main__":
    main()