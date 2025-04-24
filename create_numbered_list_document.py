#!/usr/bin/env python3
"""
Create a test ELISA Kit datasheet document that has numbered lists in the
Preparations Before Assay section for testing our enhanced document processing.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path

def create_test_document():
    """Create a test ELISA Kit datasheet with numbered lists."""
    doc = Document()
    
    # Add main properties and sections that we need
    doc.add_heading("Mouse KLK1 Kallikrein 1 ELISA Kit", 0)
    doc.add_paragraph("Catalog Number: EK1586")
    doc.add_paragraph("Lot Number: 123456")
    
    # Add Intended Use section
    doc.add_heading("INTENDED USE", 1)
    doc.add_paragraph("This Mouse KLK1 Kallikrein 1 ELISA Kit is for research use only. Not for diagnostic procedures.")
    
    # Add Background
    doc.add_heading("BACKGROUND", 1)
    doc.add_paragraph("Kallikrein-1 (KLK1) is a member of the kallikrein subfamily of serine proteases. Kallikreins are involved in post-translational processing of many polypeptides. Kallikrein 1 (KLK1), also known as tissue kallikrein, is one of the 15 known human kallikreins and has a key role in the cardiovascular system.")
    
    # Create Overview Table
    doc.add_heading("OVERVIEW", 1)
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    
    # Fill in the table with some sample data
    overview_specs = [
        ("Product Name", "Mouse KLK1 Kallikrein 1 ELISA Kit"),
        ("Reactive Species", "Mouse"),
        ("Size", "96T"),
        ("Description", "Mouse KLK1 Kallikrein 1 ELISA Kit for Serum, Plasma, Cell Culture Supernatants, Urine"),
        ("Sensitivity", "< 1 pg/ml"),
        ("Detection Range", "1.56-100 pg/ml"),
        ("Storage Instructions", "Store at 4°C for 6 months, -20°C for 12 months"),
        ("Uniprot ID", "P15947")
    ]
    
    for i, (key, value) in enumerate(overview_specs):
        row = table.rows[i].cells
        row[0].text = key
        row[1].text = value
    
    # Add some technical details
    doc.add_heading("TECHNICAL DETAILS", 1)
    tech_table = doc.add_table(rows=4, cols=2)
    tech_table.style = 'Table Grid'
    
    tech_specs = [
        ("Capture/Detection Antibodies", "Rabbit"),
        ("Specificity", "Natural and recombinant Mouse KLK1"),
        ("Standard Protein", "Recombinant Mouse KLK1"),
        ("Cross-reactivity", "No detectable cross-reactivity with other relevant proteins")
    ]
    
    for i, (key, value) in enumerate(tech_specs):
        row = tech_table.rows[i].cells
        row[0].text = key
        row[1].text = value
    
    # Add preparations before assay section with numbered lists
    doc.add_heading("PREPARATIONS BEFORE ASSAY", 1)
    
    # Add introduction paragraphs
    doc.add_paragraph("Please read the following instructions before starting the experiment. This section includes both regular paragraphs and numbered instructions.")
    doc.add_paragraph("Before beginning, make sure to thoroughly read the entire protocol.")
    
    # Add numbered lists - we'll use explicit numbering
    for i, step in enumerate([
        "Prepare all reagents, samples, and standards according to the instructions.",
        "Confirm that you have the appropriate non-supplied equipment available.",
        "Spin down all components to the bottom of the tube before opening.",
        "Don't let the 96-well plate dry out as this will inactivate active components.",
        "Don't reuse tips and tubes to avoid cross-contamination."
    ], 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {step}")
    
    # Add a few more regular paragraphs
    doc.add_paragraph("Avoid using reagents from different batches together.")
    doc.add_paragraph("The kit should not be used beyond the expiration date on the kit label.")
    
    # Add Kit Components section
    doc.add_heading("KIT COMPONENTS", 1)
    kit_table = doc.add_table(rows=5, cols=4)
    kit_table.style = 'Table Grid'
    
    # Add headers
    headers = ["Description", "Quantity", "Volume", "Storage"]
    for i, header in enumerate(headers):
        kit_table.rows[0].cells[i].text = header
    
    # Add some sample components
    components = [
        ("Anti-Mouse KLK1 Precoated 96-well strip microplate", "1", "96 wells", "4°C"),
        ("Mouse KLK1 Standard", "2", "10ng/tube", "-20°C"),
        ("Anti-Mouse KLK1 Detection Antibody", "1", "130 μL", "-20°C"),
        ("HRP-Streptavidin", "1", "130 μL", "4°C")
    ]
    
    for i, (desc, qty, vol, storage) in enumerate(components, 1):
        row = kit_table.rows[i].cells
        row[0].text = desc
        row[1].text = qty
        row[2].text = vol
        row[3].text = storage
    
    # Add a few more sections to make the document more complete
    doc.add_heading("MATERIALS REQUIRED BUT NOT PROVIDED", 1)
    doc.add_paragraph("Microplate reader capable of measuring absorbance at 450nm")
    doc.add_paragraph("Precision pipettes and pipette tips")
    doc.add_paragraph("Distilled or deionized water")
    doc.add_paragraph("Tubes for standard and sample dilution")
    
    # Save the document
    output_path = "attached_assets/elisa_test_with_numbered_lists.docx"
    doc.save(output_path)
    print(f"Test document created at: {output_path}")
    return output_path

if __name__ == "__main__":
    create_test_document()