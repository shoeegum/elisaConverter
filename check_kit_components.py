#!/usr/bin/env python3
"""
Check the Kit Components table in the output document.
"""

import logging
from docx import Document
from pathlib import Path

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def check_kit_components(output_path="output_populated_template.docx"):
    """Check the kit components table in the output document."""
    doc = Document(output_path)
    logger.info(f"Checking kit components table in document at {output_path}")
    
    # Find the kit components section in the document
    kit_components_section_idx = None
    kit_components_table_idx = None
    
    # Look for the Kit Components section header
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        if "kit components" in text or "materials provided" in text:
            logger.info(f"Found Kit Components section at paragraph {i}: {para.text}")
            kit_components_section_idx = i
            break
            
    # If we found the section, look for tables after it
    if kit_components_section_idx is not None:
        # Try to find the correct kit components table based on content
        correct_table_found = False
        for i, table in enumerate(doc.tables):
            # Print each table's details
            rows = len(table.rows)
            cols = len(table.columns)
            logger.info(f"Table {i+1}: {rows} rows x {cols} columns")
            
            if cols >= 3:  # Look for the multi-column table (should have Description, Quantity, Volume, Storage)
                # Check if the first row contains expected headers
                try:
                    header_row = [cell.text.strip().lower() for cell in table.rows[0].cells]
                    if len(header_row) >= 3 and any("description" in h or "component" in h for h in header_row):
                        logger.info(f"Found likely kit components table at index {i}")
                        kit_components_table_idx = i
                        correct_table_found = True
                        break
                except:
                    pass
        
            # As a fallback, check for reagent-related content in any table
            if not correct_table_found:
                # Sample the table to look for reagent keywords
                sample_content = ""
                for r in range(min(3, rows)):
                    for c in range(min(3, cols)):
                        try:
                            sample_content += table.rows[r].cells[c].text.lower() + " "
                        except:
                            pass
                
                # Look for common kit component terms
                if any(keyword in sample_content for keyword in 
                       ["microplate", "antibody", "standard", "buffer", "substrate", "diluent", "wash"]):
                    logger.info(f"Found likely kit components table at index {i} based on content")
                    kit_components_table_idx = i
                    break
        
        # If we haven't found a table yet, just use the first table
        if kit_components_table_idx is None and len(doc.tables) > 0:
            logger.warning("Using first table as kit components table")
            kit_components_table_idx = 0
    
    # If no section or table found, use the first table
    if kit_components_table_idx is None and len(doc.tables) > 0:
        kit_components_table_idx = 0
    
    if kit_components_table_idx is None:
        logger.error("No tables found in document")
        return
        
    # Examine all tables in the document
    logger.info(f"Examining all {len(doc.tables)} tables:")
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns)
        logger.info(f"Table {i+1}: {rows} rows x {cols} columns")
        
        # Print sample from table (first 2 rows)
        print(f"\nTable {i+1} Content Sample:")
        print("-" * 50)
        for row_idx in range(min(2, rows)):
            cells = [cell.text.strip() for cell in table.rows[row_idx].cells]
            print(f"Row {row_idx+1}: {cells}")
    
    # Now examine the selected kit components table
    kit_table = doc.tables[kit_components_table_idx]
    rows = len(kit_table.rows)
    cols = len(kit_table.columns)
    
    print(f"\nSelected Kit Components Table (Table {kit_components_table_idx+1}):")
    print("-" * 50)
    print(f"Dimensions: {rows} rows x {cols} columns")
    
    # Print the table contents
    for row_idx, row in enumerate(kit_table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"Row {row_idx+1}: {cells}")
        
    # Check for reagent rows with content
    filled_rows = 0
    for row_idx in range(1, rows): # Skip header row
        if row_idx < len(kit_table.rows):
            row = kit_table.rows[row_idx]
            if row.cells[0].text.strip():
                filled_rows += 1
                
    logger.info(f"Found {filled_rows} filled reagent rows")
    print(f"\nFilled reagent rows: {filled_rows}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        check_kit_components(sys.argv[1])
    else:
        check_kit_components()