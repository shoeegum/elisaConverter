#!/usr/bin/env python3
"""
Create an updated enhanced template to include missing sections:
- ASSAY PRINCIPLE
- SAMPLE PREPARATION AND STORAGE
- SAMPLE COLLECTION NOTES
- SAMPLE DILUTION GUIDELINE
- DATA ANALYSIS
"""

import logging
from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def create_heading(doc, text, level=1):
    """Create a heading with the specified text and level."""
    heading = doc.add_paragraph(text)
    heading.style = f'Heading {level}'
    
    # Apply custom formatting for all headings
    if level == 2:
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for run in heading.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 70, 180)
    elif level == 1:
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in heading.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(36)
            run.font.bold = True
    
    return heading

def create_paragraph(doc, text="", style="Normal"):
    """Create a paragraph with the specified text and style."""
    para = doc.add_paragraph(text)
    para.style = style
    
    # Apply consistent formatting
    for run in para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    
    return para

def create_enhanced_template():
    """
    Create an enhanced template with all required sections.
    """
    doc = Document()
    
    # Set the document margins (narrow)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Add title and catalog/lot number
    title = create_paragraph(doc, "{{ kit_name|default('ELISA Kit') }}")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.size = Pt(36)
        run.font.bold = True
    
    catalog_lot = create_paragraph(doc, "Catalog Number: {{ catalog_number|default('') }}\nLot Number: {{ lot_number|default('') }}")
    catalog_lot.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add sections with placeholders
    create_heading(doc, "INTENDED USE", level=2)
    create_paragraph(doc, "{{ intended_use|default('') }}")
    
    create_heading(doc, "TECHNICAL DETAILS", level=2)
    
    # Create technical details table (4 rows, 2 columns)
    tech_table = doc.add_table(rows=4, cols=2)
    tech_table.style = 'Table Grid'
    
    # Fill the technical details table with placeholders
    headers = ['Capture/Detection Antibodies', 'Specificity', 'Standard Protein', 'Cross-reactivity']
    for i, header in enumerate(headers):
        cell = tech_table.cell(i, 0)
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add more sections
    create_heading(doc, "OVERVIEW", level=2)
    
    # Create the overview specifications table (8 rows, 2 columns)
    overview_table = doc.add_table(rows=8, cols=2)
    overview_table.style = 'Table Grid'
    
    # Fill the overview specifications table with placeholders
    overview_headers = [
        'Product Name', 'Reactive Species', 'Range', 'Sensitivity', 
        'Sample Type', 'Sample Volume', 'Assay Time', 'Protocol'
    ]
    for i, header in enumerate(overview_headers):
        cell = overview_table.cell(i, 0)
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        
        # Add placeholder for values
        value_cell = overview_table.cell(i, 1)
        if header == 'Product Name':
            value_cell.text = "{{ kit_name|default('Mouse KLK1/Kallikrein 1 ELISA Kit') }}"
    
    create_heading(doc, "BACKGROUND", level=2)
    create_paragraph(doc, "{{ background_text|default('') }}")
    
    # New section for assay principle
    create_heading(doc, "ASSAY PRINCIPLE", level=2)
    create_paragraph(doc, "{{ assay_principle|default('') }}")
    
    create_heading(doc, "KIT COMPONENTS", level=2)
    
    # Create kit components table (8 rows, 4 columns)
    kit_table = doc.add_table(rows=12, cols=4)
    kit_table.style = 'Table Grid'
    
    # Fill the kit components table header
    kit_table.cell(0, 0).text = "Description"
    kit_table.cell(0, 1).text = "Quantity"
    kit_table.cell(0, 2).text = "Volume"
    kit_table.cell(0, 3).text = "Storage of opened/reconstituted material"
    
    # Bold the header row
    for i in range(4):
        cell = kit_table.cell(0, i)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    create_heading(doc, "MATERIALS REQUIRED BUT NOT PROVIDED", level=2)
    create_paragraph(doc, "{{ required_materials_with_bullets|default('') }}")
    
    create_heading(doc, "REAGENT PREPARATION", level=2)
    create_paragraph(doc, "{{ reagent_preparation|default('') }}")
    
    create_heading(doc, "DILUTION OF STANDARD", level=2)
    create_paragraph(doc, "{{ dilution_of_standard|default('') }}")
    
    create_heading(doc, "PREPARATIONS BEFORE ASSAY", level=2)
    
    # Add numbered list for preparations
    prep_list = [
        "Prepare all reagents, samples, and standards according to the instructions.",
        "Confirm that you have the appropriate non-supplied equipment available.",
        "Spin down all components to the bottom of the tube before opening.",
        "Don't let the 96-well plate dry out as this will inactivate active components.",
        "Don't reuse tips and tubes to avoid cross-contamination. Avoid using reagents from different batches."
    ]
    
    for i, item in enumerate(prep_list, 1):
        p = create_paragraph(doc, f"{i}. {item}")
    
    # New section for sample preparation
    create_heading(doc, "SAMPLE PREPARATION AND STORAGE", level=2)
    create_paragraph(doc, "{{ sample_preparation_and_storage|default('') }}")
    
    # New section for sample collection
    create_heading(doc, "SAMPLE COLLECTION NOTES", level=2)
    create_paragraph(doc, "{{ sample_collection_notes|default('') }}")
    
    # New section for sample dilution
    create_heading(doc, "SAMPLE DILUTION GUIDELINE", level=2)
    create_paragraph(doc, "{{ sample_dilution_guideline|default('') }}")
    
    create_heading(doc, "ASSAY PROTOCOL", level=2)
    create_paragraph(doc, "{{ assay_protocol_numbered|default('') }}")
    
    create_heading(doc, "TYPICAL DATA / STANDARD CURVE", level=2)
    create_paragraph(doc, "This standard curve is for demonstration only. A standard curve must be run with each assay.")
    
    # Create standard curve table (2 rows, 9 columns)
    std_table = doc.add_table(rows=2, cols=9)
    std_table.style = 'Table Grid'
    
    # Fill the standard curve table header
    std_table.cell(0, 0).text = "Concentration (pg/ml)"
    for i in range(1, 9):
        std_table.cell(0, i).text = "{{ standard_curve.concentrations[" + str(i-1) + "]|default('') }}"
        std_table.cell(1, i).text = "{{ standard_curve.od_values[" + str(i-1) + "]|default('') }}"
    
    std_table.cell(1, 0).text = "O.D."
    
    # Bold the header row and first column
    for i in range(9):
        cell = std_table.cell(0, i)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    for i in range(2):
        cell = std_table.cell(i, 0)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    create_heading(doc, "INTRA/INTER-ASSAY VARIABILITY", level=2)
    create_paragraph(doc, "Three samples of known concentration were tested on one plate to assess intra-assay precision.")
    
    # Create intra-assay table (4 rows, 5 columns)
    intra_table = doc.add_table(rows=4, cols=5)
    intra_table.style = 'Table Grid'
    
    # Fill the intra-assay table header
    intra_headers = ["Sample", "n", "Mean (pg/ml)", "Standard Deviation", "CV (%)"]
    for i, header in enumerate(intra_headers):
        intra_table.cell(0, i).text = header
        for paragraph in intra_table.cell(0, i).paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Fill the intra-assay table with placeholders
    for i in range(1, 4):
        intra_table.cell(i, 0).text = str(i)
        intra_table.cell(i, 1).text = "{{ variability.intra_assay.sample_" + str(i) + ".n|default('24') }}"
        intra_table.cell(i, 2).text = "{{ variability.intra_assay.sample_" + str(i) + ".mean|default('') }}"
        intra_table.cell(i, 3).text = "{{ variability.intra_assay.sample_" + str(i) + ".sd|default('') }}"
        intra_table.cell(i, 4).text = "{{ variability.intra_assay.sample_" + str(i) + ".cv|default('') }}"
    
    create_paragraph(doc, "Three samples of known concentration were tested in separate assays to assess inter-assay precision.")
    
    # Create inter-assay table (4 rows, 5 columns)
    inter_table = doc.add_table(rows=4, cols=5)
    inter_table.style = 'Table Grid'
    
    # Fill the inter-assay table header
    for i, header in enumerate(intra_headers):
        inter_table.cell(0, i).text = header
        for paragraph in inter_table.cell(0, i).paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Fill the inter-assay table with placeholders
    for i in range(1, 4):
        inter_table.cell(i, 0).text = str(i)
        inter_table.cell(i, 1).text = "{{ variability.inter_assay.sample_" + str(i) + ".n|default('24') }}"
        inter_table.cell(i, 2).text = "{{ variability.inter_assay.sample_" + str(i) + ".mean|default('') }}"
        inter_table.cell(i, 3).text = "{{ variability.inter_assay.sample_" + str(i) + ".sd|default('') }}"
        inter_table.cell(i, 4).text = "{{ variability.inter_assay.sample_" + str(i) + ".cv|default('') }}"
    
    create_heading(doc, "REPRODUCIBILITY", level=2)
    create_paragraph(doc, "Samples were tested in four different assay lots to assess reproducibility.")
    
    # Create reproducibility table (4 rows, 7 columns) with standard deviation
    repro_table = doc.add_table(rows=4, cols=7)
    repro_table.style = 'Table Grid'
    
    # Fill the reproducibility table header
    repro_headers = ["", "Lot 1", "Lot 2", "Lot 3", "Lot 4", "Standard Deviation", "Mean", "CV (%)"]
    for i, header in enumerate(repro_headers):
        if i < len(repro_table.columns):  # Make sure we don't go out of bounds
            repro_table.cell(0, i).text = header
            for paragraph in repro_table.cell(0, i).paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
    
    # Fill the reproducibility table with placeholders
    for i in range(1, 4):
        if i < len(repro_table.rows):  # Make sure we don't go out of bounds
            repro_table.cell(i, 0).text = "Sample " + str(i)
    
    # Add DATA ANALYSIS section
    create_heading(doc, "DATA ANALYSIS", level=2)
    create_paragraph(doc, "{{ data_analysis|default('') }}")
    
    # Save the template
    template_path = Path('templates_docx/enhanced_template_complete.docx')
    doc.save(template_path)
    logger.info(f"Created updated template at {template_path}")
    
    # Copy to the original template location
    import shutil
    destination_path = Path('templates_docx/enhanced_template.docx')
    shutil.copy(template_path, destination_path)
    logger.info(f"Replaced original template at {destination_path}")

if __name__ == "__main__":
    create_enhanced_template()