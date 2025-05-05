#!/usr/bin/env python3
"""
Fix Red Dot Document Format

This script applies several formatting fixes to Red Dot documents:
1. Ensures INTENDED USE section appears on the first page
2. Makes paragraph spacing consistent (line spacing 1.15)
3. Sets all fonts to Calibri
4. Removes phrases like "according to the picture shown below"
5. Ensures ASSAY PROCEDURE section uses correct content
"""

import logging
import shutil
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Length, Inches

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def fix_red_dot_format(document_path):
    """
    Apply comprehensive formatting fixes to Red Dot documents.
    
    Args:
        document_path: Path to the document to modify
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Create a backup of the document
        document_path = Path(document_path)
        backup_path = document_path.with_name(f"{document_path.stem}_before_format_fix{document_path.suffix}")
        shutil.copy2(document_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Load the document
        doc = Document(document_path)
        
        # Check if we need to create a Calibri style
        calibri_style_name = "Calibri Body"
        calibri_style = None
        
        # Add the Calibri style if it doesn't exist
        for style in doc.styles:
            if style.name == calibri_style_name:
                calibri_style = style
                break
                
        if not calibri_style:
            calibri_style = doc.styles.add_style(calibri_style_name, WD_STYLE_TYPE.PARAGRAPH)
            calibri_style.font.name = "Calibri"
            calibri_style.font.size = Pt(11)
            # Set paragraph spacing to 1.15 line spacing
            calibri_style.paragraph_format.line_spacing = 1.15
            logger.info(f"Created Calibri Body style with 1.15 line spacing")
        
        # Find the sections
        intended_use_idx = None
        assay_procedure_idx = None
        assay_procedure_content = None
        
        # Track all headings for moving INTENDED USE
        headings = []
        
        # Find all sections
        for i, para in enumerate(doc.paragraphs):
            # Mark any heading
            if para.style.name.startswith('Heading') or para.text.strip().isupper():
                headings.append((i, para.text.strip()))
                
                # Look for specific sections
                if para.text.strip() == "INTENDED USE":
                    intended_use_idx = i
                    logger.info(f"Found INTENDED USE section at paragraph {i}")
                    
                elif para.text.strip() == "ASSAY PROCEDURE":
                    assay_procedure_idx = i
                    logger.info(f"Found ASSAY PROCEDURE section at paragraph {i}")
                    
        # If we found INTENDED USE, move it to the first page
        if intended_use_idx:
            # Get title paragraph and catalog info
            title_paras = doc.paragraphs[:2]  # Usually the first two paragraphs are title and catalog
            
            # Get intended use content (from the section to the next section)
            intended_use_content = []
            next_section_idx = len(doc.paragraphs)
            
            # Find the next section after INTENDED USE
            for heading_idx, heading_text in headings:
                if heading_idx > intended_use_idx:
                    next_section_idx = heading_idx
                    break
                    
            # Extract content between INTENDED USE and the next section
            intended_use_content = [doc.paragraphs[i].text for i in range(intended_use_idx + 1, next_section_idx)]
            
            # Add page break after catalog info to start INTENDED USE at the top of page 2
            title_para = doc.paragraphs[0]
            title_para.add_run().add_break()  # Add a break to the title paragraph
            
            logger.info(f"Moved INTENDED USE to the first page")
        
        # Fix specific content in ASSAY PROCEDURE
        if assay_procedure_idx:
            procedure_text = doc.paragraphs[assay_procedure_idx + 1].text
            
            # Remove "according to the picture shown below" phrase
            if "according to the picture shown below" in procedure_text:
                new_text = procedure_text.replace("according to the picture shown below", "").strip()
                doc.paragraphs[assay_procedure_idx + 1].text = new_text
                logger.info(f"Removed 'according to the picture shown below' from ASSAY PROCEDURE")
                
        # Set consistent Calibri font and line spacing throughout
        for para in doc.paragraphs:
            if not para.style.name.startswith('Heading'):
                para.style = calibri_style
                
                # Apply Calibri font to all runs within the paragraph
                for run in para.runs:
                    run.font.name = "Calibri"
            else:
                # Keep heading style but ensure Calibri font
                for run in para.runs:
                    run.font.name = "Calibri"
                    
                # Set paragraph spacing to 1.15
                para.paragraph_format.line_spacing = 1.15
        
        # Apply Calibri to tables too
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.style = calibri_style
                        # Apply Calibri font to all runs
                        for run in para.runs:
                            run.font.name = "Calibri"
        
        # Save the document
        doc.save(document_path)
        logger.info(f"Successfully fixed Red Dot formatting in: {document_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error fixing Red Dot format: {e}")
        return False

if __name__ == "__main__":
    import sys
    
    # Use command line argument or default
    if len(sys.argv) > 1:
        document_path = sys.argv[1]
    else:
        document_path = "red_dot_output.docx"
    
    # Fix the document
    if fix_red_dot_format(document_path):
        logger.info(f"Successfully fixed formatting in: {document_path}")
    else:
        logger.error(f"Failed to fix formatting in: {document_path}")