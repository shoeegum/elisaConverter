#!/usr/bin/env python3
"""
Check the ASSAY PROCEDURE SUMMARY section in the output document.
"""

import logging
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO,
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def check_assay_summary(document_path):
    """
    Check the ASSAY PROCEDURE SUMMARY section in the output document.
    
    Args:
        document_path: Path to the document to check
    """
    # Load the document
    doc = Document(document_path)
    
    # Just identify each section and print its paragraph index
    logger.info("=== Document Structure ===")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() and para.text.strip().isupper() and len(para.text.strip()) < 50:
            logger.info(f"Section at P{i}: {para.text}")
    logger.info("========================")

    # Dictionary to store section content
    sections = {}
    current_section = None
    
    # First pass: identify sections and their content
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and text.isupper() and len(text) < 50:
            # This is likely a section heading
            current_section = text
            sections[current_section] = []
        elif current_section and text:
            # This is content for the current section
            sections[current_section].append(text)
    
    # Print the content of specific sections we're interested in
    target_sections = ["REAGENTS PROVIDED", "ASSAY PROCEDURE SUMMARY"]
    for section in target_sections:
        if section in sections:
            logger.info(f"\n=== {section} Content ===")
            for line in sections[section]:
                logger.info(line)
            logger.info("=" * (len(section) + 14))
        else:
            logger.warning(f"Section '{section}' not found in document")
    
    # Look specifically for tables in the document
    logger.info(f"\nDocument contains {len(doc.tables)} tables")
    for i, table in enumerate(doc.tables):
        logger.info(f"Table {i} has {len(table.rows)} rows and {len(table.rows[0].cells) if table.rows else 0} columns")
        if table.rows:
            for j, row in enumerate(table.rows):
                if j < 3:  # Print just the first few rows for inspection
                    row_text = " | ".join([cell.text for cell in row.cells])
                    logger.info(f"Row {j}: {row_text}")

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        check_assay_summary(sys.argv[1])
    else:
        check_assay_summary("improved_red_dot_output.docx")