#!/usr/bin/env python3
"""
Check the output document to verify that all sections and tables were populated correctly.
"""

import logging
from docx import Document
from pathlib import Path

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def check_output(output_path="output_populated_template.docx"):
    """Check the output document for properly populated content."""
    doc = Document(output_path)
    logger.info(f"Checking output document at {output_path}")
    
    # Check sections
    sections_found = {
        "kit_name": False,
        "catalog_number": False,
        "intended_use": False,
        "materials_required": False,
        "standard_curve": False,
        "variability": False,
        "reproducibility": False
    }
    
    # Check all paragraphs
    for i, para in enumerate(doc.paragraphs):
        # Check for kit name - should be in the first few paragraphs
        if i < 5 and len(para.text) > 10 and "KLK1" in para.text:
            sections_found["kit_name"] = True
            logger.info(f"Found kit name at paragraph {i}: {para.text}")
        
        # Check for catalog number
        if "EK1586" in para.text:
            sections_found["catalog_number"] = True
            logger.info(f"Found catalog number at paragraph {i}: {para.text}")
        
        # Check section headings
        if "INTENDED USE" in para.text:
            sections_found["intended_use"] = True
            logger.info(f"Found intended use section at paragraph {i}")
        elif "MATERIALS REQUIRED" in para.text:
            sections_found["materials_required"] = True
            logger.info(f"Found materials required section at paragraph {i}")
        elif "STANDARD CURVE" in para.text:
            sections_found["standard_curve"] = True
            logger.info(f"Found standard curve section at paragraph {i}")
        elif "VARIABILITY" in para.text:
            sections_found["variability"] = True
            logger.info(f"Found variability section at paragraph {i}")
        elif "REPRODUCIBILITY" in para.text:
            sections_found["reproducibility"] = True
            logger.info(f"Found reproducibility section at paragraph {i}")
    
    # Check tables in the document
    standard_curve_table = None
    intra_assay_table = None
    inter_assay_table = None
    reproducibility_table = None
    
    for i, table in enumerate(doc.tables):
        # Check table dimensions
        rows = len(table.rows)
        cols = len(table.columns)
        logger.info(f"Table {i+1}: {rows} rows x {cols} columns")
        
        # Examine cell content to determine table type
        if rows > 0 and cols > 0:
            header_cell = table.cell(0, 0).text
            
            if "Concentration" in header_cell and cols > 8:
                standard_curve_table = i
                logger.info(f"Found standard curve table at index {i} ({rows}x{cols})")
                logger.info(f"  Value examples: {table.cell(1, 1).text}, {table.cell(1, 2).text}")
            elif "Sample" in header_cell and cols >= 5:
                if intra_assay_table is None:
                    intra_assay_table = i
                    logger.info(f"Found intra-assay table at index {i} ({rows}x{cols})")
                    if rows > 1 and cols > 4:
                        logger.info(f"  Values: {table.cell(1, 2).text}, CV: {table.cell(1, 4).text}")
                else:
                    inter_assay_table = i
                    logger.info(f"Found inter-assay table at index {i} ({rows}x{cols})")
                    if rows > 1 and cols > 4:
                        logger.info(f"  Values: {table.cell(1, 2).text}, CV: {table.cell(1, 4).text}")
            elif cols >= 7 and rows >= 4:
                # Check if this might be the reproducibility table
                if "Lot 1" in table.cell(0, 1).text if cols > 1 else "":
                    reproducibility_table = i
                    logger.info(f"Found reproducibility table at index {i} ({rows}x{cols})")
                    if rows > 1 and cols > 5:
                        logger.info(f"  Values: {table.cell(1, 1).text}, CV: {table.cell(1, 6).text}")
    
    # Check materials section content
    materials_count = 0
    for i, para in enumerate(doc.paragraphs):
        if sections_found["materials_required"]:
            # Count bullet points after the materials section
            if "•" in para.text and len(para.text.strip()) > 3:
                materials_count += 1
                logger.info(f"Found material bullet point: {para.text}")
    
    logger.info(f"Found {materials_count} material bullet points")
    
    # Print summary of findings
    print("\nOutput Document Summary:")
    print(f"Kit name: {'Found' if sections_found['kit_name'] else 'Missing'}")
    print(f"Catalog number: {'Found' if sections_found['catalog_number'] else 'Missing'}")
    print(f"Intended use section: {'Found' if sections_found['intended_use'] else 'Missing'}")
    print(f"Materials required section: {'Found' if sections_found['materials_required'] else 'Missing'}")
    print(f"Material bullet points: {materials_count}")
    print(f"Standard curve section: {'Found' if sections_found['standard_curve'] else 'Missing'}")
    print(f"Standard curve table: {'Found' if standard_curve_table is not None else 'Missing'}")
    print(f"Variability section: {'Found' if sections_found['variability'] else 'Missing'}")
    print(f"Intra-assay table: {'Found' if intra_assay_table is not None else 'Missing'}")
    print(f"Inter-assay table: {'Found' if inter_assay_table is not None else 'Missing'}")
    print(f"Reproducibility section: {'Found' if sections_found['reproducibility'] else 'Missing'}")
    print(f"Reproducibility table: {'Found' if reproducibility_table is not None else 'Missing'}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        check_output(sys.argv[1])
    else:
        check_output()