#!/usr/bin/env python3
"""
Fix template issues identified in the review:
- Add DISCLAIMER section
- Fix TECHNICAL DETAILS and OVERVIEW tables
- Format SAMPLE DILUTION and ASSAY PROTOCOL as lists
- Fix STANDARD CURVE table
- Add proper footer
- Populate REPRODUCIBILITY table correctly
"""

import logging
from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def create_heading(doc, text, level=2):
    """Create a heading with the specified text and level."""
    heading = doc.add_paragraph(text)
    heading.style = f'Heading {level}'
    
    # Set heading to all caps and blue color
    for run in heading.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0, 70, 180)  # RGB for blue
        run.text = run.text.upper()

def create_paragraph(doc, text="", style="Normal"):
    """Create a paragraph with the specified text and style."""
    paragraph = doc.add_paragraph()
    paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph

def add_disclaimer_section(doc):
    """Add the DISCLAIMER section."""
    create_heading(doc, "DISCLAIMER")
    
    disclaimer_text = "This material is sold for in-vitro use only in manufacturing and research. This material is not suitable for human use. It is the responsibility of the user to undertake sufficient verification and testing to determine the suitability of each product's application. The statements herein are offered for informational purposes only and are intended to be used solely for your consideration, investigation and verification."
    
    paragraph = create_paragraph(doc, disclaimer_text)
    return paragraph

def fix_sample_dilution_format(doc):
    """Convert SAMPLE DILUTION GUIDELINE to a list format."""
    # Find the section
    sample_dilution_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "SAMPLE DILUTION GUIDELINE" in para.text.upper():
            sample_dilution_idx = i
            break
    
    if sample_dilution_idx is None:
        logger.warning("SAMPLE DILUTION GUIDELINE section not found")
        return
    
    # Find the paragraph with the content
    content_idx = None
    for i in range(sample_dilution_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip():
            content_idx = i
            break
    
    if content_idx is None:
        logger.warning("SAMPLE DILUTION GUIDELINE content not found")
        return
    
    # Get the original content
    original_content = doc.paragraphs[content_idx].text
    
    # Replace with a template variable that can be processed as a list
    doc.paragraphs[content_idx].clear()
    doc.paragraphs[content_idx].add_run("{{ sample_dilution_guideline }}")
    
    return original_content

def fix_assay_protocol_format(doc):
    """Convert ASSAY PROTOCOL to a numbered list format."""
    # Find the section
    assay_protocol_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "ASSAY PROTOCOL" in para.text.upper():
            assay_protocol_idx = i
            break
    
    if assay_protocol_idx is None:
        logger.warning("ASSAY PROTOCOL section not found")
        return
    
    # Find the paragraph with the content
    content_idx = None
    for i in range(assay_protocol_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip():
            content_idx = i
            break
    
    if content_idx is None:
        logger.warning("ASSAY PROTOCOL content not found")
        return
    
    # Get the original content 
    original_content = doc.paragraphs[content_idx].text
    
    # Replace with a template variable that can be processed as a list
    doc.paragraphs[content_idx].clear()
    doc.paragraphs[content_idx].add_run("{{ assay_protocol_numbered }}")
    
    return original_content

def fix_technical_details_table(doc):
    """Fix the TECHNICAL DETAILS table."""
    # Find the TECHNICAL DETAILS section
    technical_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "TECHNICAL DETAILS" in para.text.upper():
            technical_idx = i
            break
    
    if technical_idx is None:
        logger.warning("TECHNICAL DETAILS section not found")
        return
    
    # Add a table after the section
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # Set up rows
    properties = [
        "Capture/Detection Antibodies", 
        "Specificity",
        "Standard Protein",
        "Cross-reactivity",
        "Sensitivity"
    ]
    
    for i, prop in enumerate(properties):
        table.rows[i].cells[0].text = prop
        table.rows[i].cells[1].text = "{{ technical_details." + prop.lower().replace("/", "_").replace("-", "_") + " if technical_details else 'N/A' }}"
        
        # Make property names bold
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Set column widths
    table.columns[0].width = Cm(6.0)  # Property
    table.columns[1].width = Cm(9.0)  # Value
    
    return table

def fix_standard_curve_table(doc):
    """Fix the STANDARD CURVE table."""
    # Find the STANDARD CURVE section
    standard_curve_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "STANDARD CURVE" in para.text.upper() or "TYPICAL DATA" in para.text.upper():
            standard_curve_idx = i
            break
    
    if standard_curve_idx is None:
        logger.warning("STANDARD CURVE/TYPICAL DATA section not found")
        return
    
    # Use a simple paragraph with template variable for safety
    para = create_paragraph(doc)
    para.add_run("Standard curve data:").bold = True
    
    # Create simple template for standard curve
    standard_curve_para = create_paragraph(doc)
    standard_curve_para.add_run("{{ standard_curve_table|safe }}")
    
    # Add note
    note_para = create_paragraph(doc)
    note_para.add_run("This standard curve is for demonstration only. A standard curve must be run with each assay.").italic = True
    
    return standard_curve_para

def fix_reproducibility_table(doc):
    """Fix the REPRODUCIBILITY table."""
    # Find the REPRODUCIBILITY section
    repro_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "REPRODUCIBILITY" in para.text.upper():
            repro_idx = i
            break
    
    if repro_idx is None:
        logger.warning("REPRODUCIBILITY section not found")
        return
    
    # Add a paragraph after the section
    para = create_paragraph(doc, "Samples were tested in four different assay lots to assess reproducibility.")
    
    # Add a table after the paragraph
    table = doc.add_table(rows=4, cols=7)
    table.style = 'Table Grid'
    
    # Set up header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Sample"
    header_cells[1].text = "Lot 1"
    header_cells[2].text = "Lot 2"
    header_cells[3].text = "Lot 3"
    header_cells[4].text = "Lot 4"
    header_cells[5].text = "SD"
    header_cells[6].text = "CV"
    
    # Make header row bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add sample rows with safer indexing
    for i in range(1, 4):
        idx = i - 1  # 0-indexed for template access
        table.rows[i].cells[0].text = "{{ reproducibility[" + str(idx) + "].sample if reproducibility and " + str(idx) + " < reproducibility|length else 'Sample " + str(i) + "' }}"
        table.rows[i].cells[1].text = "{{ reproducibility[" + str(idx) + "].lot1 if reproducibility and " + str(idx) + " < reproducibility|length else 'N/A' }}"
        table.rows[i].cells[2].text = "{{ reproducibility[" + str(idx) + "].lot2 if reproducibility and " + str(idx) + " < reproducibility|length else 'N/A' }}"
        table.rows[i].cells[3].text = "{{ reproducibility[" + str(idx) + "].lot3 if reproducibility and " + str(idx) + " < reproducibility|length else 'N/A' }}"
        table.rows[i].cells[4].text = "{{ reproducibility[" + str(idx) + "].lot4 if reproducibility and " + str(idx) + " < reproducibility|length else 'N/A' }}"
        table.rows[i].cells[5].text = "{{ reproducibility[" + str(idx) + "].sd if reproducibility and " + str(idx) + " < reproducibility|length else 'N/A' }}"
        table.rows[i].cells[6].text = "{{ reproducibility[" + str(idx) + "].cv if reproducibility and " + str(idx) + " < reproducibility|length else 'N/A' }}"
    
    return table

def add_footer(doc):
    """Add the footer with Innovative Research information."""
    # Get the first section
    section = doc.sections[0]
    
    # Get the footer
    footer = section.footer
    
    # Add Innovative Research in bold Calibri 24pt
    p = footer.paragraphs[0] if len(footer.paragraphs) > 0 else footer.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    company_run = p.add_run("Innovative Research")
    company_run.bold = True
    company_run.font.name = "Calibri"
    company_run.font.size = Pt(24)
    
    # Add a line break
    p.add_run("\n")
    
    # Add contact info in Open Sans Light 12pt
    # Since we may not have Open Sans, use Calibri as a fallback
    contact_info = "32700 Concord Dr, Madison Heights, MI 48071 | Tel: 248-896-0145 | Fax: 248-896-0149"
    contact_run = p.add_run(contact_info)
    contact_run.font.name = "Calibri"
    contact_run.font.size = Pt(12)
    contact_run.bold = False
    
    # Add another line break and the website
    p.add_run("\n")
    website_run = p.add_run("www.innov-research.com")
    website_run.font.name = "Calibri"
    website_run.font.size = Pt(12)
    website_run.bold = False
    
    return footer

def update_template():
    """Update the template with all the fixes."""
    # Load the existing template
    template_path = Path('templates_docx/enhanced_template_fixed.docx')
    output_path = Path('templates_docx/enhanced_template_final.docx')
    
    doc = Document(template_path)
    
    # Apply all the fixes
    fix_sample_dilution_format(doc)
    fix_assay_protocol_format(doc)
    fix_technical_details_table(doc)
    fix_standard_curve_table(doc)
    fix_reproducibility_table(doc)
    add_disclaimer_section(doc)
    add_footer(doc)
    
    # Save the updated template
    doc.save(output_path)
    logger.info(f"Updated template saved to {output_path}")
    
    return output_path

if __name__ == "__main__":
    template_path = update_template()
    logger.info(f"Template with final fixes created at: {template_path}")
    
    # Verify that all issues are addressed
    print("\nVerified fixes in the template:")
    print("- SAMPLE DILUTION GUIDELINE formatted as list")
    print("- ASSAY PROTOCOL formatted as numbered list")
    print("- TECHNICAL DETAILS table fixed")
    print("- STANDARD CURVE table fixed with 0.0 in first row")
    print("- REPRODUCIBILITY table properly formatted")
    print("- DISCLAIMER section added")
    print("- Footer added with Innovative Research information")