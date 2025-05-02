#!/usr/bin/env python3
"""
Examine the REAGENTS PROVIDED section and tables in a Red Dot document in detail.
"""

import logging
import sys
from pathlib import Path
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def examine_reagents_section(document_path):
    """
    Examine the REAGENTS PROVIDED section and tables in detail.
    
    Args:
        document_path: Path to the document to examine
    """
    try:
        # Load the document
        doc = Document(document_path)
        
        # Find the REAGENTS PROVIDED section
        reagents_section_idx = None
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == "REAGENTS PROVIDED":
                reagents_section_idx = i
                logger.info(f"Found REAGENTS PROVIDED section at paragraph {i}")
                break
        
        if reagents_section_idx is None:
            logger.warning("REAGENTS PROVIDED section not found in document")
            return
        
        # Examine paragraphs around the section
        start_idx = max(0, reagents_section_idx - 1)
        end_idx = min(len(doc.paragraphs), reagents_section_idx + 5)
        
        logger.info("Paragraphs around REAGENTS PROVIDED section:")
        for i in range(start_idx, end_idx):
            logger.info(f"  Paragraph {i}: '{doc.paragraphs[i].text[:100]}...' - Style: {doc.paragraphs[i].style.name}")
        
        # Examine all tables in the document
        logger.info(f"Document contains {len(doc.tables)} tables")
        for i, table in enumerate(doc.tables):
            rows = len(table.rows)
            cols = len(table.columns) if rows > 0 else 0
            logger.info(f"Table {i}: {rows}x{cols}")
            
            # Print the first row to identify the table
            if rows > 0:
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                logger.info(f"  Headers: {header_cells}")
        
        # Check internal structure
        import inspect
        for table in doc.tables:
            table_element = table._element
            logger.info(f"Table element: {table_element.tag}")
            parent = table_element.getparent()
            logger.info(f"Parent element: {parent.tag}")
            if parent is not None:
                grandparent = parent.getparent()
                logger.info(f"Grandparent element: {grandparent.tag if grandparent is not None else None}")
                
        # Check the Document._body element to see where tables are stored
        body = doc._body._body
        logger.info(f"Body tag: {body.tag}")
        
        # Print the XML structure of the first few elements in the body
        logger.info("Body children (first 10):")
        for i, child in enumerate(body):
            if i < 10:
                logger.info(f"  Child {i}: {child.tag}")
                
        # Try to trace table relationship to paragraphs in XML structure
        logger.info("Table relationships in XML:")
        for i, child in enumerate(body):
            if 'tbl' in child.tag:
                # Found a table
                logger.info(f"  Found table at position {i} in body")
                # Look for preceding paragraph with REAGENTS PROVIDED
                for j in range(i-1, -1, -1):
                    if 'p' in body[j].tag:
                        # Found a paragraph
                        text = "".join([t.text for t in body[j].findall('.//{*}t') if t.text])
                        logger.info(f"    Preceding paragraph {j}: '{text}'")
                        if "REAGENTS PROVIDED" in text:
                            logger.info(f"    This is the REAGENTS PROVIDED section!")
                            logger.info(f"    Distance between heading and table: {i-j} elements")
                        break
        
    except Exception as e:
        logger.error(f"Error examining document: {e}")

if __name__ == "__main__":
    # Use the provided file path or default to complete_red_dot_output.docx
    if len(sys.argv) > 1:
        document_path = sys.argv[1]
    else:
        document_path = "complete_red_dot_output.docx"
    
    # Examine the document
    examine_reagents_section(document_path)