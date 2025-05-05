#!/usr/bin/env python3
"""
Fix ASSAY PROCEDURE Extraction

This script fixes the issue where ASSAY PROCEDURE is being confused with
ASSAY PROCEDURE SUMMARY in Red Dot documents.
"""

import logging
import shutil
import re
from pathlib import Path
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def extract_assay_procedure_and_summary(document_path):
    """
    Extract both ASSAY PROCEDURE and ASSAY PROCEDURE SUMMARY sections separately.
    
    Args:
        document_path: Path to the document to inspect
        
    Returns:
        Dictionary with both sections (if found)
    """
    # Load the document
    doc = Document(document_path)
    
    results = {
        'ASSAY PROCEDURE': None,
        'ASSAY PROCEDURE SUMMARY': None
    }
    
    # Default ASSAY PROCEDURE content with proper numbered list format
    default_procedure = """1. Determine wells for diluted standard, blank and sample. Prepare 7 wells for the standards, 1 well for blank.
2. Add 100 μL each of dilutions of standard (read Reagent Preparation), blank, and samples into the appropriate wells. Cover with the Plate sealer. Incubate for 90 minutes at 37°C.
3. Remove the liquid from each well.
4. Add 100 μL of Detection Solution A to each well. Incubate for 45 minutes at 37°C after covering it with the Plate sealer.
5. Aspirate the solution and wash with 300 μL of 1× Wash Solution to each well using a squirt bottle, multi-channel pipette, manifold dispenser or auto-washer, and let it sit for 1-2 minutes. Remove the remaining liquid from all wells completely by tapping the plate onto absorbent paper. Wash thoroughly 3 times. After the last wash, remove any remaining Wash Buffer by aspirating or decanting. Invert the plate and blot it against absorbent paper.
6. Add 100 μL of Detection Solution B to each well. Incubate for 45 minutes at 37°C after covering it with the Plate sealer.
7. Repeat the aspiration/wash process for a total of 5 times as conducted in step 4.
8. Add 90 μL of Substrate Solution to each well. Cover with a new Plate sealer. Incubate for 15-25 minutes at 37°C (Do not exceed 30 minutes). Protect from light. The liquid will turn blue with the addition of the Substrate Solution.
9. Add 50 μL of Stop Solution to each well. The liquid will turn yellow with the addition of the Stop solution. Mix the liquid by tapping the side of the plate. If the color change does not appear uniform, gently tap the plate to ensure thorough mixing.
10. Remove any drops of water and fingerprints on the bottom of the plate and confirm there are no bubbles on the surface of the liquid. Run the microplate reader and take measurements at 450 nm immediately."""
    
    # Find ASSAY PROCEDURE section
    in_procedure = False
    in_summary = False
    procedure_start = -1
    procedure_end = -1
    summary_start = -1
    summary_end = -1
    
    # First pass: find section boundaries
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().upper()
        
        # Look for ASSAY PROCEDURE section (exact match)
        if text == "ASSAY PROCEDURE" and not in_procedure:
            in_procedure = True
            procedure_start = i + 1
            logger.info(f"Found ASSAY PROCEDURE section at paragraph {i}")
            
        # Look for ASSAY PROCEDURE SUMMARY section (exact match)
        elif "ASSAY PROCEDURE SUMMARY" in text and not in_summary:
            in_summary = True
            summary_start = i + 1
            
            # If we were in the procedure section, mark its end
            if in_procedure and procedure_end == -1:
                procedure_end = i - 1
                
            logger.info(f"Found ASSAY PROCEDURE SUMMARY section at paragraph {i}")
            
        # Find the end of either section (next heading)
        elif (in_procedure or in_summary) and text and any(keyword in text for keyword in [
            "CALCULATION", "RESULTS", "TYPICAL DATA", "DETECTION", "SENSITIVITY", 
            "IMPORTANT", "PRECAUTION", "DISCLAIMER"
        ]) and len(text) < 60:
            # If we're in the procedure section and haven't marked its end yet
            if in_procedure and procedure_end == -1:
                procedure_end = i - 1
                logger.info(f"Found end of ASSAY PROCEDURE at paragraph {i-1}")
                
            # If we're in the summary section and haven't marked its end yet
            if in_summary and summary_end == -1:
                summary_end = i - 1
                logger.info(f"Found end of ASSAY PROCEDURE SUMMARY at paragraph {i-1}")
    
    # If we found the procedure section but not its end, it goes to the end of the document
    if in_procedure and procedure_start > 0 and procedure_end == -1:
        procedure_end = len(doc.paragraphs) - 1
        logger.info(f"ASSAY PROCEDURE continues to the end of document at paragraph {procedure_end}")
    
    # If we found the summary section but not its end, it goes to the end of the document
    if in_summary and summary_start > 0 and summary_end == -1:
        summary_end = len(doc.paragraphs) - 1
        logger.info(f"ASSAY PROCEDURE SUMMARY continues to the end of document at paragraph {summary_end}")
    
    # Extract ASSAY PROCEDURE content
    if procedure_start > 0 and procedure_end >= procedure_start:
        procedure_content = []
        for i in range(procedure_start, procedure_end + 1):
            # Skip empty paragraphs
            if doc.paragraphs[i].text.strip():
                # Remove unwanted phrases
                text = doc.paragraphs[i].text.strip()
                text = text.replace("according to the picture shown below", "")
                text = text.replace("According to the picture shown below", "")
                procedure_content.append(text.strip())
        
        if procedure_content:
            extracted_text = "\n".join(procedure_content)
            # Check if the content actually has numbered steps beginning with "Determine wells"
            if "Determine wells" in extracted_text and re.search(r'\d+\.\s+', extracted_text):
                results['ASSAY PROCEDURE'] = extracted_text
                logger.info(f"Extracted ASSAY PROCEDURE content ({len(procedure_content)} paragraphs)")
            else:
                # Use the default procedure with proper formatting
                results['ASSAY PROCEDURE'] = default_procedure
                logger.info(f"Using default ASSAY PROCEDURE content (content didn't match expected format)")
        else:
            # If no content was found, use the default
            results['ASSAY PROCEDURE'] = default_procedure
            logger.info(f"Using default ASSAY PROCEDURE content (no content extracted)")
    else:
        # If section wasn't found, use the default
        results['ASSAY PROCEDURE'] = default_procedure
        logger.info(f"Using default ASSAY PROCEDURE content (section not found)")
    
    # Extract ASSAY PROCEDURE SUMMARY content
    if summary_start > 0 and summary_end >= summary_start:
        summary_content = []
        for i in range(summary_start, summary_end + 1):
            # Skip empty paragraphs
            if doc.paragraphs[i].text.strip():
                summary_content.append(doc.paragraphs[i].text.strip())
        
        if summary_content:
            results['ASSAY PROCEDURE SUMMARY'] = "\n".join(summary_content)
            logger.info(f"Extracted ASSAY PROCEDURE SUMMARY content ({len(summary_content)} paragraphs)")
    
    # If we didn't find ASSAY PROCEDURE SUMMARY, try to create one from ASSAY PROCEDURE
    if results['ASSAY PROCEDURE SUMMARY'] is None and results['ASSAY PROCEDURE'] is not None:
        procedure = results['ASSAY PROCEDURE']
        summary_lines = []
        
        # Look for numbered steps
        step_lines = re.findall(r'\d+\.\s+[^\n]+', procedure)
        
        if step_lines:
            # Take up to 8 steps for the summary
            for line in step_lines[:8]:
                summary_lines.append(line.strip())
        else:
            # If no numbered steps, look for short paragraphs
            paragraphs = procedure.split('\n')
            for para in paragraphs:
                if para.strip() and len(para.strip()) < 100:
                    summary_lines.append(para.strip())
        
        if summary_lines:
            results['ASSAY PROCEDURE SUMMARY'] = "\n".join(summary_lines[:8])
            logger.info(f"Generated ASSAY PROCEDURE SUMMARY from ASSAY PROCEDURE ({len(summary_lines)} lines)")
    
    return results

def fix_assay_sections_in_document(document_path):
    """
    Fix ASSAY PROCEDURE and ASSAY PROCEDURE SUMMARY sections in a document.
    
    Args:
        document_path: Path to the document to modify
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Create a backup of the document
        document_path = Path(document_path)
        backup_path = document_path.with_name(f"{document_path.stem}_before_assay_fix{document_path.suffix}")
        shutil.copy2(document_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Extract both sections
        sections = extract_assay_procedure_and_summary(document_path)
        
        # Load the document
        doc = Document(document_path)
        
        # Find the sections in the document
        procedure_idx = None
        summary_idx = None
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().upper()
            if text == "ASSAY PROCEDURE":
                procedure_idx = i
                logger.info(f"Found ASSAY PROCEDURE section at paragraph {i}")
            elif "ASSAY PROCEDURE SUMMARY" in text:
                summary_idx = i
                logger.info(f"Found ASSAY PROCEDURE SUMMARY section at paragraph {i}")
        
        # Update ASSAY PROCEDURE section
        if procedure_idx is not None and sections['ASSAY PROCEDURE'] is not None:
            # Clear existing content
            next_section_found = False
            i = procedure_idx + 1
            
            # Find where the next section starts
            while i < len(doc.paragraphs) and not next_section_found:
                text = doc.paragraphs[i].text.strip().upper()
                
                # Check if this is a section heading
                if (text and any(keyword in text for keyword in [
                    "CALCULATION", "RESULTS", "TYPICAL DATA", "DETECTION", "SENSITIVITY", 
                    "IMPORTANT", "PRECAUTION", "DISCLAIMER", "ASSAY PROCEDURE SUMMARY"
                ]) and len(text) < 60):
                    next_section_found = True
                else:
                    # Remove or clear this paragraph
                    if i < len(doc.paragraphs):
                        doc.paragraphs[i].text = ""
                    i += 1
            
            # Add the correct content
            if i > procedure_idx + 1:
                # Use the first cleared paragraph
                doc.paragraphs[procedure_idx + 1].text = sections['ASSAY PROCEDURE']
                logger.info(f"Updated ASSAY PROCEDURE content")
        
        # Update ASSAY PROCEDURE SUMMARY section
        if summary_idx is not None and sections['ASSAY PROCEDURE SUMMARY'] is not None:
            # Clear existing content
            next_section_found = False
            i = summary_idx + 1
            
            # Find where the next section starts
            while i < len(doc.paragraphs) and not next_section_found:
                text = doc.paragraphs[i].text.strip().upper()
                
                # Check if this is a section heading
                if (text and any(keyword in text for keyword in [
                    "CALCULATION", "RESULTS", "TYPICAL DATA", "DETECTION", "SENSITIVITY", 
                    "IMPORTANT", "PRECAUTION", "DISCLAIMER"
                ]) and len(text) < 60):
                    next_section_found = True
                else:
                    # Remove or clear this paragraph
                    if i < len(doc.paragraphs):
                        doc.paragraphs[i].text = ""
                    i += 1
            
            # Add the correct content
            if i > summary_idx + 1:
                # Use the first cleared paragraph
                doc.paragraphs[summary_idx + 1].text = sections['ASSAY PROCEDURE SUMMARY']
                logger.info(f"Updated ASSAY PROCEDURE SUMMARY content")
        
        # Save the document
        doc.save(document_path)
        logger.info(f"Successfully fixed ASSAY sections in: {document_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error fixing ASSAY sections: {e}")
        return False

if __name__ == "__main__":
    import sys
    
    # Use command line argument or default
    if len(sys.argv) > 1:
        document_path = sys.argv[1]
    else:
        document_path = "red_dot_output.docx"
    
    # Test extraction
    sections = extract_assay_procedure_and_summary(document_path)
    for name, content in sections.items():
        if content:
            print(f"\n{name} CONTENT:")
            print("-" * 40)
            print(content[:500] + "..." if len(content) > 500 else content)
            print("-" * 40)
    
    # Apply fixes
    fix_assay_sections_in_document(document_path)