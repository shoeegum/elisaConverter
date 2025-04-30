#!/usr/bin/env python3
"""
Check the header format and first page layout in the output document.

This script verifies if the document header is properly sized at 36pt with Title style
and that the first page only contains the title, catalog number, lot number, and intended use.
"""

import docx
from docx.shared import Pt

def check_header_format(document_path="output_populated_template.docx"):
    """
    Check the header format in the output document.
    
    Args:
        document_path: Path to the document to check
    """
    print(f"Checking header format in {document_path}...")
    
    # Open the document
    doc = docx.Document(document_path)
    
    # Check style definitions first
    print("\nChecking style definitions:")
    if 'Title' in doc.styles:
        title_style = doc.styles['Title']
        if hasattr(title_style.font, 'size') and title_style.font.size is not None:
            print(f"Title style font size: {title_style.font.size.pt}pt")
        else:
            print("Title style font size: Not explicitly set")
        print(f"Title style font name: {title_style.font.name}")
        print(f"Title style bold: {title_style.font.bold}")
    else:
        print("No 'Title' style found in document")
    
    # Check the first paragraph (should be the title)
    print("\nChecking title paragraph:")
    if len(doc.paragraphs) > 0:
        title_para = doc.paragraphs[0]
        print(f"Title text: {title_para.text}")
        print(f"Title style: {title_para.style.name}")
        
        # Check if any runs in the title paragraph
        if len(title_para.runs) > 0:
            for i, run in enumerate(title_para.runs):
                if hasattr(run.font, 'size') and run.font.size is not None:
                    size_pt = run.font.size.pt
                    print(f"Run {i} font size: {size_pt}pt")
                else:
                    print(f"Run {i} font size: Not explicitly set (inherits from style)")
                print(f"Run {i} font name: {run.font.name}")
                print(f"Run {i} bold: {run.font.bold}")
        else:
            print("No runs found in title paragraph")
    
    # Check for page breaks
    print("\nChecking for page breaks to verify first page layout...")
    for i, para in enumerate(doc.paragraphs[:20]):  # Check the first 20 paragraphs
        for run in para.runs:
            if hasattr(run, '_element') and run._element.xpath('.//w:br[@w:type="page"]'):
                print(f"Found page break after paragraph {i}: '{para.text[:50]}...'")
    
    # Print the content of the first few paragraphs to check first page content
    print("\nContent of first few paragraphs:")
    for i, para in enumerate(doc.paragraphs[:10]):
        print(f"Paragraph {i}: {para.text[:50]}...")

if __name__ == "__main__":
    check_header_format()