#!/usr/bin/env python3
"""
Add Assay Principle Section

This script adds an ASSAY PRINCIPLE section to the document before the TECHNICAL DETAILS section.
Since we couldn't find the ASSAY PRINCIPLE section in the original document, we'll create it
with standard ELISA assay principle content.
"""

import logging
from pathlib import Path
import shutil
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def add_assay_principle(document_path):
    """
    Add an ASSAY PRINCIPLE section to the document before the TECHNICAL DETAILS section.
    
    Args:
        document_path: Path to the document to modify
    """
    try:
        # Make a backup of the document
        document_path = Path(document_path)
        backup_path = document_path.with_name(f"{document_path.stem}_before_adding_principle{document_path.suffix}")
        shutil.copy2(document_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Load the document
        doc = Document(document_path)
        
        # Find the TECHNICAL DETAILS section
        technical_details_idx = None
        for i, para in enumerate(doc.paragraphs):
            if "TECHNICAL DETAILS" in para.text.upper():
                technical_details_idx = i
                logger.info(f"Found TECHNICAL DETAILS section at paragraph {i}")
                break
                
        if technical_details_idx is None:
            logger.warning("Could not find TECHNICAL DETAILS section")
            return False
            
        # Create a new document to build the updated content
        temp_doc = Document()
        
        # Copy paragraphs up to TECHNICAL DETAILS
        for i in range(technical_details_idx):
            para = temp_doc.add_paragraph(doc.paragraphs[i].text)
            para.style = doc.paragraphs[i].style
            
        # Add ASSAY PRINCIPLE section
        principle_heading = temp_doc.add_paragraph("ASSAY PRINCIPLE")
        if 'Heading 2' in doc.styles:
            principle_heading.style = 'Heading 2'
            # Apply blue color to match other headings
            for run in principle_heading.runs:
                run.font.color.rgb = RGBColor(0, 70, 180)
                
        # Standard ELISA assay principle content
        principle_content = [
            "This ELISA Kit uses the Sandwich-ELISA principle. The micro ELISA plate provided in this kit has been pre-coated with an antibody specific to Mouse KLK1/Kallikrein 1. Standards or samples are added to the micro ELISA plate wells and combined with the specific antibody.",
            "Then a biotinylated detection antibody specific for Mouse KLK1/Kallikrein 1 and Avidin-Horseradish Peroxidase (HRP) conjugate are added successively to each micro plate well and incubated. Free components are washed away. The substrate solution is added to each well. Only those wells that contain Mouse KLK1/Kallikrein 1, biotinylated detection antibody and Avidin-HRP conjugate will appear blue in color. The enzyme-substrate reaction is terminated by the addition of stop solution and the color turns yellow.",
            "The optical density (OD) is measured spectrophotometrically at a wavelength of 450 nm ± 2 nm. The OD value is proportional to the concentration of Mouse KLK1/Kallikrein 1. You can calculate the concentration of Mouse KLK1/Kallikrein 1 in the samples by comparing the OD of the samples to the standard curve."
        ]
        
        # Add principle content paragraphs with 1.15 spacing
        for content in principle_content:
            para = temp_doc.add_paragraph(content)
            para.paragraph_format.line_spacing = 1.15
            # Add some space between paragraphs
            para.paragraph_format.space_after = Pt(6)
        
        # Now add the TECHNICAL DETAILS paragraph and all subsequent paragraphs
        for i in range(technical_details_idx, len(doc.paragraphs)):
            para = temp_doc.add_paragraph(doc.paragraphs[i].text)
            para.style = doc.paragraphs[i].style
            
        # Copy all tables 
        for table in doc.tables:
            # Get the dimensions of the table
            rows = len(table.rows)
            cols = len(table.rows[0].cells) if rows > 0 else 0
            
            # Create a new table with the same dimensions
            if rows > 0 and cols > 0:
                new_table = temp_doc.add_table(rows=rows, cols=cols)
                new_table.style = table.style
                
                # Copy cell content
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                            new_table.rows[i].cells[j].text = cell.text
                            
                            # Apply formatting to table cells
                            for para in new_table.rows[i].cells[j].paragraphs:
                                # Apply paragraph formatting
                                para.paragraph_format.line_spacing = 1.15
                                
                                # Apply font to all runs
                                for run in para.runs:
                                    run.font.name = "Calibri"
            
        # Save the document
        temp_path = document_path.with_name(f"{document_path.stem}_temp{document_path.suffix}")
        temp_doc.save(temp_path)
        
        # Now use the proper formatting function to ensure consistent styling
        from format_document import apply_document_formatting
        apply_document_formatting(temp_path)
        
        # Replace the original with our temporary document
        shutil.copy2(temp_path, document_path)
        
        # Clean up
        if temp_path.exists():
            import os
            os.remove(temp_path)
            
        logger.info(f"Successfully added ASSAY PRINCIPLE section before TECHNICAL DETAILS: {document_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error adding ASSAY PRINCIPLE section: {e}")
        return False

if __name__ == "__main__":
    # Add ASSAY PRINCIPLE to the current output document
    add_assay_principle("output_populated_template.docx")