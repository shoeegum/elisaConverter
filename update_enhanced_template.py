#!/usr/bin/env python3
"""
Update the enhanced template to include missing sections:
- ASSAY PRINCIPLE
- SAMPLE PREPARATION AND STORAGE
- SAMPLE COLLECTION NOTES
- SAMPLE DILUTION GUIDELINE
- DATA ANALYSIS
"""

import logging
from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def update_enhanced_template():
    """
    Update the enhanced template by adding the missing sections.
    """
    template_path = Path('templates_docx/enhanced_template.docx')
    new_template_path = Path('templates_docx/enhanced_template_updated.docx')
    
    # Load the template
    doc = Document(template_path)
    
    # Find the correct paragraph indices to insert new sections
    background_idx = None
    kit_components_idx = None
    assay_protocol_idx = None
    data_analysis_idx = None
    reproducibility_idx = None
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == 'BACKGROUND':
            background_idx = i
        elif para.text.strip() == 'KIT COMPONENTS':
            kit_components_idx = i
        elif para.text.strip() == 'ASSAY PROTOCOL':
            assay_protocol_idx = i
        elif para.text.strip() == 'REPRODUCIBILITY':
            reproducibility_idx = i
    
    if background_idx is None or kit_components_idx is None or assay_protocol_idx is None:
        logger.error("Could not find required sections in the template document")
        return
    
    # Add ASSAY PRINCIPLE section after BACKGROUND
    # Find the paragraph with BACKGROUND text plus 2 to get to the background content
    insert_after = background_idx + 2
    
    # Insert ASSAY PRINCIPLE heading
    assay_principle_heading = doc.paragraphs[insert_after-1]._element.addnext(
        doc.paragraphs[background_idx]._element.deepcopy()
    )
    assay_principle_para = doc.add_paragraph()  # Create an empty paragraph for reference, will be removed later
    
    # Insert ASSAY PRINCIPLE content placeholder
    assay_principle_content = doc.paragraphs[insert_after]._element.addnext(
        doc.paragraphs[background_idx + 1]._element.deepcopy()
    )
    
    # Add SAMPLE PREPARATION section before ASSAY PROTOCOL
    # Insert SAMPLE PREPARATION AND STORAGE heading
    sample_prep_heading = doc.paragraphs[assay_protocol_idx-1]._element.addnext(
        doc.paragraphs[background_idx]._element.deepcopy()
    )
    sample_prep_para = doc.add_paragraph()  # Create an empty paragraph for reference, will be removed later
    
    # Insert SAMPLE PREPARATION content placeholder
    sample_prep_content = doc.paragraphs[assay_protocol_idx]._element.addnext(
        doc.paragraphs[background_idx + 1]._element.deepcopy()
    )
    
    # Add SAMPLE COLLECTION NOTES section
    # Insert SAMPLE COLLECTION NOTES heading
    sample_notes_heading = doc.paragraphs[assay_protocol_idx+1]._element.addnext(
        doc.paragraphs[background_idx]._element.deepcopy()
    )
    sample_notes_para = doc.add_paragraph()  # Create an empty paragraph for reference, will be removed later
    
    # Insert SAMPLE COLLECTION NOTES content placeholder
    sample_notes_content = doc.paragraphs[assay_protocol_idx+2]._element.addnext(
        doc.paragraphs[background_idx + 1]._element.deepcopy()
    )
    
    # Add SAMPLE DILUTION GUIDELINE section
    # Insert SAMPLE DILUTION GUIDELINE heading
    sample_dilution_heading = doc.paragraphs[assay_protocol_idx+3]._element.addnext(
        doc.paragraphs[background_idx]._element.deepcopy()
    )
    sample_dilution_para = doc.add_paragraph()  # Create an empty paragraph for reference, will be removed later
    
    # Insert SAMPLE DILUTION GUIDELINE content placeholder
    sample_dilution_content = doc.paragraphs[assay_protocol_idx+4]._element.addnext(
        doc.paragraphs[background_idx + 1]._element.deepcopy()
    )
    
    # Add DATA ANALYSIS section after REPRODUCIBILITY
    # Insert DATA ANALYSIS heading
    data_analysis_heading = doc.paragraphs[reproducibility_idx+2]._element.addnext(
        doc.paragraphs[background_idx]._element.deepcopy()
    )
    data_analysis_para = doc.add_paragraph()  # Create an empty paragraph for reference, will be removed later
    
    # Insert DATA ANALYSIS content placeholder
    data_analysis_content = doc.paragraphs[reproducibility_idx+3]._element.addnext(
        doc.paragraphs[background_idx + 1]._element.deepcopy()
    )
    
    # Now fix the reproducibility table to add a standard deviation column
    if len(doc.tables) >= 7:  # Make sure we have enough tables
        repro_table = doc.tables[6]  # Reproducibility table is the 7th table (index 6)
        
        # Check if it already has a standard deviation column
        if len(repro_table.columns) < 7:
            # Insert a new column before the "Mean" column (which is the 5th column, index 4)
            for i, row in enumerate(repro_table.rows):
                cell = row.cells[4]  # Cell before which to insert the new column
                new_cell = row._tr.add_tc_before(cell._tc)
                
                # Add content to the new header cell
                if i == 0:
                    new_cell.text = "Standard Deviation"
                    
                    # Apply the same style as other header cells
                    for paragraph in new_cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
    
    # Update ASSAY PROTOCOL section to use a numbered list template instead of paragraphs
    if assay_protocol_idx is not None:
        # Find the paragraph containing the assay protocol content
        protocol_para_idx = assay_protocol_idx + 1
        if protocol_para_idx < len(doc.paragraphs):
            protocol_para = doc.paragraphs[protocol_para_idx]
            protocol_para.text = "{{ assay_protocol_numbered|default('') }}"
    
    # Save the document with proper paragraph text
    for i, para in enumerate(doc.paragraphs):
        if i == background_idx + 3:  # Assay Principle heading
            para.text = "ASSAY PRINCIPLE"
            para.style = 'Heading 2'
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 70, 180)
        elif i == background_idx + 4:  # Assay Principle content
            para.text = "{{ assay_principle|default('') }}"
            para.style = 'Normal'
        elif i == assay_protocol_idx + 2:  # Sample preparation heading
            para.text = "SAMPLE PREPARATION AND STORAGE"
            para.style = 'Heading 2'
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 70, 180)
        elif i == assay_protocol_idx + 3:  # Sample preparation content
            para.text = "{{ sample_preparation_and_storage|default('') }}"
            para.style = 'Normal'
        elif i == assay_protocol_idx + 5:  # Sample collection heading
            para.text = "SAMPLE COLLECTION NOTES"
            para.style = 'Heading 2'
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 70, 180)
        elif i == assay_protocol_idx + 6:  # Sample collection content
            para.text = "{{ sample_collection_notes|default('') }}"
            para.style = 'Normal'
        elif i == assay_protocol_idx + 8:  # Sample dilution heading
            para.text = "SAMPLE DILUTION GUIDELINE"
            para.style = 'Heading 2'
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 70, 180)
        elif i == assay_protocol_idx + 9:  # Sample dilution content
            para.text = "{{ sample_dilution_guideline|default('') }}"
            para.style = 'Normal'
        elif i == reproducibility_idx + 4:  # Data analysis heading
            para.text = "DATA ANALYSIS"
            para.style = 'Heading 2'
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 70, 180)
        elif i == reproducibility_idx + 5:  # Data analysis content
            para.text = "{{ data_analysis|default('') }}"
            para.style = 'Normal'
    
    # Save the updated template
    doc.save(new_template_path)
    logger.info(f"Updated template saved to {new_template_path}")
    
    # Create a copy of the template in the templates_docx folder
    import shutil
    shutil.copy(new_template_path, template_path)
    logger.info(f"Replaced original template at {template_path}")

if __name__ == "__main__":
    update_enhanced_template()