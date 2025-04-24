#!/usr/bin/env python3
"""
Check tables in DOCX files to inspect their structure.
"""

from docx import Document
import sys

def check_tables(document_path):
    doc = Document(document_path)
    print(f"Analyzing tables in document: {document_path}")
    print("-" * 80)
    
    print(f"Total tables: {len(doc.tables)}")
    
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i+1}: {len(table.rows)} rows x {len(table.columns)} columns")
        
        # Print header row contents
        if len(table.rows) > 0:
            print("  Header row contents:")
            for j, cell in enumerate(table.rows[0].cells):
                print(f"    Cell {j+1}: {cell.text}")
        
        # Print data row contents (first data row)
        if len(table.rows) > 1:
            print("  First data row contents:")
            for j, cell in enumerate(table.rows[1].cells):
                print(f"    Cell {j+1}: {cell.text}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python check_tables.py <docx_file>")
        sys.exit(1)
    
    check_tables(sys.argv[1])