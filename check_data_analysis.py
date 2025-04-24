"""
Check the data analysis section in the output document for publication citation text.
"""

import docx

def check_data_analysis(document_path="output_populated_template.docx"):
    """
    Check the data analysis section for publication citation text.
    
    Args:
        document_path: Path to the document to check
    """
    # Load the document
    doc = docx.Document(document_path)
    
    # Find the data analysis section
    section_start = -1
    
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().upper() == "DATA ANALYSIS":
            section_start = i
            break
            
    if section_start == -1:
        print("DATA ANALYSIS section not found")
        return
        
    # Find the next section to determine the end
    section_end = len(doc.paragraphs)
    for i in range(section_start + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].style.name.startswith('Heading'):
            section_end = i
            break
            
    # Extract section content
    section_text = "\n".join([doc.paragraphs[i].text for i in range(section_start + 1, section_end) if doc.paragraphs[i].text.strip()])
    
    print("=== DATA ANALYSIS SECTION CONTENT ===")
    print(section_text)
    print("====================================")
    
    # Check for publication citation text
    publication_phrases = [
        "Publications Citing This Product",
        "PubMed ID:",
        "html to see all",
        "publications"
    ]
    
    for phrase in publication_phrases:
        if phrase in section_text:
            print(f"WARNING: Found publication phrase '{phrase}' in DATA ANALYSIS section")
        else:
            print(f"OK: No '{phrase}' in DATA ANALYSIS section")

if __name__ == "__main__":
    check_data_analysis()