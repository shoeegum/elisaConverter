#!/usr/bin/env python3
"""
Check the REAGENTS PROVIDED section in the Innovative Research template output.

This script examines how the reagents information is processed 
and added to the output document.
"""

import logging
from pathlib import Path
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO,
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def check_reagents_provided(source_path, output_path):
    """
    Compare REAGENTS PROVIDED section between source and output documents.
    
    Args:
        source_path: Path to the source document
        output_path: Path to the output document
    """
    # Load both documents
    source_doc = Document(source_path)
    output_doc = Document(output_path)
    
    # Find the reagents section in the source document
    source_reagents = extract_reagents_from_source(source_doc)
    logger.info(f"Found reagents section in source document: {len(source_reagents)} paragraphs")
    
    # Find the reagents section in the output document
    output_reagents = extract_section_from_doc(output_doc, "REAGENTS PROVIDED")
    logger.info(f"Found reagents section in output document: {len(output_reagents)} paragraphs")
    
    # Compare the content
    logger.info("\n=== Source Document Reagents ===")
    for para in source_reagents:
        logger.info(para[:100] + "..." if len(para) > 100 else para)
    
    logger.info("\n=== Output Document Reagents ===")
    for para in output_reagents:
        logger.info(para[:100] + "..." if len(para) > 100 else para)
    
    # Check for tables in both documents
    logger.info("\n=== Source Document Tables ===")
    print_tables(source_doc)
    
    logger.info("\n=== Output Document Tables ===")
    print_tables(output_doc)

def extract_reagents_from_source(doc):
    """Extract reagents section from source document."""
    reagents = []
    in_reagents_section = False
    
    # Look for relevant section titles in the source document
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Check for section starts
        if not in_reagents_section and any(keyword in text.upper() for keyword in 
                                          ["REAGENTS PROVIDED", "KIT COMPONENTS", 
                                           "REAGENTS AND MATERIALS PROVIDED"]):
            in_reagents_section = True
            logger.info(f"Found reagents section in source at: {text}")
            continue
        
        # Check for section ends
        if in_reagents_section and text and any(keyword in text.upper() for keyword in 
                                              ["MATERIALS REQUIRED", "OTHER SUPPLIES", 
                                               "STORAGE", "SAMPLE COLLECTION", "SPECIMEN"]):
            logger.info(f"End of reagents section in source at: {text}")
            break
        
        # Collect section content
        if in_reagents_section and text:
            reagents.append(text)
    
    return reagents

def extract_section_from_doc(doc, section_name):
    """Extract content from a specific section in document."""
    content = []
    in_section = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Check for section start
        if section_name.upper() in text.upper() and not in_section:
            in_section = True
            logger.info(f"Found {section_name} section at: {text}")
            continue
        
        # Check for section end (next section start)
        if in_section and text and text.isupper() and len(text) < 50:
            logger.info(f"End of {section_name} section at: {text}")
            break
        
        # Collect section content
        if in_section and text:
            content.append(text)
    
    return content

def print_tables(doc):
    """Print information about tables in the document."""
    for i, table in enumerate(doc.tables):
        logger.info(f"Table {i+1}: {len(table.rows)} rows x {len(table.columns)} columns")
        
        # Print the first few rows of each table
        for j, row in enumerate(table.rows):
            if j < 3:  # Limit to first 3 rows
                row_content = " | ".join([cell.text for cell in row.cells])
                logger.info(f"  Row {j+1}: {row_content}")
        
        if len(table.rows) > 3:
            logger.info(f"  ... and {len(table.rows) - 3} more rows")

if __name__ == "__main__":
    import sys
    
    source_path = "attached_assets/RDR-LMNB2-Hu.docx"
    output_path = "improved_red_dot_output.docx"
    
    check_reagents_provided(source_path, output_path)