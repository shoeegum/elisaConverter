"""
Check the background section in the output document for publication citation text.
"""

import docx

def check_background(document_path="output_populated_template.docx"):
    """
    Check the background section for publication citation text.
    
    Args:
        document_path: Path to the document to check
    """
    # Load the document
    doc = docx.Document(document_path)
    
    # Find the background section
    background_section = None
    background_start = -1
    
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().upper() == "BACKGROUND":
            background_start = i
            break
            
    if background_start == -1:
        print("Background section not found")
        return
        
    # Find the next section to determine the end of background
    background_end = len(doc.paragraphs)
    for i in range(background_start + 1, len(doc.paragraphs)):
        if paragraph.style.name.startswith('Heading'):
            background_end = i
            break
            
    # Extract background content
    background_text = "\n".join([doc.paragraphs[i].text for i in range(background_start + 1, background_end) if doc.paragraphs[i].text.strip()])
    
    print("=== BACKGROUND SECTION CONTENT ===")
    print(background_text)
    print("=================================")
    
    # Check for publication citation text
    publication_phrases = [
        "Publications Citing This Product",
        "PubMed ID:",
        "html to see all",
        "publications"
    ]
    
    for phrase in publication_phrases:
        if phrase in background_text:
            print(f"WARNING: Found publication phrase '{phrase}' in background section")
        else:
            print(f"OK: No '{phrase}' in background section")

if __name__ == "__main__":
    check_background()