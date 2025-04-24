#!/usr/bin/env python3
"""
Check if the disclaimer has been added to the document.
"""

import docx

def check_disclaimer(document_path="output_populated_template.docx"):
    """
    Check if the disclaimer exists in the document and verify its text.
    
    Args:
        document_path: Path to the document to check
    """
    print(f"Checking disclaimer in {document_path}...")
    
    # Open the document
    doc = docx.Document(document_path)
    
    # Look for DISCLAIMER heading and section
    found_disclaimer_heading = False
    disclaimer_text = ""
    
    for i, para in enumerate(doc.paragraphs):
        if "DISCLAIMER" in para.text:
            found_disclaimer_heading = True
            print(f"Found DISCLAIMER heading at paragraph {i}")
            
            # Check next paragraph for disclaimer text
            if i + 1 < len(doc.paragraphs):
                disclaimer_text = doc.paragraphs[i+1].text
                break
    
    if found_disclaimer_heading:
        print("\nDisclaimer text:")
        print(disclaimer_text)
        
        # Check if the disclaimer text matches what's expected
        expected_text = "This material is sold for in-vitro use only in manufacturing and research. This material is not suitable for human use. It is the responsibility of the user to undertake sufficient verification and testing to determine the suitability of each product's application. The statements herein are offered for informational purposes only and are intended to be used solely for your consideration, investigation and verification."
        
        if expected_text in disclaimer_text:
            print("\nSUCCESS: The disclaimer text matches the expected text.")
        else:
            print("\nWARNING: The disclaimer text does not match exactly.")
            print("Expected: " + expected_text)
    else:
        print("ERROR: DISCLAIMER heading not found in the document.")

if __name__ == "__main__":
    check_disclaimer()