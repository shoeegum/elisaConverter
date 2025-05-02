#!/usr/bin/env python3
"""
Fix Red Dot documents by replacing company names and ensuring proper table placement.

This module provides functions that can be called from the post-processing step
of the Red Dot template populator to fix common issues with the output document.
"""

import logging
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def fix_company_names(document_path):
    """
    Replace all instances of 'Reddot Biotech INC.' with 'Innovative Research, Inc.'
    and 'Reddot Biotech' with 'Innovative Research' in the document.
    
    Args:
        document_path: Path to the document to modify
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Create a backup of the document
        document_path = Path(document_path)
        backup_path = document_path.with_name(f"{document_path.stem}_before_name_changes{document_path.suffix}")
        shutil.copy2(document_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Load the document
        doc = Document(document_path)
        
        # Define replacements
        replacements = [
            ('Reddot Biotech INC.', 'Innovative Research, Inc.'),
            ('Reddot Biotech', 'Innovative Research'),
        ]
        
        # Replace in paragraphs
        para_count = 0
        for para in doc.paragraphs:
            original_text = para.text
            modified_text = original_text
            
            for old_text, new_text in replacements:
                if old_text in modified_text:
                    modified_text = modified_text.replace(old_text, new_text)
            
            if modified_text != original_text:
                para.text = modified_text
                para_count += 1
        
        # Replace in tables
        table_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original_text = para.text
                        modified_text = original_text
                        
                        for old_text, new_text in replacements:
                            if old_text in modified_text:
                                modified_text = modified_text.replace(old_text, new_text)
                        
                        if modified_text != original_text:
                            para.text = modified_text
                            table_count += 1
        
        # Save the document
        doc.save(document_path)
        logger.info(f"Replaced company names in {para_count} paragraphs and {table_count} table cells")
        return True
        
    except Exception as e:
        logger.error(f"Error fixing company names: {e}")
        return False

def fix_table_position(document_path):
    """
    Fix the position of the REAGENTS PROVIDED table in the document.
    
    Args:
        document_path: Path to the document to modify
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Create a backup of the document
        document_path = Path(document_path)
        backup_path = document_path.with_name(f"{document_path.stem}_before_table_fix{document_path.suffix}")
        shutil.copy2(document_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Load the document
        doc = Document(document_path)
        
        # Find the REAGENTS PROVIDED section
        reagents_section_idx = None
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == "REAGENTS PROVIDED":
                reagents_section_idx = i
                logger.info(f"Found REAGENTS PROVIDED section at paragraph {i}")
                break
        
        if reagents_section_idx is None:
            logger.warning("REAGENTS PROVIDED section not found in document")
            return False
        
        # Find the next section after REAGENTS PROVIDED
        next_section_idx = None
        for i in range(reagents_section_idx + 1, len(doc.paragraphs)):
            if doc.paragraphs[i].style.name.startswith('Heading'):
                next_section_idx = i
                logger.info(f"Found next section at paragraph {i}: '{doc.paragraphs[i].text}'")
                break
        
        # Find tables in the document
        reagents_table = None
        reagents_table_idx = None
        
        for i, table in enumerate(doc.tables):
            if len(table.rows) > 0:
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                if 'Reagents' in header_cells and 'Quantity' in header_cells:
                    reagents_table = table
                    reagents_table_idx = i
                    logger.info(f"Found reagents table at index {i}")
                    break
        
        if reagents_table is None:
            logger.warning("Reagents table not found in document")
            return False
        
        # Extract table data before removing it
        table_data = []
        for row in reagents_table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        
        logger.info(f"Extracted data from table: {len(table_data)} rows")
        
        # Remove the original table from the document
        element = reagents_table._element
        element.getparent().remove(element)
        
        # Target position: right after the REAGENTS PROVIDED heading
        target_idx = reagents_section_idx + 1
        
        # If the next paragraph is empty or just whitespace, use it
        # Otherwise, add a new paragraph
        if target_idx < len(doc.paragraphs) and not doc.paragraphs[target_idx].text.strip():
            target_para = doc.paragraphs[target_idx]
            target_para.text = ""
        else:
            # Insert a new paragraph after the REAGENTS PROVIDED heading
            p_element = doc.paragraphs[reagents_section_idx]._element
            new_p = p_element.__class__()
            p_element.addnext(new_p)
            
            # Refresh the paragraphs list after adding the new paragraph
            doc = Document(document_path)
            target_idx = reagents_section_idx + 1
        
        # Add the table at the target position
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Table Grid'
        
        # Populate the table
        for i, row_data in enumerate(table_data):
            for j, cell_text in enumerate(row_data):
                if j < len(table.columns):
                    cell = table.cell(i, j)
                    cell.text = cell_text
        
        # Bold the header row
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Save the modified document
        doc.save(document_path)
        logger.info(f"Successfully moved table to the correct position in: {document_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error fixing table position: {e}")
        return False

def fix_document(document_path):
    """
    Apply all fixes to the document.
    
    Args:
        document_path: Path to the document to modify
        
    Returns:
        True if successful, False otherwise
    """
    success = True
    
    # First fix company names
    if not fix_company_names(document_path):
        logger.warning("Failed to fix company names")
        success = False
    
    # Then fix table position
    if not fix_table_position(document_path):
        logger.warning("Failed to fix table position")
        success = False
    
    return success

def process_output_document(document_path):
    """
    Process the output document after template population to fix common issues.
    This function is called from the Red Dot template populator's post-processing logic.
    
    Args:
        document_path: Path to the document to fix
        
    Returns:
        True if successful, False otherwise
    """
    return fix_document(document_path)

if __name__ == "__main__":
    import sys
    
    # Use command line argument or default
    if len(sys.argv) > 1:
        document_path = sys.argv[1]
    else:
        document_path = "complete_red_dot_output.docx"
    
    # Apply fixes
    if fix_document(document_path):
        logger.info(f"Successfully fixed: {document_path}")
    else:
        logger.error(f"Failed to fix all issues in: {document_path}")