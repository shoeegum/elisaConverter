#!/usr/bin/env python3
"""
Fix tables in the enhanced template
- Add a kit components table
- Add a standard curve table 
- Add a technical details table
- Fix the variability and reproducibility tables
"""

import logging
from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def add_kit_components_table(doc):
    """Add a kit components table to the document."""
    # Find the KIT COMPONENTS section
    kit_components_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "KIT COMPONENTS" in para.text.upper():
            kit_components_idx = i
            break
    
    if kit_components_idx is None:
        logger.warning("KIT COMPONENTS section not found")
        return
    
    # Add a paragraph after the KIT COMPONENTS heading
    para = doc.add_paragraph()
    
    # Add a table with 4 columns and 12 rows (1 header + 11 reagents)
    table = doc.add_table(rows=12, cols=4)
    table.style = 'Table Grid'
    
    # Set up header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Description"
    header_cells[1].text = "Quantity"
    header_cells[2].text = "Volume"
    header_cells[3].text = "Storage"
    
    # Make header row bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add reagent placeholders rows
    for i in range(1, 12):
        table.rows[i].cells[0].text = f"{{{{ reagent_{i}_name }}}}"
        table.rows[i].cells[1].text = f"{{{{ reagent_{i}_quantity }}}}"
        table.rows[i].cells[2].text = f"{{{{ reagent_{i}_volume }}}}"
        table.rows[i].cells[3].text = f"{{{{ reagent_{i}_storage }}}}"
    
    # Set column widths
    table.columns[0].width = Cm(5.0)  # Description
    table.columns[1].width = Cm(2.5)  # Quantity
    table.columns[2].width = Cm(2.5)  # Volume
    table.columns[3].width = Cm(5.0)  # Storage
    
    return table

def add_technical_details_table(doc):
    """Add a technical details table to the document."""
    # Find the TECHNICAL DETAILS section
    technical_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "TECHNICAL DETAILS" in para.text.upper():
            technical_idx = i
            break
    
    if technical_idx is None:
        logger.warning("TECHNICAL DETAILS section not found")
        return
    
    # Add a paragraph after the TECHNICAL DETAILS heading
    para = doc.add_paragraph()
    
    # Add a table with 2 columns and 5 rows
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # Set up rows
    properties = [
        "Capture/Detection Antibodies", 
        "Specificity",
        "Standard Protein",
        "Cross-reactivity",
        "Sensitivity"
    ]
    
    for i, prop in enumerate(properties):
        table.rows[i].cells[0].text = prop
        # Use safer access with default fallback if index doesn't exist
        table.rows[i].cells[1].text = "{{ technical_details_table[" + str(i) + "].value if technical_details_table and " + str(i) + " < technical_details_table|length else 'N/A' }}"
        
        # Make property names bold
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Set column widths
    table.columns[0].width = Cm(6.0)  # Property
    table.columns[1].width = Cm(9.0)  # Value
    
    return table

def fix_variability_tables(doc):
    """Fix the variability tables in the document."""
    # Find the INTRA/INTER-ASSAY VARIABILITY section
    variability_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "INTRA/INTER-ASSAY VARIABILITY" in para.text.upper():
            variability_idx = i
            break
    
    if variability_idx is None:
        logger.warning("VARIABILITY section not found")
        return
    
    # Add a paragraph with intra-assay text
    para = doc.add_paragraph("Three samples of known concentration were tested on one plate to assess intra-assay precision.")
    
    # Add intra-assay variability table
    intra_table = doc.add_table(rows=4, cols=4)
    intra_table.style = 'Table Grid'
    
    # Set up header row
    header_cells = intra_table.rows[0].cells
    header_cells[0].text = "Sample"
    header_cells[1].text = "n"
    header_cells[2].text = "Mean (pg/ml)"
    header_cells[3].text = "Standard Deviation"
    
    # Make header row bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add sample rows
    for i in range(1, 4):
        intra_table.rows[i].cells[0].text = f"Sample {i}"
        intra_table.rows[i].cells[1].text = "{{ variability.intra_assay.sample_" + str(i) + ".n if variability and variability.intra_assay else 'N/A' }}"
        intra_table.rows[i].cells[2].text = "{{ variability.intra_assay.sample_" + str(i) + ".mean if variability and variability.intra_assay else 'N/A' }}"
        intra_table.rows[i].cells[3].text = "{{ variability.intra_assay.sample_" + str(i) + ".sd if variability and variability.intra_assay else 'N/A' }}"
    
    # Add a paragraph with inter-assay text
    para = doc.add_paragraph("Three samples of known concentration were tested in separate assays to assess inter-assay precision.")
    
    # Add inter-assay variability table
    inter_table = doc.add_table(rows=4, cols=4)
    inter_table.style = 'Table Grid'
    
    # Set up header row
    header_cells = inter_table.rows[0].cells
    header_cells[0].text = "Sample"
    header_cells[1].text = "n"
    header_cells[2].text = "Mean (pg/ml)"
    header_cells[3].text = "Standard Deviation"
    
    # Make header row bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add sample rows
    for i in range(1, 4):
        inter_table.rows[i].cells[0].text = f"Sample {i}"
        inter_table.rows[i].cells[1].text = "{{ variability.inter_assay.sample_" + str(i) + ".n if variability and variability.inter_assay else 'N/A' }}"
        inter_table.rows[i].cells[2].text = "{{ variability.inter_assay.sample_" + str(i) + ".mean if variability and variability.inter_assay else 'N/A' }}"
        inter_table.rows[i].cells[3].text = "{{ variability.inter_assay.sample_" + str(i) + ".sd if variability and variability.inter_assay else 'N/A' }}"
    
    return intra_table, inter_table

def fix_reproducibility_table(doc):
    """Fix the reproducibility table in the document."""
    # Find the REPRODUCIBILITY section
    reproducibility_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "REPRODUCIBILITY" in para.text.upper():
            reproducibility_idx = i
            break
    
    if reproducibility_idx is None:
        logger.warning("REPRODUCIBILITY section not found")
        return
    
    # Add a paragraph with reproducibility text
    para = doc.add_paragraph("Samples were tested in four different assay lots to assess reproducibility.")
    
    # Add reproducibility table
    repro_table = doc.add_table(rows=4, cols=7)
    repro_table.style = 'Table Grid'
    
    # Set up header row
    header_cells = repro_table.rows[0].cells
    header_cells[0].text = "Sample"
    header_cells[1].text = "Lot 1"
    header_cells[2].text = "Lot 2"
    header_cells[3].text = "Lot 3"
    header_cells[4].text = "Lot 4"
    header_cells[5].text = "SD"
    header_cells[6].text = "CV"
    
    # Make header row bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add sample rows
    for i in range(1, 4):
        idx = i - 1  # 0-indexed for template access
        # Use safe indexing with defaults
        repro_table.rows[i].cells[0].text = "{{ reproducibility[" + str(idx) + "].sample if reproducibility and " + str(idx) + " < reproducibility|length else 'Sample " + str(i) + "' }}"
        repro_table.rows[i].cells[1].text = "{{ reproducibility[" + str(idx) + "].lot1 if reproducibility and " + str(idx) + " < reproducibility|length else 'N/A' }}"
        repro_table.rows[i].cells[2].text = "{{ reproducibility[" + str(idx) + "].lot2 if reproducibility and " + str(idx) + " < reproducibility|length else 'N/A' }}"
        repro_table.rows[i].cells[3].text = "{{ reproducibility[" + str(idx) + "].lot3 if reproducibility and " + str(idx) + " < reproducibility|length else 'N/A' }}"
        repro_table.rows[i].cells[4].text = "{{ reproducibility[" + str(idx) + "].lot4 if reproducibility and " + str(idx) + " < reproducibility|length else 'N/A' }}"
        repro_table.rows[i].cells[5].text = "{{ reproducibility[" + str(idx) + "].sd if reproducibility and " + str(idx) + " < reproducibility|length else 'N/A' }}"
        repro_table.rows[i].cells[6].text = "{{ reproducibility[" + str(idx) + "].cv if reproducibility and " + str(idx) + " < reproducibility|length else 'N/A' }}"
    
    return repro_table

def fix_all_tables():
    """Fix all tables in the enhanced template."""
    # Load the enhanced template
    template_path = Path('templates_docx/enhanced_template_complete.docx')
    output_path = Path('templates_docx/enhanced_template_fixed.docx')
    
    doc = Document(template_path)
    
    # Add or fix tables
    kit_table = add_kit_components_table(doc)
    if kit_table:
        logger.info("Added kit components table")
    
    technical_table = add_technical_details_table(doc)
    if technical_table:
        logger.info("Added technical details table")
    
    result = fix_variability_tables(doc)
    if result:
        intra_table, inter_table = result
        logger.info("Fixed variability tables")
    
    repro_table = fix_reproducibility_table(doc)
    if repro_table:
        logger.info("Fixed reproducibility table")
    
    # Save the updated template
    doc.save(output_path)
    logger.info(f"Updated template saved to {output_path}")
    
    return output_path

if __name__ == "__main__":
    template_path = fix_all_tables()
    logger.info(f"Template with fixed tables created at: {template_path}")
    
    # Verify that all tables are in the template
    print("\nVerify that these tables are fixed in the template:")
    print("- KIT COMPONENTS table with 4 columns (Description, Quantity, Volume, Storage)")
    print("- TECHNICAL DETAILS table")
    print("- INTRA/INTER-ASSAY VARIABILITY tables with Standard Deviation column")
    print("- REPRODUCIBILITY table with Standard Deviation column")