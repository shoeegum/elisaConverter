#!/usr/bin/env python3
"""
Update the Red Dot Template to Use Tables

This script modifies the enhanced Red Dot template to use proper tables
for the REAGENTS PROVIDED section.
"""

import logging
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def update_red_dot_template(template_path):
    """
    Update the Red Dot template to use proper tables for REAGENTS PROVIDED.
    
    Args:
        template_path: Path to the template document to modify
    """
    # Create a backup of the document
    template_path = Path(template_path)
    backup_path = template_path.with_name(f"{template_path.stem}_before_update{template_path.suffix}")
    shutil.copy2(template_path, backup_path)
    logger.info(f"Created backup at {backup_path}")
    
    # Load the document
    doc = Document(template_path)
    
    # Find the REAGENTS PROVIDED section
    reagents_section_index = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == "REAGENTS PROVIDED":
            reagents_section_index = i
            logger.info(f"Found REAGENTS PROVIDED section at paragraph {i}")
            break
    
    if reagents_section_index is None:
        logger.warning("REAGENTS PROVIDED section not found in template")
        return False
    
    # Find the placeholder paragraph after the section heading
    placeholder_index = reagents_section_index + 1
    if placeholder_index >= len(doc.paragraphs):
        logger.warning("No placeholder paragraph found after REAGENTS PROVIDED section")
        return False
    
    # Get the placeholder text
    placeholder_text = doc.paragraphs[placeholder_index].text
    logger.info(f"Current placeholder: {placeholder_text}")
    
    # Check if this is already using Jinja2 placeholders
    if "{{" in placeholder_text and "}}" in placeholder_text:
        # Replace the text placeholder with a table structure for Jinja2
        p = doc.paragraphs[placeholder_index]
        p.text = "{% if reagents_table_data %}"
        
        # Add a new paragraph for the table
        table_para = doc.add_paragraph()
        table_para.text = "{{ reagents_table_data }}"
        
        # Save the modified template
        doc.save(template_path)
        logger.info(f"Updated REAGENTS PROVIDED section in template: {template_path}")
        return True
    else:
        logger.warning("No Jinja2 placeholders found in REAGENTS PROVIDED section")
        return False

if __name__ == "__main__":
    # Update the enhanced Red Dot template
    template_path = Path("templates_docx/enhanced_red_dot_template.docx")
    if template_path.exists():
        update_red_dot_template(template_path)
    else:
        logger.error(f"Template not found: {template_path}")