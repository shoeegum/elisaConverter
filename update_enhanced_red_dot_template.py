#!/usr/bin/env python3
"""
Update Enhanced Innovative Research Template

This script modifies the enhanced Innovative Research template to:
1. Add proper placeholders for the reagents table in the REAGENTS PROVIDED section
2. Replace all Reddot Biotech references with Innovative Research
"""

import logging
import shutil
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def update_template(template_path):
    """
    Update the enhanced Innovative Research template.
    
    Args:
        template_path: Path to the template to update
    
    Returns:
        True if successful, False otherwise
    """
    try:
        # Create a backup of the template
        template_path = Path(template_path)
        backup_path = template_path.with_name(f"{template_path.stem}_backup{template_path.suffix}")
        shutil.copy2(template_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Load the template
        doc = Document(template_path)
        
        # Replace company names
        replacements = [
            ('Reddot Biotech INC.', 'Innovative Research, Inc.'),
            ('Reddot Biotech', 'Innovative Research'),
        ]
        
        name_replacements = 0
        # Fix in paragraphs
        for para in doc.paragraphs:
            original_text = para.text
            new_text = original_text
            
            for old_text, new_text in replacements:
                if old_text in new_text:
                    new_text = new_text.replace(old_text, new_text)
                    
            if new_text != original_text:
                para.text = new_text
                name_replacements += 1
        
        # Fix in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original_text = para.text
                        new_text = original_text
                        
                        for old_text, new_text in replacements:
                            if old_text in new_text:
                                new_text = new_text.replace(old_text, new_text)
                                
                        if new_text != original_text:
                            para.text = new_text
                            name_replacements += 1
        
        logger.info(f"Made {name_replacements} company name replacements")
        
        # Find the REAGENTS PROVIDED section
        reagents_section_idx = None
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == "REAGENTS PROVIDED":
                reagents_section_idx = i
                logger.info(f"Found REAGENTS PROVIDED section at paragraph {i}")
                break
        
        if reagents_section_idx is None:
            logger.warning("REAGENTS PROVIDED section not found")
            return False
        
        # Check if there's a paragraph after the section heading
        if reagents_section_idx + 1 < len(doc.paragraphs):
            next_para = doc.paragraphs[reagents_section_idx + 1]
            logger.info(f"Paragraph after REAGENTS PROVIDED: '{next_para.text}'")
            
            # Clear the next paragraph and add our special placeholder
            next_para.text = "{{ reagents_table_placeholder }}"
            logger.info("Added reagents_table_placeholder to template")
        else:
            # Add a new paragraph for the placeholder
            p = doc.add_paragraph("{{ reagents_table_placeholder }}")
            logger.info("Added new paragraph with reagents_table_placeholder")
        
        # Save the updated template
        doc.save(template_path)
        logger.info(f"Successfully updated template: {template_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error updating template: {e}")
        return False

if __name__ == "__main__":
    # Get the template path from command line or use default
    if len(sys.argv) > 1:
        template_path = sys.argv[1]
    else:
        template_path = "templates_docx/enhanced_innovative_research_template.docx"
    
    # Update the template
    update_template(template_path)