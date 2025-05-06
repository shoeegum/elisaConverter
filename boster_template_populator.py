#!/usr/bin/env python3
"""
Boster Template Populator

This module populates the Innovative Research template with data extracted from Boster ELISA kit datasheets.
It maps extracted ELISA kit data from Boster format to the Innovative Research template format.
"""

import logging
import re
import os
from pathlib import Path
from typing import Dict, Any, List, Tuple, Optional
from docxtpl import DocxTemplate

import docx
from elisa_parser import extract_elisa_data, ELISADatasheetParser

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Mapping of Boster document sections to Innovative Research template sections
BOSTER_SECTION_MAPPING = {
    # First page paragraph -> INTENDED USE
    "INTENDED USE": "INTENDED USE",
    
    # Standard mappings
    "ASSAY PRINCIPLE": "TEST PRINCIPLE",
    "PRINCIPLE OF THE ASSAY": "TEST PRINCIPLE",
    "OVERVIEW": "OVERVIEW",
    "TECHNICAL DETAILS": "TECHNICAL DETAILS",
    "PREPARATIONS BEFORE ASSAY": "PREPARATIONS BEFORE ASSAY",
    "KIT COMPONENTS": "REAGENTS PROVIDED",
    "MATERIALS PROVIDED": "REAGENTS PROVIDED",
    "REQUIRED MATERIALS THAT ARE NOT SUPPLIED": "OTHER SUPPLIES REQUIRED",
    "MATERIALS REQUIRED BUT NOT PROVIDED": "OTHER SUPPLIES REQUIRED",
    "DILUTION OF STANDARD": "DILUTION OF STANDARD",
    "SAMPLE PREPARATION AND STORAGE": "SAMPLE PREPARATION AND STORAGE",
    "SAMPLE COLLECTION NOTES": "SAMPLE COLLECTION NOTES",
    "SAMPLE DILUTION GUIDELINE": "SAMPLE DILUTION GUIDELINE",
    "ASSAY PROTOCOL": "ASSAY PROCEDURE",
    "DATA ANALYSIS": "CALCULATION OF RESULTS",
    "BACKGROUND": "BACKGROUND",
    "INTRA/INTER-ASSAY VARIABILITY": "REPRODUCIBILITY",
    "REPRODUCIBILITY": "REPRODUCIBILITY",
    "PREPARATIONS BEFORE THE EXPERIMENT": "PREPARATIONS BEFORE ASSAY",
}

def extract_boster_data(source_path):
    """
    Extract data from a Boster document with special handling for Boster-specific sections.
    
    Args:
        source_path: Path to the Boster document
        
    Returns:
        Dictionary containing extracted data
    """
    # First extract the basic data using the standard extractor
    data = extract_elisa_data(source_path)
    
    # Now enhance with Boster-specific extraction
    doc = docx.Document(source_path)
    
    # Extract the first page paragraph as INTENDED USE
    intended_use = ""
    for i, para in enumerate(doc.paragraphs[:20]):  # Check first 20 paragraphs
        # Look for paragraphs that are not headings and have substantial content
        if len(para.text.strip()) > 100 and not para.style.name.startswith('Heading'):
            intended_use = para.text.strip()
            logger.info(f"Extracted first page paragraph as INTENDED USE: {intended_use[:50]}...")
            break
    
    if intended_use:
        data["intended_use"] = intended_use
    
    # Create a dictionary to store Boster-specific sections
    boster_sections = {}
    
    # Define all the section headings we want to extract
    sections_to_extract = [
        "ASSAY PRINCIPLE", "OVERVIEW", "TECHNICAL DETAILS", 
        "PREPARATIONS BEFORE ASSAY", "KIT COMPONENTS", "MATERIALS PROVIDED",
        "REQUIRED MATERIALS THAT ARE NOT SUPPLIED", "DILUTION OF STANDARD",
        "SAMPLE PREPARATION AND STORAGE", "SAMPLE COLLECTION NOTES",
        "SAMPLE DILUTION GUIDELINE", "ASSAY PROTOCOL", "DATA ANALYSIS",
        "BACKGROUND", "INTRA/INTER-ASSAY VARIABILITY", "REPRODUCIBILITY",
        "PREPARATIONS BEFORE THE EXPERIMENT"
    ]
    
    # Initialize current section
    current_section = None
    section_content = []
    
    # Process paragraphs to extract sections
    for para in doc.paragraphs:
        text = para.text.strip()
        upper_text = text.upper()
        
        # Check if this is a section header
        is_section_header = False
        matched_section = None
        
        for section in sections_to_extract:
            # Check for exact match or close match (e.g., "Assay Principle" vs "ASSAY PRINCIPLE")
            if upper_text == section or upper_text == section.title():
                is_section_header = True
                matched_section = section
                break
        
        # If it's a new section header
        if is_section_header:
            # Save the previous section if we were tracking one
            if current_section and section_content:
                boster_sections[current_section] = "\n".join(section_content)
                logger.info(f"Extracted {current_section} section: {len(section_content)} paragraphs")
            
            # Start tracking the new section
            current_section = matched_section
            section_content = []
        
        # If we're in a section, add the paragraph
        elif current_section and text:
            section_content.append(text)
    
    # Add the last section if we were tracking one
    if current_section and section_content:
        boster_sections[current_section] = "\n".join(section_content)
        logger.info(f"Extracted {current_section} section: {len(section_content)} paragraphs")
    
    # Store the Boster-specific sections in the data
    data["boster_sections"] = boster_sections
    
    return data

def populate_boster_template(source_path, template_path, output_path,
                           kit_name=None, catalog_number=None, lot_number=None):
    """
    Extract data from a Boster document and populate an Innovative Research template.
    
    Args:
        source_path: Path to the Boster document
        template_path: Path to the Innovative Research template
        output_path: Path to save the populated template
        kit_name: Optional kit name override
        catalog_number: Optional catalog number override
        lot_number: Optional lot number override
        
    Returns:
        True if successful, False otherwise
    """
    try:
        source_path = Path(source_path)
        template_path = Path(template_path)
        output_path = Path(output_path)
        
        logger.info(f"Extracting data from Boster document: {source_path}")
        data = extract_boster_data(source_path)
        
        # Override extracted values if provided
        if kit_name:
            data["kit_name"] = kit_name
        if catalog_number:
            data["catalog_number"] = catalog_number
        if lot_number:
            data["lot_number"] = lot_number
        
        # Create a context dictionary for the template
        context = {
            "kit_name": data.get("kit_name", "ELISA Kit"),
            "catalog_number": data.get("catalog_number", ""),
            "lot_number": data.get("lot_number", ""),
        }
        
        # Map Boster sections to template sections
        if "boster_sections" in data:
            for boster_section, content in data["boster_sections"].items():
                if boster_section in BOSTER_SECTION_MAPPING:
                    template_section = BOSTER_SECTION_MAPPING[boster_section].lower()
                    template_section = template_section.replace(" ", "_")
                    context[template_section] = content
                    logger.info(f"Mapped {boster_section} to {template_section}")
        
        # Handle INTENDED USE separately since it's from the first page paragraph
        if "intended_use" in data:
            context["intended_use"] = data["intended_use"]
        
        # Map any standard sections that weren't in the Boster sections
        standard_mappings = {
            "background": "background",
            "materials_required": "other_supplies_required",
            "reagents_provided": "reagents_provided",
            "assay_principle": "test_principle",
            "storage": "storage_of_the_kits",
            "standard_curve": "typical_data"
        }
        
        for source_key, target_key in standard_mappings.items():
            if source_key in data and target_key not in context:
                context[target_key] = data[source_key]
                logger.info(f"Mapped standard {source_key} to {target_key}")
        
        # Load the template and render with the context
        template = DocxTemplate(template_path)
        template.render(context)
        
        # Create output directory if it doesn't exist
        output_path.parent.mkdir(exist_ok=True)
        
        # Save the populated template
        template.save(output_path)
        logger.info(f"Saved populated template to: {output_path}")
        
        # Replace any remaining instances of "Boster" or "PicoKine" with "Innovative Research"
        from fix_red_dot_document_comprehensive import fix_red_dot_document
        if fix_red_dot_document(output_path):
            logger.info(f"Applied comprehensive formatting fixes to: {output_path}")
        else:
            logger.warning(f"Could not apply all formatting fixes to: {output_path}")
        
        return True
        
    except Exception as e:
        logger.error(f"Error populating Boster template: {e}")
        return False

if __name__ == "__main__":
    # Example usage
    source_path = Path("attached_assets/EK1586_Mouse_KLK1Kallikrein_1_ELISA_Kit_PicoKine_Datasheet.docx")
    template_path = Path("templates_docx/innovative_template.docx")
    output_path = Path("output_boster_template.docx")
    
    populate_boster_template(source_path, template_path, output_path)