#!/usr/bin/env python3
"""
Create a test document with a Preparations Before Assay section
that includes numbered lists to verify that our code handles
numbered lists properly.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_test_document():
    """
    Create a test document that includes a Preparations Before Assay
    section with numbered lists.
    """
    doc = Document()
    
    # Add a title
    title = doc.add_heading("Test ELISA Kit Datasheet", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add basic metadata
    doc.add_paragraph("Catalog Number: EK1586")
    doc.add_paragraph("Lot Number: 12345")
    
    # Add a "Preparations Before Assay" section with numbered lists
    doc.add_heading("Preparations Before Assay", 1)
    
    # Add introduction paragraph
    doc.add_paragraph("Please read the following instructions before starting the experiment. This section includes both regular paragraphs and numbered instructions.")
    
    # Add some paragraphs first
    doc.add_paragraph("Before beginning, make sure to thoroughly read the entire protocol.")
    doc.add_paragraph("All reagents should be brought to room temperature (20-25°C) before use.")
    
    # Add numbered lists
    for i, step in enumerate([
        "Prepare all reagents, samples, and standards according to the instructions.",
        "Confirm that you have the appropriate non-supplied equipment available.",
        "Spin down all components to the bottom of the tube before opening.",
        "Don't let the 96-well plate dry out as this will inactivate active components.",
        "Don't reuse tips and tubes to avoid cross-contamination."
    ], 1):
        p = doc.add_paragraph()
        p.text = f"{i}. {step}"
    
    # Add a few more regular paragraphs
    doc.add_paragraph("Avoid using reagents from different batches together.")
    doc.add_paragraph("The kit should not be used beyond the expiration date on the kit label.")
    
    # Add another section to mark the end
    doc.add_heading("Kit Components", 1)
    
    # Save the document
    output_path = "attached_assets/test_numbered_lists.docx"
    doc.save(output_path)
    print(f"Test document created at: {output_path}")
    return output_path

if __name__ == "__main__":
    create_test_document()