#!/usr/bin/env python3
"""
Check tables in the fixed document
"""

from docx import Document

def check_tables(document_path="output_fixed_sample_sections.docx"):
    """Check the tables in the document"""
    doc = Document(document_path)
    
    print(f"Total tables: {len(doc.tables)}")
    
    for i, table in enumerate(doc.tables):
        first_cell = table.cell(0, 0).text[:40] + "..." if len(table.cell(0, 0).text) > 40 else table.cell(0, 0).text
        print(f"Table {i}: {len(table.rows)}x{len(table.columns)} - First cell: {first_cell}")
        
        # If this might be our sample preparation table
        if "Sample Type" in table.cell(0, 0).text:
            print(f"\nFound Sample Preparation Table (Table {i}):")
            print("  Headers:")
            for j in range(len(table.columns)):
                if j < len(table.columns):
                    print(f"    Column {j}: {table.cell(0, j).text}")
            
            print("\n  Data rows:")
            for r in range(1, len(table.rows)):
                if r < len(table.rows):
                    row_data = []
                    for c in range(len(table.columns)):
                        if c < len(table.columns):
                            cell_text = table.cell(r, c).text
                            if len(cell_text) > 30:
                                cell_text = cell_text[:30] + "..."
                            row_data.append(cell_text)
                    print(f"    Row {r}: {' | '.join(row_data)}")
            print()

if __name__ == "__main__":
    check_tables()