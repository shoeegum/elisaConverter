#!/usr/bin/env python3
"""
Create a better, clean template file from scratch that should avoid corruption issues.
"""

import logging
from pathlib import Path
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def create_better_template(output_path='clean_template.docx'):
    """
    Create a better, clean template from scratch using python-docx.
    
    Args:
        output_path: Path where the template will be saved
    """
    # Create a new Document
    doc = Document()
    
    # Set document properties
    core_properties = doc.core_properties
    core_properties.title = "ELISA Kit Datasheet Template"
    core_properties.author = "Innovative Research"
    
    # Set up page format - Letter size, narrow margins
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Title - with placeholder for kit name
    heading = doc.add_heading("{{ kit_name }}", level=0)
    for run in heading.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(36)
        run.font.bold = True
        
    # Add catalog and lot number
    catalog_lot_paragraph = doc.add_paragraph()
    catalog_lot_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    catalog_run = catalog_lot_paragraph.add_run("Catalog #: {{ catalog_number }}")
    catalog_run.font.name = 'Calibri'
    catalog_run.font.size = Pt(16)
    catalog_run.font.bold = True
    
    catalog_lot_paragraph.add_run(" | ")
    
    lot_run = catalog_lot_paragraph.add_run("Lot #: {{ lot_number }}")
    lot_run.font.name = 'Calibri'
    lot_run.font.size = Pt(16)
    lot_run.font.bold = True
    
    # Add a page break
    doc.add_paragraph().add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
    
    # Intended Use Section - H2 style (blue, all caps)
    intended_use_heading = doc.add_heading("INTENDED USE", level=2)
    for run in intended_use_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    intended_use = doc.add_paragraph("{{ intended_use }}")
    intended_use.style = 'Normal'
    intended_use.paragraph_format.space_after = Pt(12)
    
    # Background Section
    background_heading = doc.add_heading("BACKGROUND", level=2)
    for run in background_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    background = doc.add_paragraph("{{ background }}")
    background.style = 'Normal'
    background.paragraph_format.space_after = Pt(12)
    
    # Assay Principle
    assay_principle_heading = doc.add_heading("PRINCIPLE OF THE ASSAY", level=2)
    for run in assay_principle_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    principle = doc.add_paragraph("{{ assay_principle }}")
    principle.style = 'Normal'
    principle.paragraph_format.space_after = Pt(12)
    
    # Kit Components / Reagents
    kit_heading = doc.add_heading("KIT COMPONENTS", level=2)
    for run in kit_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    reagents = doc.add_paragraph("{{ reagents_table_content }}")
    reagents.style = 'Normal'
    reagents.paragraph_format.space_after = Pt(12)
    
    # Materials Required
    materials_heading = doc.add_heading("MATERIALS REQUIRED BUT NOT PROVIDED", level=2)
    for run in materials_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    materials = doc.add_paragraph("{{ required_materials }}")
    materials.style = 'Normal'
    materials.paragraph_format.space_after = Pt(12)
    
    # Technical Details
    technical_heading = doc.add_heading("TECHNICAL DETAILS", level=2)
    for run in technical_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    tech_details = doc.add_paragraph("{{ technical_details }}")
    tech_details.style = 'Normal'
    tech_details.paragraph_format.space_after = Pt(12)
    
    # Assay Protocol
    protocol_heading = doc.add_heading("ASSAY PROTOCOL", level=2)
    for run in protocol_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    protocol = doc.add_paragraph("{{ assay_protocol }}")
    protocol.style = 'Normal'
    protocol.paragraph_format.space_after = Pt(12)
    
    # Footer with copyright information
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    company_run = footer_para.add_run("INNOVATIVE RESEARCH")
    company_run.font.name = 'Calibri'
    company_run.font.size = Pt(24)
    company_run.font.bold = True
    
    footer_para.add_run("\n")  # Add a line break
    
    contact_run = footer_para.add_run("35200 Schoolcraft Rd, Livonia, MI 48150 | Phone: (248) 896-0142 | Fax: (248) 896-0148")
    contact_run.font.name = 'Calibri Light'
    contact_run.font.size = Pt(12)
    
    # Save the document
    try:
        doc.save(output_path)
        logger.info(f"Clean template successfully created and saved to {output_path}")
        return True
    except Exception as e:
        logger.error(f"Error saving template: {e}")
        return False

if __name__ == "__main__":
    create_better_template()