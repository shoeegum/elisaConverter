#!/usr/bin/env python3
"""
Fix Reagents Provided Table Format

This script modifies the Red Dot template populator to ensure that
REAGENTS PROVIDED section is displayed as a proper table with individual cells.
It also ensures that the table is properly rendered in the output document.
"""

import logging
import shutil
from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def fix_reagents_table(document_path):
    """
    Fix the REAGENTS PROVIDED section to use a proper table format.
    
    Args:
        document_path: Path to the document to modify
    """
    # Create a backup of the document
    document_path = Path(document_path)
    backup_path = document_path.with_name(f"{document_path.stem}_before_table_fix{document_path.suffix}")
    shutil.copy2(document_path, backup_path)
    logger.info(f"Created backup at {backup_path}")
    
    # Load the document
    doc = Document(document_path)
    
    # Find the REAGENTS PROVIDED section
    reagents_section_index = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == "REAGENTS PROVIDED":
            reagents_section_index = i
            logger.info(f"Found REAGENTS PROVIDED section at paragraph {i}")
            break
    
    if reagents_section_index is None:
        logger.warning("REAGENTS PROVIDED section not found in document")
        return False
    
    # Find the content paragraph after the section heading
    content_index = reagents_section_index + 1
    if content_index >= len(doc.paragraphs):
        logger.warning("No content paragraph found after REAGENTS PROVIDED section")
        return False
    
    # Get the content text
    content_text = doc.paragraphs[content_index].text
    logger.info(f"REAGENTS PROVIDED content: {content_text[:100]}...")
    
    # Check if the content contains a pipe-separated table format
    if "|" in content_text:
        logger.info("Detected pipe-separated table format in REAGENTS PROVIDED section")
        
        # Parse the text into table rows
        rows = []
        for line in content_text.split('\n'):
            if "|" in line and "-" * 10 not in line:  # Skip separator lines
                cells = [cell.strip() for cell in line.split('|')]
                rows.append(cells)
        
        # Make sure we have at least one row
        if not rows:
            logger.warning("No valid table rows found in REAGENTS PROVIDED content")
            return False
            
        logger.info(f"Extracted {len(rows)} rows from table text")
        
        # Create a new table to replace the text
        # First, find how many columns we need
        max_columns = max(len(row) for row in rows)
        
        # Create a table in the document before the content paragraph
        table = doc.add_table(rows=len(rows), cols=max_columns)
        table.style = 'Table Grid'
        
        # Populate the table with the extracted data
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < max_columns:  # Make sure we don't exceed the number of columns
                    cell = row.cells[j]
                    cell.text = cell_text
                    
                    # Apply formatting to cell text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)
        
        # Remove the original content paragraph
        p = doc.paragraphs[content_index]
        p._element.getparent().remove(p._element)
        
        # Save the modified document
        doc.save(document_path)
        logger.info(f"Successfully converted pipe-separated text to table in: {document_path}")
        return True
    else:
        logger.info("REAGENTS PROVIDED section does not contain a table format")
        return False

if __name__ == "__main__":
    # Fix the REAGENTS PROVIDED table in the output document
    fix_reagents_table("improved_red_dot_output.docx")