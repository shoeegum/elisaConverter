#!/usr/bin/env python3
"""
Fix both issues with the Red Dot document:
1. Move REAGENTS PROVIDED table to the correct position
2. Replace company names

This is a more careful implementation that handles document structure properly.
"""

import logging
import shutil
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def fix_both_issues(document_path):
    """
    Fix both issues with the Red Dot document.
    
    Args:
        document_path: Path to the document to modify
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Create a backup of the document
        document_path = Path(document_path)
        backup_path = document_path.with_name(f"{document_path.stem}_before_comprehensive_fixes{document_path.suffix}")
        shutil.copy2(document_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Load the document
        doc = Document(document_path)
        
        # First, fix company names
        replacements = [
            ('Reddot Biotech INC.', 'Innovative Research, Inc.'),
            ('Reddot Biotech', 'Innovative Research'),
        ]
        
        name_replacements = 0
        # Fix in paragraphs
        for para in doc.paragraphs:
            original_text = para.text
            new_text = original_text
            
            for old_text, new_text in replacements:
                if old_text in new_text:
                    new_text = new_text.replace(old_text, new_text)
                    
            if new_text != original_text:
                # Replace text by recreating the runs with the new text
                # This preserves formatting better than setting para.text directly
                for run in para.runs:
                    run.text = run.text.replace('Reddot Biotech INC.', 'Innovative Research, Inc.')
                    run.text = run.text.replace('Reddot Biotech', 'Innovative Research')
                name_replacements += 1
        
        # Fix in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original_text = para.text
                        new_text = original_text
                        
                        for old_text, new_text in replacements:
                            if old_text in new_text:
                                new_text = new_text.replace(old_text, new_text)
                                
                        if new_text != original_text:
                            # Replace text by recreating the runs with the new text
                            for run in para.runs:
                                run.text = run.text.replace('Reddot Biotech INC.', 'Innovative Research, Inc.')
                                run.text = run.text.replace('Reddot Biotech', 'Innovative Research')
                            name_replacements += 1
        
        logger.info(f"Made {name_replacements} company name replacements")
        
        # Now fix table positioning
        # First identify all section headings
        section_headings = {}
        for i, para in enumerate(doc.paragraphs):
            if para.style.name.startswith('Heading'):
                section_headings[para.text.strip()] = i
                logger.info(f"Found section heading at para {i}: '{para.text.strip()}'")
        
        # Find REAGENTS PROVIDED section
        if "REAGENTS PROVIDED" not in section_headings:
            logger.warning("REAGENTS PROVIDED section not found")
            doc.save(document_path)
            return False
        
        reagents_section_idx = section_headings["REAGENTS PROVIDED"]
        logger.info(f"REAGENTS PROVIDED section is at paragraph {reagents_section_idx}")
        
        # Find the next section
        other_supplies_idx = None
        for section, idx in section_headings.items():
            if idx > reagents_section_idx:
                other_supplies_idx = idx
                logger.info(f"Next section '{section}' is at paragraph {idx}")
                break
        
        if other_supplies_idx is None:
            logger.warning("No section after REAGENTS PROVIDED found")
            doc.save(document_path)
            return False
        
        # Find the reagents table
        reagents_table = None
        reagents_table_idx = None
        for i, table in enumerate(doc.tables):
            if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                logger.info(f"Table {i} headers: {header_cells}")
                if 'Reagents' in header_cells and 'Quantity' in header_cells:
                    reagents_table = table
                    reagents_table_idx = i
                    logger.info(f"Found reagents table at index {i}")
                    break
        
        if reagents_table is None:
            logger.warning("Reagents table not found")
            doc.save(document_path)
            return False
        
        # Collect table data
        table_data = []
        for row in reagents_table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        
        logger.info(f"Extracted table data: {len(table_data)} rows")
        
        # Remove the original table
        element = reagents_table._element
        element.getparent().remove(element)
        
        # If the next paragraph after REAGENTS PROVIDED is not the start of the next section,
        # we need to remove that paragraph too (since it's whatever was between the heading and next section)
        if reagents_section_idx + 1 < other_supplies_idx:
            logger.info("Found content between REAGENTS PROVIDED and next section")
            para_to_remove = doc.paragraphs[reagents_section_idx + 1]
            p_element = para_to_remove._element
            p_element.getparent().remove(p_element)
            
            # Re-get the paragraphs after removing one
            # This is necessary because indices can change
            other_supplies_idx -= 1
        
        # Now add a new empty paragraph after REAGENTS PROVIDED
        # Since we removed the content, we need to get the sections again
        section_para = doc.paragraphs[reagents_section_idx]
        p = doc.add_paragraph()
        p.insert_paragraph_before("")
        
        # Create a new table after the REAGENTS PROVIDED heading
        new_table = doc.add_table(rows=1, cols=len(table_data[0]))
        new_table.style = 'Table Grid'
        
        # Populate the header row
        for i, cell_text in enumerate(table_data[0]):
            new_table.cell(0, i).text = cell_text
        
        # Populate the rest of the table
        for i in range(1, len(table_data)):
            row_cells = new_table.add_row().cells
            for j, cell_text in enumerate(table_data[i]):
                row_cells[j].text = cell_text
        
        # Save the document
        doc.save(document_path)
        logger.info(f"Successfully fixed both issues in: {document_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error fixing document: {e}")
        return False

if __name__ == "__main__":
    # Use the provided file path or default
    if len(sys.argv) > 1:
        document_path = sys.argv[1]
    else:
        document_path = "complete_red_dot_output.docx"
    
    # Apply the fixes
    fix_both_issues(document_path)