#!/usr/bin/env python3
"""
Fix Document Structure

This script ensures:
1. The ASSAY PRINCIPLE section remains on the second page
2. Tables stay within their appropriate sections
3. Document maintains proper formatting and spacing
4. Sections are arranged in the correct order according to requirements
5. Content is properly extracted from the source document

Section Order Requirements:
1. Title, Catalog Number, Lot Number, and INTENDED USE (first page)
2. ASSAY PRINCIPLE (second page)
3. BACKGROUND (from source 'Background on...' section)
4. OVERVIEW (with table)
5. TECHNICAL DETAILS
6. PREPARATIONS BEFORE ASSAY (with numbered list)
7. KIT COMPONENTS/MATERIALS PROVIDED (table)
8. REQUIRED MATERIALS THAT ARE NOT SUPPLIED (bulleted list)
9. ELISA STANDARD CURVE EXAMPLE (with table and figure)
10. INTRA/INTER-ASSAY VARIABILITY (with tables)
11. REPRODUCIBILITY (with table)
12. PREPARATION BEFORE THE EXPERIMENT (with table)
13. DILUTION OF STANDARD (numbered list)
14. SAMPLE PREPARATION AND STORAGE (with table)
15. SAMPLE COLLECTION NOTES (numbered list)
16. SAMPLE DILUTION GUIDELINE
17. ASSAY PROTOCOL (numbered list)
18. DATA ANALYSIS
19. DISCLAIMER
"""

import logging
from pathlib import Path
import shutil
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def ensure_sections_with_tables(document_path):
    """
    Fix the document structure to ensure all sections appear in the correct order
    with their tables properly positioned within those sections.
    
    Args:
        document_path: Path to the document to modify
    """
    try:
        # Make a backup of the document
        document_path = Path(document_path)
        backup_path = document_path.with_name(f"{document_path.stem}_before_fix{document_path.suffix}")
        shutil.copy2(document_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Load the document
        doc = Document(document_path)
        
        # Extract source data from EK1586 document to get ASSAY PRINCIPLE text
        source_doc_path = Path('attached_assets/EK1586_Mouse_KLK1Kallikrein_1_ELISA_Kit_PicoKine_Datasheet.docx')
        if source_doc_path.exists():
            source_doc = Document(source_doc_path)
            # Try to find ASSAY PRINCIPLE section in the source document
            assay_principle_content = []
            found_assay_principle = False
            for paragraph in source_doc.paragraphs:
                if "ASSAY PRINCIPLE" in paragraph.text.upper():
                    found_assay_principle = True
                    continue
                    
                if found_assay_principle:
                    # If we hit another section heading, stop collecting content
                    if paragraph.text.strip().isupper() and len(paragraph.text.strip()) > 10:
                        break
                        
                    # Add non-empty paragraphs to our content
                    if paragraph.text.strip():
                        assay_principle_content.append(paragraph.text.strip())
            
            logger.info(f"Extracted {len(assay_principle_content)} paragraphs from ASSAY PRINCIPLE section")
        else:
            # Default content if source document not found
            assay_principle_content = [
                "This ELISA Kit uses the Sandwich-ELISA principle. The micro ELISA plate provided in this kit has been pre-coated with an antibody specific to Mouse KLK1/Kallikrein 1. Standards or samples are added to the micro ELISA plate wells and combined with the specific antibody.",
                "Then a biotinylated detection antibody specific for Mouse KLK1/Kallikrein 1 and Avidin-Horseradish Peroxidase (HRP) conjugate are added successively to each micro plate well and incubated. Free components are washed away. The substrate solution is added to each well. Only those wells that contain Mouse KLK1/Kallikrein 1, biotinylated detection antibody and Avidin-HRP conjugate will appear blue in color. The enzyme-substrate reaction is terminated by the addition of stop solution and the color turns yellow.",
                "The optical density (OD) is measured spectrophotometrically at a wavelength of 450 nm ± 2 nm. The OD value is proportional to the concentration of Mouse KLK1/Kallikrein 1. You can calculate the concentration of Mouse KLK1/Kallikrein 1 in the samples by comparing the OD of the samples to the standard curve."
            ]
        
        # Create a completely new document to rebuild with correct structure
        new_doc = Document()
        
        # Set overall document style
        styles = new_doc.styles
        if 'Normal' in styles:
            style = styles['Normal']
            style.font.name = "Calibri"
            style.paragraph_format.line_spacing = 1.15
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        
        # 1. Add the first page (cover page) - find title and metadata paragraphs
        cover_page_elements = []
        title_paragraph = None
        intended_use_section = None
        intended_use_content = []
        
        # Find cover page elements
        for i, para in enumerate(doc.paragraphs[:20]):  # Look in first 20 paragraphs
            if i == 0:  # First paragraph is always the title
                title_paragraph = para
                cover_page_elements.append(para)
            elif "catalog" in para.text.lower() or "lot" in para.text.lower():
                cover_page_elements.append(para)
            elif "intended use" in para.text.upper():
                intended_use_section = para
                # Get the content of the INTENDED USE section
                if i+1 < len(doc.paragraphs):
                    intended_use_content.append(doc.paragraphs[i+1])
        
        # Add cover page elements to new document
        for para in cover_page_elements:
            p = new_doc.add_paragraph(para.text)
            p.style = para.style
            # Make sure title is properly formatted
            if para == title_paragraph:
                p.style = 'Title'
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(36)
                    run.font.bold = True
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add INTENDED USE section
        if intended_use_section:
            p = new_doc.add_paragraph("INTENDED USE")
            p.style = 'Heading 2'
            for content_para in intended_use_content:
                p = new_doc.add_paragraph(content_para.text)
                p.style = content_para.style
        
        # Add a section break for page 2
        section = new_doc.add_section(WD_SECTION_START.NEW_PAGE)
        
        # 2. Add ASSAY PRINCIPLE section on the second page
        p = new_doc.add_paragraph("ASSAY PRINCIPLE")
        p.style = 'Heading 2'
        for run in p.runs:
            run.font.color.rgb = RGBColor(0, 70, 180)
            
        # Add the content for ASSAY PRINCIPLE section
        # Filter out the specific sentence that should be removed per requirements
        filtered_content = []
        for content in assay_principle_content:
            if "For more information on assay principle, protocols, and troubleshooting tips" not in content:
                filtered_content.append(content)
        
        # Add the filtered content
        for content in filtered_content:
            p = new_doc.add_paragraph(content)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
        
        # 3. Try to extract and add BACKGROUND section right after ASSAY PRINCIPLE
        # First try to find "Background on" in the source document
        background_content = []
        found_background = False
        
        if source_doc_path.exists():
            for paragraph in source_doc.paragraphs:
                if paragraph.text.strip().startswith("Background on"):
                    found_background = True
                    background_content.append(paragraph.text.strip())
                    break
        
        if found_background and background_content:
            # Add the BACKGROUND section
            p = new_doc.add_paragraph("BACKGROUND")
            p.style = 'Heading 2'
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 70, 180)
                
            # Add the content for BACKGROUND section
            for content in background_content:
                p = new_doc.add_paragraph(content)
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(6)
        
        # 3. Find the remaining sections and tables
        section_map = {}  # Maps section title paragraph to index
        table_map = {}    # Maps table index to nearby section title
        
        # Define the expected section order based on requirements
        expected_section_order = [
            "BACKGROUND",
            "OVERVIEW",
            "TECHNICAL DETAILS",
            "PREPARATIONS BEFORE ASSAY",
            "KIT COMPONENTS",
            "MATERIALS REQUIRED BUT NOT PROVIDED",
            "ELISA STANDARD CURVE EXAMPLE",
            "INTRA/INTER-ASSAY VARIABILITY",
            "REPRODUCIBILITY",
            "PREPARATION BEFORE THE EXPERIMENT",
            "DILUTION OF STANDARD",
            "SAMPLE PREPARATION AND STORAGE",
            "SAMPLE COLLECTION NOTES",
            "SAMPLE DILUTION GUIDELINE",
            "ASSAY PROTOCOL",
            "DATA ANALYSIS",
            "DISCLAIMER"
        ]
        
        # First, identify all section headings
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if (text.isupper() and any(text.startswith(prefix) for prefix in [
                "TECHNICAL", "OVERVIEW", "KIT COMPONENTS", "MATERIALS REQUIRED", 
                "SAMPLE", "ASSAY PROTOCOL", "DATA ANALYSIS", "BACKGROUND", 
                "DISCLAIMER", "ELISA STANDARD", "INTRA/INTER", "REPRODUCIBILITY",
                "PREPARATION", "DILUTION"])
            ):
                section_map[text] = i
                logger.info(f"Found section '{text}' at paragraph {i}")
        
        # Find tables and associate them with sections
        section_titles = list(section_map.keys())
        # Sort by expected order rather than document index
        sorted_section_titles = []
        for expected_section in expected_section_order:
            for title in section_titles:
                if title.startswith(expected_section) and title not in sorted_section_titles:
                    sorted_section_titles.append(title)
                    break
        
        # Add any sections that weren't in our expected list at the end
        for title in section_titles:
            if title not in sorted_section_titles:
                sorted_section_titles.append(title)
        
        section_titles = sorted_section_titles
        
        current_section = None
        for i, table in enumerate(doc.tables):
            # Look at first cell to help identify the table
            first_cell_text = ""
            if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
                first_cell_text = table.rows[0].cells[0].text.strip()
            
            # Associated this table with the nearest preceding section
            for title in reversed(section_titles):
                if section_map[title] < i:  # Section appears before table
                    table_map[i] = title
                    logger.info(f"Associating table {i} with section '{title}'")
                    break
        
        # 4. Add sections in order with their associated tables
        for title in section_titles:
            if title == "ASSAY PRINCIPLE":
                continue  # Skip as we've already added this
                
            # Add section heading
            p = new_doc.add_paragraph(title)
            p.style = 'Heading 2'
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 70, 180)
                
            # Find content paragraphs for this section
            section_idx = section_map[title]
            next_section_idx = float('inf')
            
            # Find the index of the next section
            for next_title in section_titles:
                next_idx = section_map[next_title]
                if next_idx > section_idx and next_idx < next_section_idx:
                    next_section_idx = next_idx
            
            # Add content paragraphs for this section
            for i in range(section_idx + 1, min(next_section_idx, len(doc.paragraphs))):
                # Skip if paragraph is part of another section heading
                if doc.paragraphs[i].text.strip() in section_map:
                    continue
                    
                # Add the paragraph
                if doc.paragraphs[i].text.strip():
                    p = new_doc.add_paragraph(doc.paragraphs[i].text)
                    p.style = doc.paragraphs[i].style
            
            # Add tables associated with this section
            for i, table in enumerate(doc.tables):
                if i in table_map and table_map[i] == title:
                    # Get the dimensions of the table
                    rows = len(table.rows)
                    cols = len(table.rows[0].cells) if rows > 0 else 0
                    
                    # Create a new table with the same dimensions
                    if rows > 0 and cols > 0:
                        new_table = new_doc.add_table(rows=rows, cols=cols)
                        new_table.style = table.style
                        
                        # Copy cell content
                        for row_idx, row in enumerate(table.rows):
                            for col_idx, cell in enumerate(row.cells):
                                if row_idx < len(new_table.rows) and col_idx < len(new_table.rows[row_idx].cells):
                                    target_cell = new_table.rows[row_idx].cells[col_idx]
                                    
                                    # Clear any default content
                                    if target_cell.paragraphs:
                                        for para in target_cell.paragraphs:
                                            para.clear()
                                    
                                    # Copy content from source cell to target cell
                                    for para in cell.paragraphs:
                                        new_para = target_cell.add_paragraph(para.text)
                                        # Apply formatting
                                        new_para.paragraph_format.line_spacing = 1.15
                        
                        logger.info(f"Added table for section '{title}'")
        
        # Ensure DISCLAIMER is the final section, immediately after DATA ANALYSIS
        # Check if we need to add the DISCLAIMER section
        disclaimer_added = False
        for title in section_titles:
            if title.startswith("DISCLAIMER"):
                disclaimer_added = True
                break
        
        if not disclaimer_added:
            # Add the DISCLAIMER section
            p = new_doc.add_paragraph("DISCLAIMER")
            p.style = 'Heading 2'
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 70, 180)
                
            # Add the standard disclaimer text
            disclaimer_text = ("This material is sold for in-vitro use only in manufacturing and research. "
                              "This material is not suitable for human use. It is the responsibility of the user "
                              "to undertake sufficient verification and testing to determine the suitability of "
                              "each product's application. The statements herein are offered for informational "
                              "purposes only and are intended to be used solely for your consideration, investigation "
                              "and verification.")
            
            p = new_doc.add_paragraph(disclaimer_text)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
        
        # Save the new document
        new_path = document_path.with_name(f"{document_path.stem}_structured{document_path.suffix}")
        new_doc.save(new_path)
        
        # Now use the proper formatting function to ensure consistent styling
        from format_document import apply_document_formatting
        apply_document_formatting(new_path)
        
        # Replace the original with our temporary document
        shutil.copy2(new_path, document_path)
        
        # Clean up
        if new_path.exists():
            import os
            os.remove(new_path)
            
        logger.info(f"Successfully fixed document structure: {document_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error fixing document structure: {e}")
        return False

if __name__ == "__main__":
    # Fix the current output document
    ensure_sections_with_tables("output_populated_template.docx")