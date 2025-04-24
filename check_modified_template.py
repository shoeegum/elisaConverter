#!/usr/bin/env python3
"""
Check the modified template for proper placeholders and tables.
"""

import logging
from docx import Document
from pathlib import Path

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def check_template(template_path="templates_docx/enhanced_template.docx"):
    """Check the template for expected placeholders and tables."""
    doc = Document(template_path)
    logger.info(f"Checking template at {template_path}")
    
    # Check the required materials section
    materials_section = False
    materials_placeholders = []
    
    # Check the variability and reproducibility sections
    variability_section = False
    reproducibility_section = False
    
    # Check all paragraphs
    for i, para in enumerate(doc.paragraphs):
        if "MATERIALS REQUIRED" in para.text.upper():
            materials_section = True
            logger.info(f"Found materials section at paragraph {i}")
        elif "INTRA/INTER-ASSAY VARIABILITY" in para.text.upper():
            variability_section = True
            logger.info(f"Found variability section at paragraph {i}")
        elif "REPRODUCIBILITY" in para.text.upper():
            reproducibility_section = True
            logger.info(f"Found reproducibility section at paragraph {i}")
        
        # Check for material placeholders
        if "req_material" in para.text:
            materials_placeholders.append(para.text)
    
    logger.info(f"Material placeholders: {len(materials_placeholders)}")
    
    # Check tables in the document
    standard_curve_table = None
    intra_assay_table = None
    inter_assay_table = None
    reproducibility_table = None
    
    for i, table in enumerate(doc.tables):
        # Check table dimensions
        rows = len(table.rows)
        cols = len(table.columns)
        logger.info(f"Table {i+1}: {rows} rows x {cols} columns")
        
        # Check first header cell to determine table type
        header_cell = table.cell(0, 0).text if rows > 0 and cols > 0 else ""
        
        if header_cell == "Concentration (pg/ml)" and cols > 8:
            standard_curve_table = i
            logger.info(f"Found standard curve table at index {i} ({rows}x{cols})")
        elif header_cell == "Sample" and "n" in table.cell(0, 1).text if cols > 1 else "":
            if intra_assay_table is None:
                intra_assay_table = i
                logger.info(f"Found intra-assay table at index {i} ({rows}x{cols})")
            else:
                inter_assay_table = i
                logger.info(f"Found inter-assay table at index {i} ({rows}x{cols})")
        elif header_cell == "" and "Lot 1" in table.cell(0, 1).text if cols > 1 else "":
            reproducibility_table = i
            logger.info(f"Found reproducibility table at index {i} ({rows}x{cols})")
    
    # Check for placeholders in tables
    if standard_curve_table is not None:
        table = doc.tables[standard_curve_table]
        placeholders = []
        for i in range(1, min(len(table.columns), 9)):
            try:
                cell_text = table.cell(1, i).text
                if "std_od" in cell_text or "{{" in cell_text:
                    placeholders.append(cell_text)
            except IndexError:
                pass
        logger.info(f"Standard curve placeholders: {len(placeholders)}")
    
    if intra_assay_table is not None:
        table = doc.tables[intra_assay_table]
        has_placeholders = False
        for row in range(1, min(len(table.rows), 4)):
            for col in range(1, min(len(table.columns), 5)):
                try:
                    cell_text = table.cell(row, col).text
                    if "intra_var" in cell_text or "{{" in cell_text:
                        has_placeholders = True
                        break
                except IndexError:
                    pass
        logger.info(f"Intra-assay table has placeholders: {has_placeholders}")
    
    if inter_assay_table is not None:
        table = doc.tables[inter_assay_table]
        has_placeholders = False
        for row in range(1, min(len(table.rows), 4)):
            for col in range(1, min(len(table.columns), 5)):
                try:
                    cell_text = table.cell(row, col).text
                    if "inter_var" in cell_text or "{{" in cell_text:
                        has_placeholders = True
                        break
                except IndexError:
                    pass
        logger.info(f"Inter-assay table has placeholders: {has_placeholders}")
    
    if reproducibility_table is not None:
        table = doc.tables[reproducibility_table]
        has_placeholders = False
        for row in range(1, min(len(table.rows), 4)):
            for col in range(1, min(len(table.columns), 7)):
                try:
                    cell_text = table.cell(row, col).text
                    if "repro" in cell_text or "{{" in cell_text:
                        has_placeholders = True
                        break
                except IndexError:
                    pass
        logger.info(f"Reproducibility table has placeholders: {has_placeholders}")
    
    print("\nSummary:")
    print(f"Materials section: {'Found' if materials_section else 'Missing'}")
    print(f"Material placeholders: {len(materials_placeholders)}")
    print(f"Standard curve table: {'Found' if standard_curve_table is not None else 'Missing'}")
    print(f"Standard curve format: {'Correct (2 rows)' if standard_curve_table is not None and len(doc.tables[standard_curve_table].rows) == 2 else 'Incorrect'}")
    print(f"Variability section: {'Found' if variability_section else 'Missing'}")
    print(f"Intra-assay table: {'Found' if intra_assay_table is not None else 'Missing'}")
    print(f"Inter-assay table: {'Found' if inter_assay_table is not None else 'Missing'}")
    print(f"Reproducibility section: {'Found' if reproducibility_section else 'Missing'}")
    print(f"Reproducibility table: {'Found' if reproducibility_table is not None else 'Missing'}")

if __name__ == "__main__":
    check_template()