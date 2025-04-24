#!/usr/bin/env python3
"""
Generate a clean document from scratch using python-docx.
This avoids template-related issues that might be causing corruption.
"""

import os
import sys
import logging
from pathlib import Path
import argparse
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

from elisa_parser import ELISADatasheetParser

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def create_document_from_data(data, output_path, kit_name=None, catalog_number=None, lot_number=None):
    """
    Create a new document from scratch with the extracted data.
    
    Args:
        data: Dictionary containing structured data extracted from the datasheet
        output_path: Path where the output document will be saved
        kit_name: Optional kit name provided by user
        catalog_number: Optional catalog number provided by user
        lot_number: Optional lot number provided by user
    """
    # Create a new Document
    doc = Document()
    
    # Set document properties
    core_properties = doc.core_properties
    core_properties.title = "ELISA Kit Datasheet"
    core_properties.author = "Innovative Research"
    
    # Set up page format - Letter size, narrow margins
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Use kit_name from data if not provided
    title = kit_name if kit_name else data.get('kit_name', 'Mouse KLK1 ELISA Kit')
    
    # Use catalog_number from data if not provided
    catalog = catalog_number if catalog_number else data.get('catalog_number', 'N/A')
    
    # Use lot_number from data if not provided
    lot = lot_number if lot_number else data.get('lot_number', 'SAMPLE')

    # Title - bold Calibri 36pt
    heading = doc.add_heading(title, level=0)
    for run in heading.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(36)
        run.font.bold = True
        
    # Add catalog and lot number
    catalog_lot_paragraph = doc.add_paragraph()
    catalog_lot_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    catalog_run = catalog_lot_paragraph.add_run(f"Catalog #: {catalog}")
    catalog_run.font.name = 'Calibri'
    catalog_run.font.size = Pt(16)
    catalog_run.font.bold = True
    
    catalog_lot_paragraph.add_run(" | ")
    
    lot_run = catalog_lot_paragraph.add_run(f"Lot #: {lot}")
    lot_run.font.name = 'Calibri'
    lot_run.font.size = Pt(16)
    lot_run.font.bold = True
    
    # Add a page break after the title section
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    # Intended Use Section - H2 style (blue, all caps)
    intended_use_heading = doc.add_heading("INTENDED USE", level=2)
    for run in intended_use_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    intended_use = doc.add_paragraph(data.get('intended_use', 'For research use only.'))
    intended_use.style = 'Normal'
    intended_use.paragraph_format.space_after = Pt(12)
    
    # Background Section
    background_heading = doc.add_heading("BACKGROUND", level=2)
    for run in background_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    background = doc.add_paragraph(data.get('background', ''))
    background.style = 'Normal'
    background.paragraph_format.space_after = Pt(12)
    
    # Assay Principle
    assay_principle_heading = doc.add_heading("PRINCIPLE OF THE ASSAY", level=2)
    for run in assay_principle_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    principle = doc.add_paragraph(data.get('assay_principle', ''))
    principle.style = 'Normal'
    principle.paragraph_format.space_after = Pt(12)
    
    # Technical Details Section (as a table)
    tech_heading = doc.add_heading("TECHNICAL DETAILS", level=2)
    for run in tech_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    # Add technical details as a table
    tech_table = doc.add_table(rows=5, cols=2)
    tech_table.style = 'Table Grid'
    tech_table.autofit = True
    
    # Set up headers and populate table
    headers = [
        'Capture/Detection Antibodies', 
        'Specificity',
        'Standard Protein',
        'Cross-reactivity',
        'Sensitivity'
    ]
    
    tech_details = data.get('technical_details', {})
    if isinstance(tech_details, dict) and 'technical_table' in tech_details:
        # Extract values from technical_table
        tech_values = {}
        for item in tech_details['technical_table']:
            tech_values[item['property']] = item['value']
    else:
        # Use individual fields if available
        tech_values = {
            'Capture/Detection Antibodies': '',
            'Specificity': data.get('specificity', ''),
            'Standard Protein': data.get('standard', ''),
            'Cross-reactivity': data.get('cross_reactivity', ''),
            'Sensitivity': data.get('sensitivity', '')
        }
    
    # Fill in the table
    for i, header in enumerate(headers):
        cell = tech_table.cell(i, 0)
        cell.text = header
        # Make header bold
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Calibri'
                
        # Add value
        value_cell = tech_table.cell(i, 1)
        value_cell.text = tech_values.get(header, 'N/A')
        for paragraph in value_cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Calibri'
    
    # Add space after the table
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Kit Components / Reagents
    kit_heading = doc.add_heading("KIT COMPONENTS", level=2)
    for run in kit_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    reagents = data.get('reagents', [])
    
    # Create a table for reagents
    if reagents:
        reagents_table = doc.add_table(rows=len(reagents)+1, cols=2)
        reagents_table.style = 'Table Grid'
        
        # Add headers
        header_row = reagents_table.rows[0]
        header_row.cells[0].text = "Component"
        header_row.cells[1].text = "Quantity"
        
        # Make header bold
        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Calibri'
        
        # Add reagents
        for i, reagent in enumerate(reagents):
            row = reagents_table.rows[i+1]
            row.cells[0].text = reagent.get('name', '')
            row.cells[1].text = reagent.get('quantity', '')
            
            # Apply formatting to cells
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
    else:
        # If no reagents, add a paragraph
        no_reagents = doc.add_paragraph("Kit components information not available")
        no_reagents.style = 'Normal'
    
    # Add space after reagents
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Materials Required
    materials_heading = doc.add_heading("MATERIALS REQUIRED BUT NOT PROVIDED", level=2)
    for run in materials_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    materials = data.get('required_materials', [])
    if materials:
        for material in materials:
            # Add as bullet points
            p = doc.add_paragraph(material, style='List Bullet')
            p.style = 'List Bullet'
            for run in p.runs:
                run.font.name = 'Calibri'
    else:
        no_materials = doc.add_paragraph("Materials information not available")
        no_materials.style = 'Normal'
    
    # Add space after materials
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Assay Protocol
    protocol_heading = doc.add_heading("ASSAY PROTOCOL", level=2)
    for run in protocol_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    protocol_steps = data.get('assay_protocol', [])
    if protocol_steps:
        for i, step in enumerate(protocol_steps):
            # Add as numbered list
            p = doc.add_paragraph(f"{i+1}. {step}")
            p.style = 'Normal'
            for run in p.runs:
                run.font.name = 'Calibri'
            p.paragraph_format.space_after = Pt(6)
    else:
        no_protocol = doc.add_paragraph("Protocol information not available")
        no_protocol.style = 'Normal'
    
    # Footer with copyright information
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    company_run = footer_para.add_run("INNOVATIVE RESEARCH")
    company_run.font.name = 'Calibri'
    company_run.font.size = Pt(24)
    company_run.font.bold = True
    
    footer_para.add_run("\\n")  # Add a line break
    
    contact_run = footer_para.add_run("35200 Schoolcraft Rd, Livonia, MI 48150 | Phone: (248) 896-0142 | Fax: (248) 896-0148")
    contact_run.font.name = 'Calibri Light'
    contact_run.font.size = Pt(12)
    
    # Save the document
    try:
        doc.save(output_path)
        logger.info(f"Document successfully created and saved to {output_path}")
        return True
    except Exception as e:
        logger.error(f"Error saving document: {e}")
        return False

def parse_arguments():
    """Parse command line arguments"""
    parser = argparse.ArgumentParser(
        description='Generate a clean document from scratch with data from an ELISA kit datasheet'
    )
    
    parser.add_argument(
        '--source', 
        required=True,
        help='Path to the source ELISA kit datasheet DOCX file'
    )
    
    parser.add_argument(
        '--output', 
        required=True,
        help='Path where the generated document will be saved'
    )
    
    parser.add_argument(
        '--kit-name', 
        help='Name of the ELISA kit (e.g., "Mouse KLK1 ELISA Kit")'
    )
    
    parser.add_argument(
        '--catalog-number', 
        help='Catalog number of the ELISA kit (e.g., "EK1586")'
    )
    
    parser.add_argument(
        '--lot-number', 
        help='Lot number of the ELISA kit'
    )
    
    return parser.parse_args()

def main():
    """Main entry point"""
    args = parse_arguments()
    
    try:
        # Parse the ELISA datasheet
        logger.info(f"Parsing ELISA datasheet: {args.source}")
        parser = ELISADatasheetParser(args.source)
        data = parser.extract_data()
        
        # Create a new document from scratch with the extracted data
        logger.info(f"Creating new document with extracted data")
        success = create_document_from_data(
            data, 
            args.output, 
            kit_name=args.kit_name, 
            catalog_number=args.catalog_number, 
            lot_number=args.lot_number
        )
        
        if success:
            logger.info(f"Document successfully generated at: {args.output}")
            return 0
        else:
            logger.error("Failed to generate document")
            return 1
    
    except Exception as e:
        logger.exception(f"Error processing files: {e}")
        return 1

if __name__ == "__main__":
    sys.exit(main())