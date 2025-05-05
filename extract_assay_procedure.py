#!/usr/bin/env python3
"""
Extract ASSAY PROCEDURE Section

This script examines the source document to locate and extract the 
ASSAY PROCEDURE section for Red Dot documents, separate from the
ASSAY PROCEDURE SUMMARY section.
"""

import logging
from pathlib import Path
from docx import Document
import re

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def extract_assay_procedure(document_path):
    """
    Extract the ASSAY PROCEDURE section from a document,
    separate from the ASSAY PROCEDURE SUMMARY section.

    Args:
        document_path: Path to the document to inspect
    
    Returns:
        Text content of the ASSAY PROCEDURE section if found, None otherwise
    """
    # Load the document
    doc = Document(document_path)
    
    # Flags to track section boundaries
    in_assay_procedure = False
    start_idx = -1
    end_idx = -1
    
    # First pass: identify section boundaries
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().upper()
        
        # Check for section start - use an exact match to avoid confusion with SUMMARY
        if (text == "ASSAY PROCEDURE" or text == "ASSAY PROTOCOL") and not in_assay_procedure:
            in_assay_procedure = True
            start_idx = i + 1  # Start after the section heading
            logger.info(f"Found ASSAY PROCEDURE section at paragraph {i}: {para.text}")
        
        # Check for section end (next section starts)
        elif in_assay_procedure and text and any(keyword in text for keyword in [
            "CALCULATION", "RESULTS", "ASSAY PROCEDURE SUMMARY", "TYPICAL", 
            "DETECTION", "SENSITIVITY", "ASSAY SUMMARY"
        ]) and len(text) < 100:  # Likely a new section header
            end_idx = i - 1  # End before the next section
            logger.info(f"Found end of ASSAY PROCEDURE at paragraph {i-1}")
            break
    
    # If we found the start but not the end, assume it continues to the end of the document
    if in_assay_procedure and start_idx > 0 and end_idx == -1:
        end_idx = len(doc.paragraphs) - 1
        logger.info(f"ASSAY PROCEDURE continues to the end of document at paragraph {end_idx}")
    
    # Extract section content
    if start_idx > 0 and end_idx >= start_idx:
        # Join the paragraphs
        content = []
        for i in range(start_idx, end_idx + 1):
            # Skip empty paragraphs
            if doc.paragraphs[i].text.strip():
                # Remove "according to the picture shown below" phrase
                text = doc.paragraphs[i].text.strip()
                text = text.replace("according to the picture shown below", "")
                text = text.replace("According to the picture shown below", "")
                content.append(text.strip())
        
        # Return the section content
        if content:
            full_content = "\n".join(content)
            logger.info(f"Extracted ASSAY PROCEDURE: {full_content[:100]}...")
            return full_content
        else:
            logger.warning("ASSAY PROCEDURE section found but no content extracted")
            return None
    else:
        logger.warning("Could not find ASSAY PROCEDURE section")
        return None

if __name__ == "__main__":
    # Extract from the source document
    content = extract_assay_procedure("attached_assets/EK1586_Mouse_KLK1Kallikrein_1_ELISA_Kit_PicoKine_Datasheet.docx")
    if content:
        print("\nFULL ASSAY PROCEDURE:")
        print("-" * 40)
        print(content)
        print("-" * 40)