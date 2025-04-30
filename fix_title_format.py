#!/usr/bin/env python3
"""
Fix the Title formatting in the output document.

This script ensures the document title is properly formatted with:
- Calibri font
- 36pt size
- Bold style
- Center alignment
"""

import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Set up logging
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def fix_title_format(document_path="output_populated_template.docx"):
    """
    Fix the title formatting in a document.
    
    Args:
        document_path: Path to the document to fix
    """
    logger.info(f"Fixing title formatting in {document_path}")
    
    # Create backup
    backup_path = f"{document_path}_backup"
    try:
        import shutil
        shutil.copy2(document_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
    except Exception as e:
        logger.warning(f"Could not create backup: {e}")
        
    # Open the document
    doc = Document(document_path)
    
    # Check and fix the title style
    if 'Title' in doc.styles:
        title_style = doc.styles['Title']
        # Set style properties
        title_style.font.name = "Calibri"
        title_style.font.size = Pt(36)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        logger.info("Updated Title style in document")
    
    # Find and fix the title paragraph (first paragraph)
    if len(doc.paragraphs) > 0:
        title_para = doc.paragraphs[0]
        
        # Apply the Title style
        title_para.style = 'Title'
        
        # Also apply direct formatting
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Apply direct formatting to all runs in the title
        for run in title_para.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(36)
            run.font.bold = True
            
        logger.info(f"Updated title paragraph: {title_para.text}")
        
        # If there are no runs in the paragraph (sometimes happens),
        # add the content as a new run with proper formatting
        if len(title_para.runs) == 0:
            title_text = title_para.text
            title_para.clear()
            new_run = title_para.add_run(title_text)
            new_run.font.name = "Calibri"
            new_run.font.size = Pt(36)
            new_run.font.bold = True
            logger.info(f"Added new formatted run with text: {title_text}")
    
    # Save the document
    doc.save(document_path)
    logger.info(f"Saved document with fixed title formatting to {document_path}")
    
if __name__ == "__main__":
    fix_title_format()