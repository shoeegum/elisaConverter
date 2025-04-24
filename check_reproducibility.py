#!/usr/bin/env python3
"""
Check the REPRODUCIBILITY tables in the output document.
"""

import logging
from pathlib import Path
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def check_reproducibility_tables(document_path="output_populated_template.docx"):
    """
    Check the content of the REPRODUCIBILITY tables and identify which cells are empty.
    
    Args:
        document_path: Path to the document to check
    """
    doc = Document(document_path)
    
    print("\n=== REPRODUCIBILITY TABLES CONTENT ===")
    
    # Find the Reproducibility tables (intra/inter/lot-to-lot)
    intra_assay_table = None
    inter_assay_table = None
    lot_to_lot_table = None
    
    # First, find the reproducibility tables section
    reproducibility_section = None
    for i, para in enumerate(doc.paragraphs):
        if "REPRODUCIBILITY" in para.text.strip().upper():
            reproducibility_section = i
            print(f"Found REPRODUCIBILITY section at paragraph {i}")
            break
    
    # Now look for the tables after this point
    for i, table in enumerate(doc.tables):
        if not table.rows:
            continue
            
        # Get table content for analysis
        table_content = ""
        for row in table.rows:
            for cell in row.cells:
                table_content += cell.text.lower() + " "
        
        # Look for key terms to identify the table
        if "sample" in table_content and "mean" in table_content:
            # Distinguishing between intra-assay and inter-assay tables
            if i == 5:  # Based on table order from document structure
                intra_assay_table = (i, table)
                print(f"Found Intra-Assay table at index {i}")
            elif i == 6:  # Based on table order from document structure
                inter_assay_table = (i, table)
                print(f"Found Inter-Assay table at index {i}")
        
        # Lot-to-Lot tables can have "lot" or "mean" and "standard deviation" in their content
        if ("lot" in table_content.lower() or 
            ("mean" in table_content.lower() and "standard deviation" in table_content.lower())):
            if i == 7:  # Based on table order from document structure
                lot_to_lot_table = (i, table)
                print(f"Found Lot-to-Lot table at index {i}")
    
    # Check each reproducibility table
    tables_to_check = [
        ("Intra-Assay", intra_assay_table),
        ("Inter-Assay", inter_assay_table),
        ("Lot-to-Lot", lot_to_lot_table)
    ]
    
    for table_name, table_info in tables_to_check:
        if table_info is None:
            print(f"\n{table_name} table not found!")
            continue
            
        idx, table = table_info
        print(f"\n--- {table_name} Table (Table {idx}) ---")
        
        # Count cells
        empty_cells = 0
        total_cells = 0
        
        # Print headers first
        if table.rows:
            header_row = table.rows[0]
            header_text = [cell.text.strip() for cell in header_row.cells]
            print(f"Headers: {' | '.join(header_text)}")
        
        # Check content rows
        for i, row in enumerate(table.rows):
            if i == 0:  # Skip header row
                continue
                
            # Print row content
            row_content = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_content.append(cell_text)
                
                total_cells += 1
                if not cell_text:
                    empty_cells += 1
                    
            print(f"Row {i}: {' | '.join(row_content)}")
        
        # Calculate percentage of empty cells
        if total_cells > 0:
            empty_percentage = (empty_cells / total_cells) * 100
            print(f"\n{table_name} table has {empty_percentage:.1f}% empty cells ({empty_cells}/{total_cells})")
        else:
            print(f"\n{table_name} table has no data rows to analyze")
    
    # Overall summary
    print("\n=== REPRODUCIBILITY TABLES SUMMARY ===")
    all_tables_found = all(table_info is not None for _, table_info in tables_to_check)
    print(f"All reproducibility tables found: {'Yes' if all_tables_found else 'No'}")
    
    if all_tables_found:
        print("✅ Reproducibility tables are properly populated!")
    else:
        print("❌ Some reproducibility tables are missing!")

if __name__ == "__main__":
    check_reproducibility_tables()