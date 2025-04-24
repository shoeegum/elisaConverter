#!/usr/bin/env python3
"""
Examine the enhanced template output to verify formatting
"""

import docx
import sys

def check_template(filename):
    """Check specific sections in the enhanced template output"""
    try:
        doc = docx.Document(filename)
        print(f"\nExamining enhanced template output: {filename}")
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        
        # Check first page content
        print("\n=== FIRST PAGE CONTENT ===")
        first_page = []
        for i, para in enumerate(doc.paragraphs[:10]):
            if para.text.strip():
                style = para.style.name if hasattr(para.style, 'name') else "None"
                first_page.append(f"Para {i}: [{style}] {para.text[:80]}...")
        
        for line in first_page:
            print(line)
            
        # Check specific sections
        sections = ["MATERIALS REQUIRED", "REAGENT PREPARATION", "ASSAY PROTOCOL"]
        for section in sections:
            section_found = False
            for i, para in enumerate(doc.paragraphs):
                if section in para.text:
                    section_found = True
                    print(f"\n=== {section} SECTION ===")
                    style = para.style.name if hasattr(para.style, 'name') else "None"
                    print(f"Header style: [{style}] {para.text}")
                    
                    # Print next few paragraphs to see formatting
                    for j in range(1, 5):
                        if i+j < len(doc.paragraphs) and doc.paragraphs[i+j].text.strip():
                            next_style = doc.paragraphs[i+j].style.name if hasattr(doc.paragraphs[i+j].style, 'name') else "None"
                            print(f"Para {i+j}: [{next_style}] {doc.paragraphs[i+j].text[:100]}...")
                    break
            
            if not section_found:
                print(f"\nSection '{section}' not found in document.")
                
        # Check tables
        print("\n=== TABLES ===")
        print(f"Total tables: {len(doc.tables)}")
        
        for i, table in enumerate(doc.tables[:2]):
            print(f"\nTable {i}:")
            print(f"  Rows: {len(table.rows)}")
            print(f"  Columns: {len(table.rows[0].cells) if table.rows else 0}")
            
            # Show headers
            if table.rows:
                header_cells = []
                for cell in table.rows[0].cells:
                    header_cells.append(cell.text)
                print(f"  Headers: {header_cells}")
                
                # Show first data row
                if len(table.rows) > 1:
                    data_cells = []
                    for cell in table.rows[1].cells:
                        data_cells.append(cell.text[:50] + "..." if len(cell.text) > 50 else cell.text)
                    print(f"  First data row: {data_cells}")
                    
        return True
    
    except Exception as e:
        print(f"Error examining document: {e}")
        return False

if __name__ == "__main__":
    if len(sys.argv) > 1:
        if sys.argv[1] == "--materials":
            # Only print materials section
            doc = docx.Document("IMSKLK1KT-20250424.docx")
            print("\n=== MATERIALS REQUIRED SECTION (FULL TEXT) ===\n")
            materials_section = False
            for para in doc.paragraphs:
                if "MATERIALS REQUIRED" in para.text:
                    materials_section = True
                    print(f"[{para.style.name}] {para.text}")
                elif materials_section and para.text.strip():
                    if "REAGENT PREPARATION" in para.text:
                        break
                    print(f"[{para.style.name}] {para.text}")
        else:
            check_template(sys.argv[1])
    else:
        check_template("IMSKLK1KT-20250424.docx")