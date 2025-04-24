#!/usr/bin/env python3
"""
Check materials section in the output document.
"""

import logging
from docx import Document
from pathlib import Path

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def check_materials(document_path="updated_output.docx"):
    """Check materials section for properly formatted bullet points."""
    doc = Document(document_path)
    logger.info(f"Checking materials section in {document_path}")
    
    # Find the materials section
    materials_section = None
    for i, para in enumerate(doc.paragraphs):
        if "MATERIALS REQUIRED" in para.text.upper():
            materials_section = i
            logger.info(f"Found materials section at paragraph {i}: {para.text}")
            break
    
    if materials_section:
        # Check the next 20 paragraphs for material content
        material_items = []
        for i in range(materials_section + 1, min(materials_section + 20, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            # Skip empty paragraphs
            if not para.text.strip():
                continue
                
            # Check if this contains material text
            style_name = para.style.name if para.style else "None"
            run_texts = [r.text for r in para.runs]
            has_bullet = any("•" in r.text for r in para.runs)
            
            logger.info(f"Paragraph {i}: '{para.text}', Style: {style_name}")
            logger.info(f"  Runs: {run_texts}")
            logger.info(f"  Has bullet: {has_bullet}")
            
            if has_bullet or style_name == "List Bullet" or para.text.strip():
                material_items.append({
                    "index": i,
                    "text": para.text,
                    "style": style_name,
                    "has_bullet": has_bullet
                })
        
        # Summarize what we found
        print(f"\nFound {len(material_items)} potential material items:")
        for item in material_items:
            print(f"  Paragraph {item['index']}: '{item['text']}', Style: {item['style']}, Has bullet: {item['has_bullet']}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        check_materials(sys.argv[1])
    else:
        check_materials()