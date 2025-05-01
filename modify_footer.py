#!/usr/bin/env python3
"""
Modify Footer Text

This script changes the footer text from 'All rights reserved' to 'Made by Sophie Gall'
and removes any copyright symbols (©) from the document footer.
"""

import logging
from pathlib import Path
import shutil
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def modify_footer_text(document_path):
    """
    Modifies the footer text in the document.
    
    Changes:
    - Replaces "All rights reserved" with "Made by Sophie Gall"
    - Removes the copyright symbol (©)
    - Adds "Made by Sophie Gall" if not present
    
    Args:
        document_path: Path to the document to modify
    """
    try:
        # Make a backup of the document
        document_path = Path(document_path)
        backup_path = document_path.with_name(f"{document_path.stem}_before_footer_change{document_path.suffix}")
        shutil.copy2(document_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Load the document
        doc = Document(document_path)
        
        # Track if we made any changes to the footer
        made_changes = False
        
        # Get a list of all sections
        sections = list(doc.sections)
        logger.info(f"Found {len(sections)} sections in the document")
        
        # Process each section's footer
        for i, section in enumerate(sections):
            # Skip if linked to previous (except the first section)
            if i > 0 and section.footer.is_linked_to_previous:
                continue
            
            logger.info(f"Processing section {i+1} footer")
                
            # Modify each paragraph in the footer
            footer_paragraphs = list(section.footer.paragraphs)
            logger.info(f"Section {i+1} has {len(footer_paragraphs)} footer paragraphs")
            
            for paragraph in footer_paragraphs:
                # Replace the text "All rights reserved" with "Made by Sophie Gall"
                if "All rights reserved" in paragraph.text:
                    paragraph.text = paragraph.text.replace("All rights reserved", "Made by Sophie Gall")
                    logger.info(f"Replaced 'All rights reserved' with 'Made by Sophie Gall'")
                    made_changes = True
                
                # Remove the copyright symbol
                if "©" in paragraph.text:
                    paragraph.text = paragraph.text.replace("©", "")
                    logger.info(f"Removed copyright symbol")
                    made_changes = True
                
                # Handle when footer has runs with different formatting
                for run in paragraph.runs:
                    if "All rights reserved" in run.text:
                        run.text = run.text.replace("All rights reserved", "Made by Sophie Gall")
                        logger.info(f"Replaced 'All rights reserved' with 'Made by Sophie Gall' in run")
                        made_changes = True
                    
                    if "©" in run.text:
                        run.text = run.text.replace("©", "")
                        logger.info(f"Removed copyright symbol from run")
                        made_changes = True
                        
            # If the footer is empty or doesn't contain "Made by Sophie Gall", add it
            has_sophie_text = False
            for paragraph in footer_paragraphs:
                if "Made by Sophie Gall" in paragraph.text:
                    has_sophie_text = True
                    break
                    
                # Also check runs
                for run in paragraph.runs:
                    if "Made by Sophie Gall" in run.text:
                        has_sophie_text = True
                        break
                        
            if not has_sophie_text:
                # If the footer is completely empty, add a new paragraph
                if len(footer_paragraphs) == 0 or all(not p.text.strip() for p in footer_paragraphs):
                    logger.info("Footer is empty, adding new paragraph with 'Made by Sophie Gall'")
                    new_para = section.footer.add_paragraph("Made by Sophie Gall")
                    made_changes = True
                else:
                    # Otherwise, append to the last paragraph
                    last_para = footer_paragraphs[-1]
                    if last_para.text.strip():
                        last_para.text += " - Made by Sophie Gall"
                    else:
                        last_para.text = "Made by Sophie Gall"
                    logger.info(f"Appended 'Made by Sophie Gall' to last paragraph: '{last_para.text}'")
                    made_changes = True
        
        # Save the document
        doc.save(document_path)
        logger.info(f"Successfully modified footer in: {document_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error modifying footer: {e}")
        return False

if __name__ == "__main__":
    # Apply to the current output document
    modify_footer_text("output_populated_template.docx")