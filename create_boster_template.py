#!/usr/bin/env python3
"""
Create Boster Template

This script creates a template for Boster ELISA kit datasheets,
with structure and placeholders specific to Boster's format.
"""

import os
import logging
from pathlib import Path

import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def create_boster_template():
    """
    Create a template DOCX file for Boster ELISA kit datasheets.
    
    Returns:
        Path to the created template
    """
    # Create a new document
    doc = Document()
    
    # Set default font for all styles
    for style in doc.styles:
        if hasattr(style, 'font'):
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
    
    # Create a title style with Calibri 36pt
    title_style = doc.styles['Title']
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(36)
    title_style.font.bold = True
    
    # Create Heading 2 style with blue color for section headings
    heading2_style = doc.styles['Heading 2']
    heading2_style.font.name = 'Calibri'
    heading2_style.font.size = Pt(12)
    heading2_style.font.color.rgb = RGBColor(0, 70, 180)  # RGB values match the required blue
    
    # Add styles for custom formatting
    calibri_normal = doc.styles.add_style('Calibri Normal', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    calibri_normal.font.name = 'Calibri'
    calibri_normal.font.size = Pt(11)
    calibri_normal.paragraph_format.space_after = Pt(6)
    calibri_normal.paragraph_format.line_spacing = 1.15
    
    # Create a custom style for 1.15 line spacing
    calibri_115 = doc.styles.add_style('Calibri115', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    calibri_115.font.name = 'Calibri'
    calibri_115.font.size = Pt(11)
    calibri_115.paragraph_format.line_spacing = 1.15
    
    # Use US Letter size (8.5 x 11 inches)
    section = doc.sections[0]
    section.page_height = Cm(27.94)  # 11 inches
    section.page_width = Cm(21.59)   # 8.5 inches
    
    # Set proper margins (1 inch on all sides)
    section.left_margin = Cm(2.54)   # 1 inch
    section.right_margin = Cm(2.54)  # 1 inch
    section.top_margin = Cm(2.54)    # 1 inch
    section.bottom_margin = Cm(2.54) # 1 inch
    
    # Add footer with Innovative Research contact info
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_para.add_run('Innovative Research, Inc.')
    run.font.name = 'Calibri'
    run.font.size = Pt(26)
    
    # Add website and contact info in footer
    left_para = footer.add_paragraph()
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = left_para.add_run('www.innov-research.com    Ph: 248.896.0145 | Fx: 248.896.0149')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    
    # Add title section placeholders
    title_para = doc.add_paragraph('{{ document_title }}', style='Title')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add the catalog and lot number information below the title
    catalog_para = doc.add_paragraph(style='Calibri Normal')
    catalog_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = catalog_para.add_run('Catalog Number: {{ catalog_number }}\n')
    run.font.bold = True
    run = catalog_para.add_run('Lot Number: {{ lot_number }}')
    run.font.bold = True
    
    # Add page break after title/catalog/lot section
    doc.add_page_break()
    
    # Add sections with placeholders
    # INTENDED USE
    doc.add_heading('INTENDED USE', level=2)
    doc.add_paragraph('{{ intended_use }}', style='Calibri115')
    
    # ASSAY PRINCIPLE
    doc.add_heading('ASSAY PRINCIPLE', level=2)
    doc.add_paragraph('{{ assay_principle }}', style='Calibri115')
    
    # TECHNICAL DETAILS
    doc.add_heading('TECHNICAL DETAILS', level=2)
    
    # Add technical details table with 2 columns
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # Add header row
    header_row = table.rows[0].cells
    header_row[0].text = 'Parameter'
    header_row[1].text = 'Characteristics'
    
    # Add data rows
    rows = table.rows
    rows[1].cells[0].text = 'Detection Range'
    rows[1].cells[1].text = '{{ detection_range }}'
    
    rows[2].cells[0].text = 'Sensitivity'
    rows[2].cells[1].text = '{{ sensitivity }}'
    
    rows[3].cells[0].text = 'Specificity'
    rows[3].cells[1].text = '{{ specificity }}'
    
    rows[4].cells[0].text = 'Cross Reactivity'
    rows[4].cells[1].text = '{{ cross_reactivity }}'
    
    # OVERVIEW
    doc.add_heading('OVERVIEW', level=2)
    
    # Add overview table with 4 columns
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    # Add header row
    header_row = table.rows[0].cells
    header_row[0].text = 'Sample Type'
    header_row[1].text = 'Detection Information'
    
    # Add data rows
    rows = table.rows
    rows[1].cells[0].text = 'Serum & Plasma'
    rows[1].cells[1].text = '{{ serum_plasma_detection }}'
    
    rows[2].cells[0].text = 'Cell Culture Supernatant'
    rows[2].cells[1].text = '{{ cell_culture_detection }}'
    
    rows[3].cells[0].text = 'Other Biological Fluids'
    rows[3].cells[1].text = '{{ other_fluids_detection }}'
    
    # BACKGROUND
    doc.add_heading('BACKGROUND', level=2)
    doc.add_paragraph('{{ background }}', style='Calibri115')
    
    # KIT COMPONENTS
    doc.add_heading('KIT COMPONENTS', level=2)
    
    # Add kit components table
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    # Add header row
    header_row = table.rows[0].cells
    header_row[0].text = 'Component'
    header_row[1].text = 'Quantity'
    
    # Add placeholder row
    rows = table.rows
    rows[1].cells[0].text = '{{ kit_components_list }}'
    rows[1].cells[1].text = '{{ kit_components_quantities }}'
    
    # MATERIALS REQUIRED BUT NOT PROVIDED
    doc.add_heading('MATERIALS REQUIRED BUT NOT PROVIDED', level=2)
    doc.add_paragraph('{{ materials_not_provided }}', style='Calibri115')
    
    # STORAGE
    doc.add_heading('STORAGE', level=2)
    doc.add_paragraph('{{ storage }}', style='Calibri115')
    
    # SAMPLE COLLECTION AND STORAGE
    doc.add_heading('SAMPLE COLLECTION AND STORAGE', level=2)
    doc.add_paragraph('{{ sample_collection }}', style='Calibri115')
    
    # REAGENT PREPARATION
    doc.add_heading('REAGENT PREPARATION', level=2)
    doc.add_paragraph('{{ reagent_preparation }}', style='Calibri115')
    
    # ASSAY PROCEDURE
    doc.add_heading('ASSAY PROCEDURE', level=2)
    doc.add_paragraph('{{ assay_procedure }}', style='Calibri115')
    
    # DATA ANALYSIS
    doc.add_heading('DATA ANALYSIS', level=2)
    doc.add_paragraph('{{ data_analysis }}', style='Calibri115')
    
    # DISCLAIMER
    doc.add_heading('DISCLAIMER', level=2)
    doc.add_paragraph('{{ disclaimer }}', style='Calibri115')
    
    # Save the template
    template_dir = Path('templates_docx')
    template_dir.mkdir(exist_ok=True)
    template_path = template_dir / 'boster_template.docx'
    doc.save(template_path)
    
    logger.info(f"Created Boster template at: {template_path}")
    return template_path

if __name__ == "__main__":
    create_boster_template()