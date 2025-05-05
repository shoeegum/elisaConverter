#!/usr/bin/env python3
"""
Modify Footer Text for Red Dot Documents

This script changes the footer text to the specified Red Dot format:
'Innovative Research, Inc.' in Calibri 26pt, right-aligned
"""

import logging
from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def modify_red_dot_footer(document_path):
    """
    Modifies the footer text in the Red Dot document.
    
    Changes:
    - Sets the footer text to "Innovative Research, Inc."
    - Uses Calibri 26pt font
    - Right-aligns the text
    
    Args:
        document_path: Path to the document to modify
    """
    try:
        # Make a backup of the document
        document_path = Path(document_path)
        backup_path = document_path.with_name(f"{document_path.stem}_before_footer_change{document_path.suffix}")
        shutil.copy2(document_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Load the document
        doc = Document(document_path)
        
        # Get a list of all sections
        sections = list(doc.sections)
        logger.info(f"Found {len(sections)} sections in the document")
        
        # Process each section's footer
        for i, section in enumerate(sections):
            # Skip if linked to previous (except the first section)
            if i > 0 and section.footer.is_linked_to_previous:
                continue
            
            logger.info(f"Processing section {i+1} footer")
            
            # Clear the existing footer
            for paragraph in list(section.footer.paragraphs):
                paragraph._element.getparent().remove(paragraph._element)
                
            # Create a new paragraph for the footer
            new_para = section.footer.add_paragraph()
            
            # Set paragraph alignment to right
            new_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            # Add text with Calibri 26pt font
            run = new_para.add_run("Innovative Research, Inc.")
            run.font.name = "Calibri"
            run.font.size = Pt(26)
            
            logger.info(f"Set Red Dot footer text: 'Innovative Research, Inc.' (Calibri 26pt, right-aligned)")
        
        # Save the document
        doc.save(document_path)
        logger.info(f"Successfully modified footer in: {document_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error modifying footer: {e}")
        return False

if __name__ == "__main__":
    # Apply to the improved Red Dot output document
    modify_red_dot_footer("improved_red_dot_output.docx")