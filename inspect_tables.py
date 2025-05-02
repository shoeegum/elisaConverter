#!/usr/bin/env python3
"""
Inspect tables in the Red Dot document.

This script examines all tables in the source document and attempts to identify their purpose.
"""

import logging
from pathlib import Path
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def inspect_tables(document_path):
    """
    Inspect tables in a document and print detailed information.

    Args:
        document_path: Path to the document to inspect
    """
    # Load the document
    doc = Document(document_path)
    
    # Count tables
    logger.info(f"Document contains {len(doc.tables)} tables")
    
    # Examine each table
    for i, table in enumerate(doc.tables):
        logger.info(f"\n=== TABLE {i} ===")
        logger.info(f"Rows: {len(table.rows)}")
        if len(table.rows) > 0:
            logger.info(f"Columns: {len(table.rows[0].cells)}")
        
        # Print the table content (first 5 rows)
        for j, row in enumerate(table.rows[:5]):
            row_text = [cell.text for cell in row.cells]
            logger.info(f"Row {j}: {row_text}")
        
        # If there are more rows, indicate that
        if len(table.rows) > 5:
            logger.info(f"... and {len(table.rows) - 5} more rows")
        
        # Try to determine the table type
        if len(table.rows) > 0:
            header_row = [cell.text.lower() for cell in table.rows[0].cells]
            header_text = " ".join(header_row)
            
            # Check for reagents/components table
            if any(keyword in header_text for keyword in ['component', 'reagent', 'kit', 'material', 'content']):
                logger.info("Table appears to be a REAGENTS/COMPONENTS table")
            
            # Check for assay procedure table
            elif any(keyword in header_text for keyword in ['assay', 'step', 'procedure', 'protocol']):
                logger.info("Table appears to be an ASSAY PROCEDURE table")
            
            # Check for standard curve table
            elif any(keyword in header_text for keyword in ['standard', 'curve', 'concentration', 'od']):
                logger.info("Table appears to be a STANDARD CURVE table")
            
            # Check for reproducibility table
            elif any(keyword in header_text for keyword in ['intra', 'inter', 'precision', 'cv%']):
                logger.info("Table appears to be a REPRODUCIBILITY table")
                
        logger.info("=" * 30)
        
    return len(doc.tables)

if __name__ == "__main__":
    # Inspect the Red Dot document tables
    inspect_tables("attached_assets/RDR-LMNB2-Hu.docx")