#!/usr/bin/env python3
"""
Check the structure of the generated document.
"""

import logging
from pathlib import Path
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def check_document_structure(document_path="output_populated_template.docx"):
    """
    Check the structure of the document and print a detailed layout of sections,
    paragraphs, and tables.
    
    Args:
        document_path: Path to the document to check
    """
    doc = Document(document_path)
    
    print(f"\n=== Document Structure for {document_path} ===\n")
    
    # Count paragraphs, tables, sections
    total_paragraphs = len(doc.paragraphs)
    total_tables = len(doc.tables)
    
    print(f"Total paragraphs: {total_paragraphs}")
    print(f"Total tables: {total_tables}")
    
    # Print an outline structure
    print("\n--- Document Outline ---\n")
    
    element_idx = 0
    table_idx = 0
    
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph
            para = doc.paragraphs[element_idx]
            text = para.text.strip()
            
            # Skip empty paragraphs in the outline
            if text:
                # Determine paragraph style/level for display
                if para.style.name.startswith('Heading 1'):
                    print(f"# {text}")
                elif para.style.name.startswith('Heading 2'):
                    print(f"## {text}")
                elif para.style.name.startswith('Heading 3'):
                    print(f"### {text}")
                elif para.style.name.startswith('Title'):
                    print(f"TITLE: {text}")
                elif len(text) > 100:
                    print(f"Para: {text[:100]}...")
                else:
                    print(f"Para: {text}")
                    
            element_idx += 1
            
        elif element.tag.endswith('tbl'):  # Table
            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                rows = len(table.rows)
                cols = len(table.rows[0].cells) if rows > 0 else 0
                
                # Extract table header or first row to identify it
                header_text = ""
                if rows > 0:
                    header_text = " | ".join([cell.text.strip() for cell in table.rows[0].cells])
                
                if len(header_text) > 80:
                    header_text = header_text[:77] + "..."
                    
                print(f"TABLE {table_idx}: {rows}×{cols} - {header_text}")
                table_idx += 1
    
    print("\n--- End of Document Outline ---")

if __name__ == "__main__":
    check_document_structure()