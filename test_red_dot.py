import sys
import os
import re
import docx
from pathlib import Path
import logging
from red_dot_template_populator import extract_red_dot_data, populate_red_dot_template

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def examine_red_dot_doc(doc_path):
    """Examine a Red Dot document to extract key information"""
    logger.info(f"Examining document: {doc_path}")
    
    try:
        doc = docx.Document(doc_path)
        logger.info(f"Document loaded successfully: {len(doc.paragraphs)} paragraphs found")
        
        # Check first 20 paragraphs for key identifiers
        for i, para in enumerate(doc.paragraphs[:30]):
            text = para.text.strip()
            if text:  # Only print non-empty paragraphs
                logger.info(f"Para {i}: {text[:100]}")
                # Look for Red Dot identifiers
                if "RED DOT" in text.upper() or "RDR-" in text.upper():
                    logger.info(f"RED DOT IDENTIFIER FOUND in paragraph {i}: {text}")
        
        # Get section headings
        section_keywords = [
            "INTENDED USE", "TEST PRINCIPLE", "REAGENTS", "MATERIALS", 
            "SAMPLE", "PROCEDURE", "CALCULATION", "SENSITIVITY", 
            "SPECIFICITY", "PRECISION", "STABILITY"
        ]
        
        logger.info("\nSearching for sections:")
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().upper()
            for keyword in section_keywords:
                if keyword in text and len(text) < 100:  # To avoid matching keywords in paragraphs
                    logger.info(f"Potential section found at paragraph {i}: '{text}'")
                    break
        
        # Check for tables
        logger.info(f"\nFound {len(doc.tables)} tables in document")
        for i, table in enumerate(doc.tables):
            if len(table.rows) > 0:
                header_row = " | ".join([cell.text.strip() for cell in table.rows[0].cells if cell.text.strip()])
                logger.info(f"Table {i}: {len(table.rows)} rows - Headers: {header_row[:100]}")
        
        return True
        
    except Exception as e:
        logger.error(f"Error examining document: {str(e)}")
        return False

def test_red_dot_extraction(doc_path):
    """Test the Red Dot data extraction function"""
    logger.info(f"Testing Red Dot extraction on: {doc_path}")
    
    try:
        # Extract data using our Red Dot extraction function
        data = extract_red_dot_data(Path(doc_path))
        
        # Log the results
        logger.info(f"Extraction completed. Results:")
        logger.info(f"Kit name: {data.get('kit_name', 'Not found')}")
        logger.info(f"Catalog number: {data.get('catalog_number', 'Not found')}")
        logger.info(f"Lot number: {data.get('lot_number', 'Not found')}")
        
        # Check if Red Dot sections were found
        if 'red_dot_sections' in data:
            logger.info("Red Dot sections found:")
            for section, content in data['red_dot_sections'].items():
                preview = content[:30] + '...' if content and len(content) > 30 else content
                logger.info(f"  - {section}: {preview}")
        else:
            logger.info("No Red Dot sections found")
            
        # Try to populate the template
        template_path = "templates_docx/red_dot_template.docx"
        output_path = "output_red_dot_test.docx"
        
        logger.info(f"Attempting to populate template: {template_path}")
        success = populate_red_dot_template(
            source_path=Path(doc_path),
            template_path=Path(template_path),
            output_path=Path(output_path)
        )
        
        if success:
            logger.info(f"Template populated successfully: {output_path}")
        else:
            logger.error("Failed to populate template")
        
        return data
    
    except Exception as e:
        logger.error(f"Error in Red Dot extraction test: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return None

if __name__ == "__main__":
    if len(sys.argv) > 1:
        doc_path = sys.argv[1]
    else:
        doc_path = "attached_assets/RDR-LMNB2-Hu.docx"
    
    logger.info("STEP 1: Examining document structure")
    examine_red_dot_doc(doc_path)
    
    logger.info("\nSTEP 2: Testing Red Dot data extraction")
    test_red_dot_extraction(doc_path)