#!/usr/bin/env python3
"""
Restructure Document

This script adds the Assay Principle section before Technical Details
while maintaining all other formatting and structure from the April 24th version.
"""

import logging
from pathlib import Path
import shutil
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def restructure_document(document_path):
    """
    Restructure the document to place Assay Principle before Technical Details.
    
    Args:
        document_path: Path to the document to modify
    """
    try:
        # Make a backup of the document
        document_path = Path(document_path)
        backup_path = document_path.with_name(f"{document_path.stem}_before_restructure{document_path.suffix}")
        shutil.copy2(document_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Load the document
        doc = Document(document_path)
        
        # Find ASSAY PRINCIPLE and TECHNICAL DETAILS sections
        assay_principle_idx = None
        assay_principle_content = []
        technical_details_idx = None
        
        # Find section indexes
        for i, para in enumerate(doc.paragraphs):
            if "ASSAY PRINCIPLE" in para.text.upper():
                assay_principle_idx = i
                logger.info(f"Found ASSAY PRINCIPLE section at paragraph {i}")
            elif "TECHNICAL DETAILS" in para.text.upper():
                technical_details_idx = i
                logger.info(f"Found TECHNICAL DETAILS section at paragraph {i}")
        
        # If we didn't find both sections, we can't continue
        if assay_principle_idx is None:
            logger.warning("Could not find ASSAY PRINCIPLE section")
            return False
            
        if technical_details_idx is None:
            logger.warning("Could not find TECHNICAL DETAILS section")
            return False
            
        # Extract ASSAY PRINCIPLE content - collect all paragraphs until the next section
        # Define section headings to check for
        section_headings = ["TECHNICAL DETAILS", "OVERVIEW", "KIT COMPONENTS", 
                           "MATERIALS REQUIRED", "STORAGE", "SAMPLE PREPARATION", 
                           "SAMPLE DILUTION", "ASSAY PROCEDURE", "DATA ANALYSIS"]
                           
        # Get ASSAY PRINCIPLE content
        assay_content_start = assay_principle_idx + 1  # Start after the heading
        assay_content = []
        
        # Collect all content until the next section heading
        for i in range(assay_content_start, len(doc.paragraphs)):
            para_text = doc.paragraphs[i].text.strip()
            # Check if this paragraph is the start of a new section
            if any(heading in para_text.upper() for heading in section_headings):
                break
            if para_text:  # Only include non-empty paragraphs
                assay_content.append((para_text, doc.paragraphs[i].style))
                
        logger.info(f"Extracted {len(assay_content)} paragraphs from ASSAY PRINCIPLE section")
        
        # Create a new document to build the restructured content
        temp_doc = Document()
        
        # Copy paragraphs up to TECHNICAL DETAILS but excluding ASSAY PRINCIPLE
        for i in range(len(doc.paragraphs)):
            # Skip the ASSAY PRINCIPLE section
            if i == assay_principle_idx:
                continue
                
            if assay_content_start <= i < assay_content_start + len(assay_content):
                continue
                
            # When we reach TECHNICAL DETAILS, insert ASSAY PRINCIPLE first
            if i == technical_details_idx:
                # Add ASSAY PRINCIPLE heading
                principle_heading = temp_doc.add_paragraph("ASSAY PRINCIPLE")
                principle_heading.style = doc.paragraphs[assay_principle_idx].style
                
                # Add ASSAY PRINCIPLE content
                for content_text, content_style in assay_content:
                    para = temp_doc.add_paragraph(content_text)
                    para.style = content_style
                    
                # Now add the TECHNICAL DETAILS paragraph
                para = temp_doc.add_paragraph(doc.paragraphs[i].text)
                para.style = doc.paragraphs[i].style
                
            else:
                # Add the paragraph as normal
                para = temp_doc.add_paragraph(doc.paragraphs[i].text)
                para.style = doc.paragraphs[i].style
        
        # Copy all tables 
        for table in doc.tables:
            # Get the dimensions of the table
            rows = len(table.rows)
            cols = len(table.rows[0].cells) if rows > 0 else 0
            
            # Create a new table with the same dimensions
            if rows > 0 and cols > 0:
                new_table = temp_doc.add_table(rows=rows, cols=cols)
                new_table.style = table.style
                
                # Copy cell content
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                            new_table.rows[i].cells[j].text = cell.text
                
        # Save the document
        temp_path = document_path.with_name(f"{document_path.stem}_temp{document_path.suffix}")
        temp_doc.save(temp_path)
        
        # Now use the proper formatting function to ensure consistent styling
        from format_document import apply_document_formatting
        apply_document_formatting(temp_path)
        
        # Replace the original with our temporary document
        shutil.copy2(temp_path, document_path)
        
        # Clean up
        if temp_path.exists():
            import os
            os.remove(temp_path)
            
        logger.info(f"Successfully restructured document to place ASSAY PRINCIPLE before TECHNICAL DETAILS: {document_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error restructuring document: {e}")
        return False

if __name__ == "__main__":
    # Restructure the current output document
    restructure_document("output_populated_template.docx")