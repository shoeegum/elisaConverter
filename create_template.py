"""
Create a proper template with placeholders from the sample document.
"""

import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Set up logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def create_innovative_template():
    """
    Create a proper Innovative Research template with proper styles and placeholders.
    """
    doc = Document()
    
    # Set document styles
    styles = doc.styles
    
    # Set default font for the entire document
    font = styles['Normal'].font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Create style for section headers as Heading 2
    if 'Heading 2' not in styles:
        heading2_style = styles.add_style('Heading 2', styles['Normal'].type)
    else:
        heading2_style = styles['Heading 2']
    
    heading2_style.font.name = 'Calibri'
    heading2_style.font.size = Pt(12)
    heading2_style.font.bold = True
    heading2_style.font.color.rgb = RGBColor(0, 0, 128)  # Dark blue
    
    # Also keep the original header style for other uses
    header_style = styles.add_style('Header Style', styles['Normal'].type)
    header_style.font.name = 'Calibri'
    header_style.font.size = Pt(12)
    header_style.font.bold = True
    header_style.font.color.rgb = RGBColor(0, 0, 128)  # Dark blue to match sample
    
    # Set paragraph spacing to match the sample document
    paragraph_format = styles['Normal'].paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(8)
    paragraph_format.line_spacing = 1.15  # Matches sample document spacing
    
    # Add the kit name title at the top - centered, bold, 14pt 
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("{{ kit_name }}")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add catalog/lot line - centered
    cat_lot_para = doc.add_paragraph()
    cat_lot_para.add_run("CATALOG NO: ").bold = True
    cat_lot_para.add_run("{{ catalog_number }}")
    cat_lot_para.add_run(" LOT NO: ").bold = True
    cat_lot_para.add_run("{{ lot_number }}")
    cat_lot_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add sections with proper formatting - exactly match the sample document
    
    # INTENDED USE - all caps, Heading 2 style
    intended_use_header = doc.add_paragraph("INTENDED USE", style='Heading 2')
    intended_use_para = doc.add_paragraph("{{ intended_use }}")
    
    # BACKGROUND - all caps, Heading 2 style
    background_header = doc.add_paragraph("BACKGROUND ON {{ kit_name }}", style='Heading 2')
    background_para = doc.add_paragraph("{{ background }}")
    
    # PRINCIPLE OF THE ASSAY - all caps, Heading 2 style
    principle_header = doc.add_paragraph("PRINCIPLE OF THE ASSAY", style='Heading 2')
    principle_para = doc.add_paragraph("{{ assay_principle }}")
    
    # OVERVIEW - all caps, Heading 2 style
    overview_header = doc.add_paragraph("OVERVIEW", style='Heading 2')
    
    # TECHNICAL DETAILS - all caps, Heading 2 style
    tech_details_header = doc.add_paragraph("TECHNICAL DETAILS", style='Heading 2')
    
    # PREPARATIONS BEFORE ASSAY - all caps, Heading 2 style
    prep_before_header = doc.add_paragraph("PREPARATIONS BEFORE ASSAY", style='Heading 2')
    
    # KIT COMPONENTS/MATERIALS PROVIDED - all caps, Heading 2 style
    reagents_header = doc.add_paragraph("KIT COMPONENTS/MATERIALS PROVIDED", style='Heading 2')
    # Add reagents table here if needed
    
    # REQUIRED MATERIALS THAT ARE NOT SUPPLIED - all caps, Heading 2 style
    materials_header = doc.add_paragraph("REQUIRED MATERIALS THAT ARE NOT SUPPLIED", style='Heading 2')
    materials_para = doc.add_paragraph("{{ required_materials }}")
    
    # TYPICAL DATA - all caps, Heading 2 style
    typical_data_header = doc.add_paragraph("TYPICAL DATA", style='Heading 2')
    
    # PRODUCT NAME ELISA STANDARD CURVE EXAMPLE - all caps, Heading 2 style
    curve_header = doc.add_paragraph("{{ kit_name }} STANDARD CURVE EXAMPLE", style='Heading 2')
    curve_para = doc.add_paragraph("This standard curve was generated for demonstration purpose only. A standard curve must be run with each assay.")
    
    # Add a table for standard curve data if needed
    curve_table = doc.add_table(rows=1, cols=2)
    curve_table.style = 'Table Grid'
    
    # Add headers to the curve table
    curve_headers = curve_table.rows[0].cells
    curve_headers[0].text = "Concentration"
    curve_headers[1].text = "OD Value"
    
    # Make headers bold
    for cell in curve_headers:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add placeholder rows with Jinja tags
    for i in range(8):  # We'll add 8 placeholder rows for standard curve
        row = curve_table.add_row().cells
        row[0].text = f"{{{{ standard_curve_table[{i}].concentration if {i} < standard_curve_table|length else '' }}}}"
        row[1].text = f"{{{{ standard_curve_table[{i}].od_value if {i} < standard_curve_table|length else '' }}}}"
    
    # INTRA/INTER-ASSAY VARIABILITY - all caps, Heading 2 style
    variability_header = doc.add_paragraph("INTRA/INTER-ASSAY VARIABILITY", style='Heading 2')
    
    # Intra-Assay Precision
    intra_para = doc.add_paragraph("Intra-Assay Precision (Precision within an assay): Three samples of known concentration were tested on one plate to assess intra-assay precision.")
    
    # Inter-Assay Precision
    inter_para = doc.add_paragraph("Inter-Assay Precision (Precision across assays): Three samples of known concentration were tested in separate assays to assess inter- assay precision.")
    
    # REPRODUCIBILITY - all caps, Heading 2 style
    repro_header = doc.add_paragraph("REPRODUCIBILITY", style='Heading 2')
    repro_para = doc.add_paragraph("*number of samples for each test n=16.")
    
    # PREPARATION BEFORE THE EXPERIMENT - all caps, Heading 2 style
    procedural_header = doc.add_paragraph("PREPARATION BEFORE THE EXPERIMENT", style='Heading 2')
    procedural_para = doc.add_paragraph("{{ procedural_notes }}")
    
    # DILUTION OF PRODUCT STANDARD - all caps, Heading 2 style
    dilution_header = doc.add_paragraph("DILUTION OF {{ kit_name }} STANDARD", style='Heading 2')
    dilution_para = doc.add_paragraph("{{ dilution_of_standard }}")
    
    # SAMPLE PREPARATION AND STORAGE - all caps, Heading 2 style
    sample_header = doc.add_paragraph("SAMPLE PREPARATION AND STORAGE", style='Heading 2')
    sample_para = doc.add_paragraph("{{ sample_collection_notes }}")
    
    # SAMPLE COLLECTION NOTES - all caps, Heading 2 style
    sample_notes_header = doc.add_paragraph("SAMPLE COLLECTION NOTES", style='Heading 2')
    
    # SAMPLE DILUTION GUIDELINE - all caps, Heading 2 style
    sample_dilution_header = doc.add_paragraph("SAMPLE DILUTION GUIDELINE", style='Heading 2')
    
    # ASSAY PROCEDURE - all caps, Heading 2 style
    assay_procedure_header = doc.add_paragraph("ASSAY PROCEDURE", style='Heading 2')
    
    # Instead of fixed number of steps, render the protocol dynamically
    assay_para = doc.add_paragraph("{{ '{% for step in assay_protocol %}' }}")
    assay_para = doc.add_paragraph("{{ '{{ step }}' }}")
    assay_para = doc.add_paragraph("{{ '{% endfor %}' }}")
    
    # DATA ANALYSIS - all caps, Heading 2 style
    data_analysis_header = doc.add_paragraph("DATA ANALYSIS", style='Heading 2')
    data_analysis_para = doc.add_paragraph("{{ data_analysis }}")
    
    # DISCLAIMER - all caps, Heading 2 style
    disclaimer_header = doc.add_paragraph("DISCLAIMER", style='Heading 2')
    disclaimer_para = doc.add_paragraph("This material is sold for in-vitro use only in manufacturing and research. This material is not suitable for human use. It is the responsibility of the user to undertake sufficient verification and testing to determine the suitability of each product's application. The statements herein are offered for informational purposes only and are intended to be used solely for your consideration, investigation and verification.")
    
    # Add footer
    footer_section = doc.sections[0]
    footer = footer_section.footer
    
    # Add footer paragraphs
    footer_website = footer.paragraphs[0]
    footer_website.text = "www.innov-research.com"
    footer_website.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    footer_website.style = 'Header Style'
    
    # Create a new paragraph for contact info
    footer_contact = footer.add_paragraph("Ph: 248.896.0145 | Fx: 248.896.0149")
    footer_contact.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Add company name to the right side
    footer_company = footer.add_paragraph("Innovative Research, Inc.")
    footer_company.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Save the template
    template_path = Path('templates_docx/innovative_exact_template.docx')
    doc.save(template_path)
    
    logger.info(f"Created template at {template_path}")
    return template_path

if __name__ == "__main__":
    create_innovative_template()