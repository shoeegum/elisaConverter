#!/usr/bin/env python3
"""
Check the content of the OVERVIEW table to identify which cells are empty.
"""

from docx import Document

def check_overview_table(document_path="output_populated_template.docx"):
    """
    Check the content of the OVERVIEW table and identify which cells are empty.
    
    Args:
        document_path: Path to the document to check
    """
    doc = Document(document_path)
    
    # Overview table is typically the second table (index 1)
    if len(doc.tables) > 1:
        table = doc.tables[1]
        print("=== OVERVIEW TABLE CONTENT ===")
        
        empty_cells = 0
        for i, row in enumerate(table.rows):
            if len(row.cells) >= 2:
                header = row.cells[0].text.strip()
                value = row.cells[1].text.strip()
                print(f"Row {i}: '{header}': '{value}'")
                
                if not value:
                    empty_cells += 1
                    print(f"  ⚠️ Empty value for '{header}'")
        
        total_rows = len(table.rows)
        if total_rows > 0:
            empty_percentage = (empty_cells / total_rows) * 100
            print(f"\nOverview table has {empty_percentage:.1f}% empty cells ({empty_cells}/{total_rows})")
    else:
        print("OVERVIEW table not found in the document (expects it at index 1)")

if __name__ == "__main__":
    check_overview_table()