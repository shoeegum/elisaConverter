#!/usr/bin/env python3
"""
Fix Reagents Table - Post Processing

This script runs after a Red Dot document has been generated to convert the
REAGENTS PROVIDED section from pipe-separated text format to a proper Word table.
It should be run on the final output document.
"""

import logging
import shutil
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def convert_text_to_table(document_path):
    """
    Convert pipe-separated text in REAGENTS PROVIDED section to a proper Word table.
    
    Args:
        document_path: Path to the document to modify
    
    Returns:
        True if successful, False otherwise
    """
    try:
        # Create a backup of the document
        document_path = Path(document_path)
        backup_path = document_path.with_name(f"{document_path.stem}_before_table_conversion{document_path.suffix}")
        shutil.copy2(document_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Load the document
        doc = Document(document_path)
        
        # Find the REAGENTS PROVIDED section
        reagents_section_idx = None
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == "REAGENTS PROVIDED":
                reagents_section_idx = i
                logger.info(f"Found REAGENTS PROVIDED section at paragraph {i}")
                break
        
        if reagents_section_idx is None:
            logger.warning("REAGENTS PROVIDED section not found in document")
            return False
        
        # Find the content paragraph after the section heading
        content_idx = reagents_section_idx + 1
        if content_idx >= len(doc.paragraphs):
            logger.warning("No content paragraph after REAGENTS PROVIDED")
            return False
        
        # Get the content text
        content_text = doc.paragraphs[content_idx].text
        logger.info(f"Current REAGENTS PROVIDED content: {content_text[:100]}...")
        
        # Check if the content contains a pipe-separated table format
        if "|" in content_text:
            logger.info("Detected pipe-separated table format in REAGENTS PROVIDED")
            
            # Parse the text into table rows
            rows = []
            current_row_data = []
            
            for line in content_text.split('\n'):
                # Skip separator lines (dashed lines)
                if line.strip() and not line.strip().startswith('-----'):
                    # Pipes indicate cells on this line
                    if "|" in line:
                        cells = [cell.strip() for cell in line.split('|')]
                        # Filter out empty cells
                        cells = [cell for cell in cells if cell]
                        if cells:
                            current_row_data = cells
                            rows.append(current_row_data)
            
            # Get the maximum number of columns
            max_cols = max(len(row) for row in rows) if rows else 0
            logger.info(f"Extracted {len(rows)} rows with max {max_cols} columns")
            
            if rows and max_cols > 0:
                # Create a new table after the heading paragraph
                table = doc.add_table(rows=len(rows), cols=max_cols)
                table.style = 'Table Grid'
                
                # Populate the table with the extracted data
                for i, row_data in enumerate(rows):
                    for j, cell_text in enumerate(row_data):
                        if j < max_cols:  # Make sure we don't exceed the number of columns
                            cell = table.rows[i].cells[j]
                            cell.text = cell_text
                            
                            # Format the text in the cell (Calibri, 11pt)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(11)
                
                # Set table width to page width
                table.autofit = False
                for col in table.columns:
                    col.width = Cm(2.5)  # Set each column to 2.5 cm
                
                # Apply a highlight to the header row if it exists
                if len(rows) > 0:
                    for cell in table.rows[0].cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                
                # Remove the original content paragraph
                p = doc.paragraphs[content_idx]
                p._element.getparent().remove(p._element)
                
                # Save the modified document
                doc.save(document_path)
                logger.info(f"Successfully converted text to table in: {document_path}")
                return True
            else:
                logger.warning("Failed to extract valid table data from content")
                return False
        else:
            logger.info("REAGENTS PROVIDED content does not contain pipe-separated text")
            return False
            
    except Exception as e:
        logger.error(f"Error converting text to table: {e}")
        return False

if __name__ == "__main__":
    import sys
    
    # Use the provided file path or default to improved_red_dot_output.docx
    if len(sys.argv) > 1:
        document_path = sys.argv[1]
    else:
        document_path = "improved_red_dot_output.docx"
    
    # Run the conversion
    convert_text_to_table(document_path)