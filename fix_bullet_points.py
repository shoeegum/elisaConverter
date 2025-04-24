#!/usr/bin/env python3
"""
Fix bullet points in the enhanced template for material items.
"""

import logging
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
from pathlib import Path

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def add_bullet_to_paragraph(paragraph):
    """Add a bullet character to the start of a paragraph."""
    # First remove existing runs
    for _ in range(len(paragraph.runs)):
        paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
    
    # Now add just the bullet
    run = paragraph.add_run("• ")
    run.font.size = Pt(11)
    
def fix_template_bullet_points(template_path="templates_docx/enhanced_template.docx"):
    """Fix bullet points in the template for materials section."""
    doc = Document(template_path)
    backup_path = Path("templates_docx/enhanced_template_backup_bullets.docx")
    
    # First make a backup
    import shutil
    shutil.copy(template_path, backup_path)
    logger.info(f"Created backup at {backup_path}")
    
    # Find the materials section
    materials_section = None
    for i, para in enumerate(doc.paragraphs):
        if "MATERIALS REQUIRED" in para.text.upper():
            materials_section = i
            logger.info(f"Found materials section at paragraph {i}")
            break
    
    if materials_section is not None:
        # Check the paragraphs after the materials section title
        material_paragraphs = []
        for i in range(materials_section + 1, min(materials_section + 25, len(doc.paragraphs))):
            para = doc.paragraphs[i]
            if "{{ req_material_" in para.text:
                material_paragraphs.append(i)
                # Check if paragraph has a style applied
                style_name = para.style.name if para.style else "No Style"
                logger.info(f"Found material paragraph {i}: {para.text} (Style: {style_name})")
        
        # Now process material paragraphs
        for i in material_paragraphs:
            para = doc.paragraphs[i]
            # Skip if paragraph contains jinja control statements
            if "{%" in para.text:
                continue
                
            # Clear any existing runs in the paragraph (to avoid duplicating content)
            for _ in range(len(para.runs)):
                para.runs[0]._element.getparent().remove(para.runs[0]._element)
            
            # Add bullet point and set style
            para.style = doc.styles['List Bullet']
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.first_line_indent = Inches(0)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # Add bullet character and variable text
            add_bullet_to_paragraph(para)
            variable_name = para.text.strip()
            para.add_run(variable_name)
            
            logger.info(f"Updated paragraph {i} with bullet point")
    
    # Save the updated document
    doc.save(template_path)
    logger.info(f"Saved updated template to {template_path}")
    return True

if __name__ == "__main__":
    fix_template_bullet_points()