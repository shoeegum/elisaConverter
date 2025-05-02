#!/usr/bin/env python3
"""
Check the position of tables relative to sections in a document.
"""

import logging
import sys
from pathlib import Path
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def check_table_position(document_path):
    """
    Check the position of tables relative to sections in a document.
    
    Args:
        document_path: Path to the document to check
    """
    try:
        # Load the document
        doc = Document(document_path)
        
        # Map each paragraph to its index
        para_map = {}
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if para_text:
                para_map[i] = para_text
        
        # Find section headings
        section_headings = {}
        for i, para in enumerate(doc.paragraphs):
            if para.style.name.startswith('Heading'):
                section_headings[para.text.strip()] = i
        
        # Find section ranges
        section_ranges = {}
        sections = list(section_headings.keys())
        for i, section in enumerate(sections):
            start = section_headings[section]
            if i < len(sections) - 1:
                end = section_headings[sections[i+1]] - 1
            else:
                end = len(doc.paragraphs) - 1
            section_ranges[section] = (start, end)
        
        # Find tables
        tables = doc.tables
        
        # Check the XML structure to find where tables are positioned
        import inspect
        from lxml import etree
        
        body = doc._body._body
        elements = []
        
        # Walk through the body elements and identify paragraphs and tables
        for i, element in enumerate(body):
            if element.tag.endswith('}p'):
                # Paragraph
                text = "".join([t.text for t in element.findall('.//{*}t') if t.text])
                if text.strip():
                    elements.append(('paragraph', text.strip()))
            elif element.tag.endswith('}tbl'):
                # Table
                elements.append(('table', f"Table with {len(element.findall('.//{*}tr'))} rows"))
        
        # Print the document structure
        logger.info("=== Document Structure ===")
        for i, (element_type, content) in enumerate(elements):
            prefix = ""
            if element_type == 'paragraph':
                if i > 0 and elements[i-1][0] == 'paragraph':
                    prefix = "│ "
                logger.info(f"{prefix}Paragraph: {content[:50]}...")
            else:
                logger.info(f"TABLE: {content}")
                
        # Specifically look for REAGENTS PROVIDED section and tables
        reagents_idx = None
        for i, (element_type, content) in enumerate(elements):
            if element_type == 'paragraph' and content == 'REAGENTS PROVIDED':
                reagents_idx = i
                logger.info(f"Found REAGENTS PROVIDED at position {i}")
                
        if reagents_idx is not None:
            # Look at the next few elements
            for i in range(reagents_idx + 1, min(reagents_idx + 5, len(elements))):
                logger.info(f"Element after REAGENTS PROVIDED at position {i}: {elements[i][0]} - {elements[i][1][:50]}...")
            
            # Find the first table after REAGENTS PROVIDED
            for i in range(reagents_idx + 1, len(elements)):
                if elements[i][0] == 'table':
                    logger.info(f"Found table {i - reagents_idx} positions after REAGENTS PROVIDED")
                    break
        
    except Exception as e:
        logger.error(f"Error checking table position: {e}")

if __name__ == "__main__":
    # Use the provided file path or default
    if len(sys.argv) > 1:
        document_path = sys.argv[1]
    else:
        document_path = "complete_red_dot_output.docx"
    
    # Check the table position
    check_table_position(document_path)