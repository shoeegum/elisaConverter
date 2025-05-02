#!/usr/bin/env python3
"""
Fix Red Dot Document Issues

This script addresses two issues with Red Dot documents:
1. Moves the REAGENTS PROVIDED table to the correct position in the document
2. Replaces all instances of 'Reddot Biotech INC.' with 'Innovative Research, Inc.'
   and 'Reddot Biotech' with 'Innovative Research'
"""

import logging
import shutil
import sys
import re
from pathlib import Path
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def fix_red_dot_document(document_path):
    """
    Fix two issues with Red Dot documents:
    1. Move the REAGENTS PROVIDED table to the correct position
    2. Replace Reddot company names with Innovative Research
    
    Args:
        document_path: Path to the document to modify
    
    Returns:
        True if successful, False otherwise
    """
    try:
        # Create a backup of the document
        document_path = Path(document_path)
        backup_path = document_path.with_name(f"{document_path.stem}_before_fixes{document_path.suffix}")
        shutil.copy2(document_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Load the document
        doc = Document(document_path)
        
        # Replace all instances of company names
        replacements = [
            ('Reddot Biotech INC.', 'Innovative Research, Inc.'),
            ('Reddot Biotech', 'Innovative Research'),  # Must be after the more specific replacement
        ]
        
        count = 0
        # Replace in paragraphs
        for para in doc.paragraphs:
            original_text = para.text
            modified_text = original_text
            for old_text, new_text in replacements:
                if old_text in modified_text:
                    modified_text = modified_text.replace(old_text, new_text)
            
            if modified_text != original_text:
                para.text = modified_text
                count += 1
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original_text = para.text
                        modified_text = original_text
                        for old_text, new_text in replacements:
                            if old_text in modified_text:
                                modified_text = modified_text.replace(old_text, new_text)
                        
                        if modified_text != original_text:
                            para.text = modified_text
                            count += 1
        
        logger.info(f"Replaced {count} instances of company names")
        
        # Find the REAGENTS PROVIDED section
        reagents_section_idx = None
        reagents_table_idx = None
        
        # Find the REAGENTS PROVIDED section
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == "REAGENTS PROVIDED":
                reagents_section_idx = i
                logger.info(f"Found REAGENTS PROVIDED section at paragraph {i}")
                break
        
        if reagents_section_idx is None:
            logger.warning("REAGENTS PROVIDED section not found")
            doc.save(document_path)
            return False
        
        # Find any table with "Reagents" in the header row
        for i, table in enumerate(doc.tables):
            if len(table.rows) > 0:
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                if 'Reagents' in header_cells and 'Quantity' in header_cells:
                    reagents_table_idx = i
                    logger.info(f"Found reagents table at index {i}")
                    break
        
        if reagents_table_idx is None:
            logger.warning("Reagents table not found")
            doc.save(document_path)
            return True  # Still return True since we fixed company names
        
        # Now we need to move the table to the correct position
        # Extract the table data first
        table = doc.tables[reagents_table_idx]
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        
        logger.info(f"Extracted table data: {len(table_data)} rows")
        
        # Get the XML element for the table and remove it
        element = table._element
        element.getparent().remove(element)
        
        # Get the paragraph right after the REAGENTS PROVIDED heading
        target_para_idx = reagents_section_idx + 1
        if target_para_idx >= len(doc.paragraphs):
            # No paragraph after the section heading, add one
            logger.info("Adding a new paragraph after REAGENTS PROVIDED")
            doc.add_paragraph()
            target_para_idx = reagents_section_idx + 1
        
        # Get reference to the target paragraph where we'll insert the table
        target_para = doc.paragraphs[target_para_idx]
        
        # If the paragraph has content, we need to be more careful
        # We'll insert the table before this paragraph
        if target_para.text.strip():
            logger.info(f"Target paragraph has content: '{target_para.text}'")
            
            # Clear the paragraph content (we'll move the table before it)
            target_para.text = ""
        
        # Create a new table at the target position
        new_table = doc.add_table(rows=1, cols=len(table_data[0]))
        new_table.style = 'Table Grid'
        
        # Populate the header row
        for i, cell_text in enumerate(table_data[0]):
            new_table.cell(0, i).text = cell_text
        
        # Populate the rest of the table
        for i in range(1, len(table_data)):
            row_cells = new_table.add_row().cells
            for j, cell_text in enumerate(table_data[i]):
                row_cells[j].text = cell_text
        
        # Format the table headings
        for cell in new_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Save the changes
        doc.save(document_path)
        logger.info(f"Successfully fixed and saved document: {document_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error fixing document: {e}")
        return False

if __name__ == "__main__":
    # Use the provided file path or default to complete_red_dot_output.docx
    if len(sys.argv) > 1:
        document_path = sys.argv[1]
    else:
        document_path = "complete_red_dot_output.docx"
    
    # Run the fix
    success = fix_red_dot_document(document_path)
    if success:
        logger.info("Fixes applied successfully")
    else:
        logger.error("Failed to apply some fixes")