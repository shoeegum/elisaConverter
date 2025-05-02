#!/usr/bin/env python3
"""
Update Red Dot Template Structure

This script modifies the Red Dot template to:
1. Ensure the REAGENTS PROVIDED section has a proper placeholder for the table
2. Fix variable names and placeholders in the template
3. Add proper variable references for table insertion
"""

import logging
import shutil
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def update_red_dot_template(template_path):
    """
    Update the Red Dot template structure for proper table placement.
    
    Args:
        template_path: Path to the template document to modify
    
    Returns:
        True if successful, False otherwise
    """
    try:
        # Create a backup of the document
        template_path = Path(template_path)
        backup_path = template_path.with_name(f"{template_path.stem}_before_structure_update{template_path.suffix}")
        shutil.copy2(template_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Load the document
        doc = Document(template_path)
        
        # Find the REAGENTS PROVIDED section
        reagents_section_idx = None
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == "REAGENTS PROVIDED":
                reagents_section_idx = i
                logger.info(f"Found REAGENTS PROVIDED section at paragraph {i}")
                break
        
        if reagents_section_idx is None:
            logger.warning("REAGENTS PROVIDED section not found in template")
            return False
        
        # Find/update content after REAGENTS PROVIDED section
        content_idx = reagents_section_idx + 1
        placeholder_found = False
        
        # If there's a paragraph after the section heading
        if content_idx < len(doc.paragraphs):
            content = doc.paragraphs[content_idx].text.strip()
            logger.info(f"Content after REAGENTS PROVIDED: '{content}'")
            
            # Check if there's already a placeholder
            if "{{" in content and "}}" in content:
                placeholder_found = True
                
                # Update the placeholder to use a table specifically
                p = doc.paragraphs[content_idx]
                p.text = ""  # Clear existing content
                
                # Add a special placeholder for the table
                p.add_run("{{ reagents_table_placeholder }}")
                logger.info("Updated placeholder for table insertion")
        
        # If no suitable placeholder found, add one
        if not placeholder_found:
            # Add a new paragraph after REAGENTS PROVIDED
            # This is a bit tricky - we need to find the paragraph XML element and add after it
            p_element = doc.paragraphs[reagents_section_idx]._element
            new_p = OxmlElement('w:p')
            p_element.addnext(new_p)
            
            # Create a paragraph wrapper around the new p element
            new_para = doc.paragraphs[reagents_section_idx + 1]
            new_para.text = "{{ reagents_table_placeholder }}"
            logger.info("Added new placeholder for table insertion")
        
        # Save the modified template
        doc.save(template_path)
        logger.info(f"Successfully updated REAGENTS PROVIDED section in template: {template_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error updating template structure: {e}")
        return False

def fix_company_names_in_template(template_path):
    """
    Replace all instances of Reddot company names with Innovative Research in the template.
    
    Args:
        template_path: Path to the template document to modify
    
    Returns:
        True if successful, False otherwise
    """
    try:
        # Load the document (no need for backup, already done in main function)
        doc = Document(template_path)
        
        replacements = [
            ('Reddot Biotech INC.', 'Innovative Research, Inc.'),
            ('Reddot Biotech', 'Innovative Research'),  # Must be after the more specific replacement
        ]
        
        count = 0
        for para in doc.paragraphs:
            for old_text, new_text in replacements:
                if old_text in para.text:
                    para.text = para.text.replace(old_text, new_text)
                    count += 1
        
        # Replace in tables too
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old_text, new_text in replacements:
                            if old_text in para.text:
                                para.text = para.text.replace(old_text, new_text)
                                count += 1
        
        # Save if changes were made
        if count > 0:
            logger.info(f"Replaced {count} instances of company names in template")
            doc.save(template_path)
            
        return True
            
    except Exception as e:
        logger.error(f"Error fixing company names in template: {e}")
        return False

if __name__ == "__main__":
    # Use the provided file path or default to the standard template
    if len(sys.argv) > 1:
        template_path = sys.argv[1]
    else:
        template_path = "templates_docx/enhanced_red_dot_template.docx"
    
    # Update the template
    if update_red_dot_template(template_path):
        # Also fix any company names in the template
        fix_company_names_in_template(template_path)