#!/usr/bin/env python3
"""
Fix the enhanced template to ensure material bullet points are correctly formatted.
This script updates the template's material section to use proper bullet point formatting.
"""

import logging
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor
from pathlib import Path

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def fix_template_bullets(template_path="templates_docx/enhanced_template.docx", output_path="templates_docx/enhanced_template_fixed.docx"):
    """
    Fix the enhanced template to ensure bullet points are correctly displayed
    in the materials section.
    """
    # Load the template
    doc = Document(template_path)
    logger.info(f"Loaded template: {template_path}")
    
    # Find the materials section
    materials_section_index = None
    for i, para in enumerate(doc.paragraphs):
        if "MATERIALS REQUIRED" in para.text.upper():
            materials_section_index = i
            logger.info(f"Found materials section at paragraph {i}: {para.text}")
            break
    
    if materials_section_index is None:
        logger.error("Could not find materials section in template")
        return False
    
    # Find any existing bullet point paragraphs and note their indices for removal
    paragraphs_to_remove = []
    for i in range(materials_section_index + 1, min(materials_section_index + 20, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        # Look for paragraphs that have both "•" and "{{" in them
        if "•" in para.text and "{{" in para.text:
            logger.info(f"Found bullet paragraph to replace at index {i}: {para.text}")
            paragraphs_to_remove.append(i)
    
    # We can't directly remove paragraphs, as that would mess up the indexing
    # So we'll mark what to clear, then clear them
    for i in sorted(paragraphs_to_remove, reverse=True):
        # Clear the text but keep the paragraph
        logger.info(f"Clearing paragraph {i}")
        for run in doc.paragraphs[i].runs:
            run.text = ""
    
    # Add new bullet point paragraphs with proper format
    for i in range(1, 11):
        # Create a bullet point paragraph
        new_para = doc.add_paragraph(style='List Bullet')
        new_para.paragraph_format.left_indent = Inches(0.25)
        new_para.paragraph_format.first_line_indent = Inches(0)
        new_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Add the bullet character
        bullet_run = new_para.add_run("• ")
        
        # Add the template placeholder
        variable_run = new_para.add_run(f"{{{{ req_material_{i} }}}}")
        
        # Only show if not empty
        condition_run = new_para.add_run(f"{{%if not req_material_{i}%}}{{{{''}}}}{{%endif%}}")
        condition_run.font.color.rgb = RGBColor(200, 200, 200)  # Light gray
    
    # Save the modified template
    doc.save(output_path)
    logger.info(f"Saved modified template to {output_path}")
    
    # Create a symlink or copy to replace the original
    if output_path != template_path:
        import shutil
        shutil.copy2(output_path, template_path)
        logger.info(f"Copied modified template to original location: {template_path}")
    
    return True

if __name__ == "__main__":
    fix_template_bullets()
    print("Fixed bullet points in the enhanced template. Try regenerating the document now.")