#!/usr/bin/env python3
"""
Check the Red Dot template structure and placeholder configuration.
"""

import logging
from pathlib import Path
from docx import Document
import re

# Configure logging
logging.basicConfig(level=logging.INFO,
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def check_red_dot_template(template_path):
    """Check the Red Dot template for placeholders and structure."""
    # Load the document
    doc = Document(template_path)
    
    # Find all section headings
    logger.info("=== Template Structure ===")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() and para.text.strip().isupper() and len(para.text.strip()) < 50:
            logger.info(f"Section at P{i}: {para.text}")
            
            # Check for placeholders in the next paragraph
            if i+1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i+1].text
                if "{{" in next_para and "}}" in next_para:
                    logger.info(f"  - Found placeholder: {next_para}")
    
    # Also look for placeholders in the document
    placeholders = set()
    for para in doc.paragraphs:
        matches = re.findall(r'\{\{([^}]+)\}\}', para.text)
        for match in matches:
            placeholders.add(match.strip())
    
    logger.info("\n=== Template Placeholders ===")
    for placeholder in sorted(placeholders):
        logger.info(f"- {placeholder}")
    
    # Check for tables in the template
    logger.info("\n=== Template Tables ===")
    for i, table in enumerate(doc.tables):
        logger.info(f"Table {i+1}: {len(table.rows)} rows x {len(table.columns) if table.rows else 0} columns")
        
        # Check for placeholders in table cells
        table_placeholders = set()
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(r'\{\{([^}]+)\}\}', cell.text)
                for match in matches:
                    table_placeholders.add(match.strip())
        
        if table_placeholders:
            logger.info(f"  - Table placeholders: {', '.join(sorted(table_placeholders))}")

if __name__ == "__main__":
    import sys
    
    template_path = "templates_docx/enhanced_red_dot_template.docx"
    
    if len(sys.argv) > 1:
        template_path = sys.argv[1]
    
    check_red_dot_template(template_path)