#!/usr/bin/env python3
"""
Updated Template Populator for ELISA Kit Datasheets

This module extends the EnhancedTemplatePopulator to add:
1. Improved Sample Preparation and Storage section with a proper table
2. Shortened Sample Dilution Guideline section
"""

import logging
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple
import re

from docx import Document
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docxtpl import DocxTemplate

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def update_template_populator(
    input_document: Path,
    template_path: Path,
    output_path: Path,
    kit_name: Optional[str] = None,
    catalog_number: Optional[str] = None,
    lot_number: Optional[str] = None
) -> None:
    """
    Process ELISA datasheet by extracting data, populating template, and fixing sample sections.
    
    Args:
        input_document: Path to the input ELISA datasheet
        template_path: Path to the enhanced template
        output_path: Path where the output will be saved
        kit_name: Optional kit name provided by user
        catalog_number: Optional catalog number provided by user
        lot_number: Optional lot number provided by user
    """
    # Import here to avoid circular imports
    from elisa_parser import ELISADatasheetParser
    from template_populator_enhanced import TemplatePopulator
    
    # Create parser and template populator instances
    parser = ELISADatasheetParser(input_document)
    populator = TemplatePopulator(template_path)
    
    try:
        # Parse the ELISA datasheet
        extracted_data = parser.extract_data()
        
        # Populate the template with extracted data
        populator.populate(
            extracted_data, 
            output_path, 
            kit_name, 
            catalog_number, 
            lot_number
        )
        
        # Fix the Sample Preparation and Sample Dilution sections
        fix_sample_sections(output_path)
        
        logger.info(f"Successfully processed document with updated sample sections: {output_path}")
        
    except Exception as e:
        logger.error(f"Error processing document: {e}")
        raise

def fix_sample_sections(document_path: Path) -> None:
    """
    Fix the Sample Preparation and Sample Dilution sections in the document.
    
    Args:
        document_path: Path to the document to fix
    """
    try:
        # Make a backup copy
        backup_path = document_path.with_name(f"{document_path.stem}_backup{document_path.suffix}")
        import shutil
        shutil.copy2(document_path, backup_path)
        
        # Load the document
        doc = Document(document_path)
        
        # Find the Sample Preparation and Sample Dilution sections
        sections = {}
        section_names = ["SAMPLE PREPARATION AND STORAGE", "SAMPLE DILUTION GUIDELINE", "ASSAY PROCEDURE", "ASSAY PROTOCOL"]
        section_indices = {}
        
        # Track tables and their positions
        table_positions = []
        
        # Find all section positions and table positions
        para_count = 0
        table_count = 0
        current_position = 0
        
        # First pass: find all sections and tables with their positions
        for element in doc.element.body:
            if element.tag.endswith('p'):  # This is a paragraph
                para = doc.paragraphs[para_count]
                text = para.text.strip().upper()
                para_count += 1
                current_position += 1
                
                # Check if this is a section we're interested in
                for section_name in section_names:
                    if section_name in text:
                        section_indices[section_name] = (para_count - 1, current_position)
                        break
                        
            elif element.tag.endswith('tbl'):  # This is a table
                table_positions.append((table_count, current_position))
                table_count += 1
                current_position += 1
        
        # Extract section positions
        sample_prep_position = section_indices.get("SAMPLE PREPARATION AND STORAGE")
        sample_dilution_position = section_indices.get("SAMPLE DILUTION GUIDELINE")
        assay_procedure_position = section_indices.get("ASSAY PROCEDURE") or section_indices.get("ASSAY PROTOCOL")
        
        if not sample_prep_position:
            logger.warning("Could not find SAMPLE PREPARATION AND STORAGE section")
            return
            
        if not sample_dilution_position:
            logger.warning("Could not find SAMPLE DILUTION GUIDELINE section")
            return
            
        if not assay_procedure_position:
            logger.warning("Could not find ASSAY PROCEDURE section")
            return
        
        # Get paragraph index and position for each section
        sample_prep_idx, sample_prep_pos = sample_prep_position
        sample_dilution_idx, sample_dilution_pos = sample_dilution_position
        assay_procedure_idx, assay_procedure_pos = assay_procedure_position
        
        logger.info(f"Found SAMPLE PREPARATION AND STORAGE at paragraph {sample_prep_idx}")
        logger.info(f"Found SAMPLE DILUTION GUIDELINE at paragraph {sample_dilution_idx}")
        logger.info(f"Found ASSAY PROCEDURE at paragraph {assay_procedure_idx}")
        
        # Keep track of which tables to preserve
        tables_to_preserve = {}
        
        # Identify tables that need to be preserved (those not between sections we're modifying)
        for table_idx, table_pos in table_positions:
            if table_pos < sample_prep_pos:
                tables_to_preserve[table_idx] = "before_sample_prep"
            elif table_pos >= assay_procedure_pos:
                tables_to_preserve[table_idx] = "after_assay_procedure"
                
        logger.info(f"Tables to preserve: {tables_to_preserve}")
        
        # Create a temporary document with our changes
        temp_path = document_path.with_name(f"{document_path.stem}_temp{document_path.suffix}")
        temp_doc = Document()
        
        # 1. Copy all content up to SAMPLE PREPARATION AND STORAGE
        for i in range(sample_prep_idx + 1):
            para = doc.paragraphs[i]
            new_para = temp_doc.add_paragraph(para.text)
            new_para.style = para.style
            
        # 2. Add our customized sample preparation content
        logger.info("Restructuring SAMPLE PREPARATION AND STORAGE section")
        temp_doc.add_paragraph("These sample collection instructions and storage conditions are intended as a general guideline. Sample stability has not been evaluated.")
        temp_doc.add_paragraph("")
        
        # Add SAMPLE COLLECTION NOTES
        sample_notes_para = temp_doc.add_paragraph("SAMPLE COLLECTION NOTES")
        sample_notes_para.style = 'Heading 3'
        
        # Add collection notes content
        temp_doc.add_paragraph("Innovative Research recommends that samples are used immediately upon preparation.")
        temp_doc.add_paragraph("Avoid repeated freeze-thaw cycles for all samples.")
        temp_doc.add_paragraph("Samples should be brought to room temperature (18-25°C) before performing the assay.")
        temp_doc.add_paragraph("")
        
        # Add a table for sample types
        table = temp_doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        # Set the table header
        table.cell(0, 0).text = "Sample Type"
        table.cell(0, 1).text = "Collection and Handling"
        
        # Set the table content
        table.cell(1, 0).text = "Cell Culture Supernatant"
        table.cell(1, 1).text = "Centrifuge at 1000 × g for 10 minutes to remove insoluble particulates. Collect supernatant."
        
        table.cell(2, 0).text = "Serum"
        table.cell(2, 1).text = "Use a serum separator tube (SST). Allow samples to clot for 30 minutes before centrifugation for 15 minutes at approximately 1000 × g. Remove serum and assay immediately or store samples at -20°C."
        
        table.cell(3, 0).text = "Plasma"
        table.cell(3, 1).text = "Collect plasma using EDTA or heparin as an anticoagulant. Centrifuge samples for 15 minutes at 1000 × g within 30 minutes of collection. Store samples at -20°C."
        
        table.cell(4, 0).text = "Cell Lysates"
        table.cell(4, 1).text = "Collect cells and rinse with ice-cold PBS. Homogenize at 1×10^7/ml in PBS with a protease inhibitor cocktail. Freeze/thaw 3 times. Centrifuge at 10,000×g for 10 min at 4°C. Aliquot the supernatant for testing and store at -80°C."
        
        # 3. Add customized Sample Dilution Guideline section
        logger.info("Restructuring SAMPLE DILUTION GUIDELINE section")
        
        dilution_para = temp_doc.add_paragraph("SAMPLE DILUTION GUIDELINE")
        dilution_para.style = 'Heading 2'
        
        # Add dilution guideline content
        temp_doc.add_paragraph("To inspect the validity of experimental operation and the appropriateness of sample dilution proportion, it is recommended to test all plates with the provided samples. Dilute the sample so the expected concentration falls near the middle of the standard curve range.")
        
        # 4. Add all content from the ASSAY PROCEDURE section to the end
        for i in range(assay_procedure_idx, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            new_para = temp_doc.add_paragraph(para.text)
            new_para.style = para.style
            
        # 5. Now add any "after_assay_procedure" tables
        tables_added = 0
        for table_idx, position in tables_to_preserve.items():
            if position == "after_assay_procedure":
                # Get the table from the original document
                orig_table = doc.tables[table_idx]
                
                # Create a new table with same dimensions
                rows = len(orig_table.rows)
                cols = len(orig_table.rows[0].cells) if rows > 0 else 0
                
                if rows > 0 and cols > 0:
                    new_table = temp_doc.add_table(rows=rows, cols=cols)
                    new_table.style = orig_table.style
                    
                    # Copy cell content
                    for i, row in enumerate(orig_table.rows):
                        for j, cell in enumerate(row.cells):
                            if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                                new_table.rows[i].cells[j].text = cell.text
                    
                    tables_added += 1
                    logger.info(f"Added table {table_idx} ({rows}x{cols}) from position {position}")
        
        # Save the temporary document
        temp_doc.save(temp_path)
        
        # Replace the original with our temporary document
        shutil.copy2(temp_path, document_path)
        
        # Clean up
        if temp_path.exists():
            os.remove(temp_path)
            
        logger.info(f"Fixed sample sections and saved to {document_path} with {tables_added} tables preserved")
        
    except Exception as e:
        logger.error(f"Error fixing sample sections: {e}")
        # Don't raise, continue as best we can

if __name__ == "__main__":
    # Example usage
    input_doc = Path("attached_assets/EK1586_Mouse_KLK1Kallikrein_1_ELISA_Kit_PicoKine_Datasheet.docx")
    template = Path("templates_docx/enhanced_template.docx")
    output = Path("output_updated_template.docx")
    
    update_template_populator(input_doc, template, output)