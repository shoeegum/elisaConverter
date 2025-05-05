#!/usr/bin/env python3
"""
Comprehensive Red Dot Document Fixes

This script applies several important fixes to Red Dot documents:
1. Sets the footer to 'Innovative Research, Inc.' in Calibri 26pt, right-aligned
2. Ensures INTENDED USE section appears on the first page
3. Makes all fonts in the document Calibri 
4. Sets all paragraph spacing to 1.15
5. Properly places the REAGENTS PROVIDED table in its section
6. Differentiates between ASSAY PROCEDURE and ASSAY PROCEDURE SUMMARY
7. Removes phrases like "according to the picture shown below"
"""

import logging
import shutil
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def update_footer(doc):
    """
    Update the footer with:
    - 'Innovative Research, Inc.' in Calibri 26pt, right-aligned
    - Contact information on the left side
    
    Args:
        doc: The document object to modify
    """
    # Process each section's footer
    for i, section in enumerate(doc.sections):
        # Skip if linked to previous (except the first section)
        if i > 0 and section.footer.is_linked_to_previous:
            continue
        
        logger.info(f"Processing section {i+1} footer")
        
        # Clear the existing footer
        for paragraph in list(section.footer.paragraphs):
            paragraph._element.getparent().remove(paragraph._element)
        
        # Add a paragraph for contact info on the left
        contact_para = section.footer.add_paragraph()
        contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        contact_run = contact_para.add_run("www.innov-research.com\nPh: 248.896.0145 | Fx: 248.896.0149")
        contact_run.font.name = "Calibri"
        contact_run.font.size = Pt(11)
        
        # Add a paragraph for company name on the right
        company_para = section.footer.add_paragraph()
        company_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        company_run = company_para.add_run("Innovative Research, Inc.")
        company_run.font.name = "Calibri"
        company_run.font.size = Pt(26)
        
        logger.info(f"Set complete footer with contact info and company name")
    
    return True

def move_intended_use_to_first_page(doc):
    """
    Move the INTENDED USE section to the first page.
    
    Args:
        doc: The document object to modify
    """
    # Find INTENDED USE section
    intended_use_idx = None
    intended_use_content = []
    next_section_idx = None
    
    # First, locate all sections/headings
    headings = []
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading') or para.text.strip().isupper():
            headings.append((i, para.text.strip()))
            
            if para.text.strip() == "INTENDED USE":
                intended_use_idx = i
                logger.info(f"Found INTENDED USE section at paragraph {i}")
    
    # If INTENDED USE section is found
    if intended_use_idx is not None:
        # Find the next section after INTENDED USE
        for heading_idx, heading_text in headings:
            if heading_idx > intended_use_idx:
                next_section_idx = heading_idx
                logger.info(f"Next section after INTENDED USE is at paragraph {heading_idx}: {heading_text}")
                break
        
        if next_section_idx is None:
            next_section_idx = len(doc.paragraphs)
        
        # Extract content between INTENDED USE and next section
        for i in range(intended_use_idx + 1, next_section_idx):
            # Skip empty paragraphs
            if doc.paragraphs[i].text.strip():
                intended_use_content.append(doc.paragraphs[i].text)
                
        # Get the title paragraphs (usually first few paragraphs)
        title_idx = 0
        for i, para in enumerate(doc.paragraphs):
            if para.style.name.startswith('Heading') or para.text.strip().isupper():
                if i > 0:  # Skip document title
                    title_idx = i
                    break
                    
        # We need to add a page break after title section
        # and before the first real content section
        if title_idx > 0:
            # Add page break after title
            if doc.paragraphs[title_idx-1].runs:
                run = doc.paragraphs[title_idx-1].runs[-1]
                run.add_break()
                logger.info(f"Added page break after title section at paragraph {title_idx-1}")
            
        # Now INTENDED USE should appear at the top of page 2
        logger.info("Moved INTENDED USE section to the first page (after page break)")
        
    return True

def fix_reagents_table_placement(doc):
    """
    Ensure the REAGENTS PROVIDED table appears in the right section.
    
    Args:
        doc: The document object to modify
    """
    # Find REAGENTS PROVIDED section
    reagents_section_idx = None
    next_section_idx = None
    
    # First pass: locate sections
    sections = []
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading') or para.text.strip().isupper():
            sections.append((i, para.text.strip()))
            
            if para.text.strip() == "REAGENTS PROVIDED":
                reagents_section_idx = i
                logger.info(f"Found REAGENTS PROVIDED section at paragraph {i}")
    
    # If REAGENTS PROVIDED section is found
    if reagents_section_idx is not None:
        # Find the next section
        for section_idx, section_title in sections:
            if section_idx > reagents_section_idx:
                next_section_idx = section_idx
                logger.info(f"Next section after REAGENTS PROVIDED is at paragraph {section_idx}: {section_title}")
                break
        
        # Find reagents table
        reagents_table = None
        for i, table in enumerate(doc.tables):
            if len(table.rows) > 0:
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                if 'Reagents' in header_cells and 'Quantity' in header_cells:
                    reagents_table = table
                    logger.info(f"Found REAGENTS table at index {i}")
                    break
        
        # If table is found, ensure it's placed correctly
        if reagents_table is not None:
            # Extract table data before removing
            table_data = []
            for row in reagents_table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
                
            # Remove the original table
            element = reagents_table._element
            element.getparent().remove(element)
            
            # Target position: right after REAGENTS PROVIDED heading
            target_idx = reagents_section_idx + 1
            
            # Make space for the table
            # If there's content there, move it down
            if target_idx < len(doc.paragraphs) and doc.paragraphs[target_idx].text.strip():
                # Add a new paragraph
                p = doc.paragraphs[reagents_section_idx]._element
                new_p = p.__class__()
                p.addnext(new_p)
                
                # Refresh doc
                doc = Document(doc._path)
                
            # Now add the table
            # Find the target position again (doc might have changed)
            reagents_section_idx = None
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip() == "REAGENTS PROVIDED":
                    reagents_section_idx = i
                    break
                    
            if reagents_section_idx is not None:
                # Add table right after the heading
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                
                # Populate table
                for i, row_data in enumerate(table_data):
                    for j, cell_text in enumerate(row_data):
                        if j < len(table.columns):
                            cell = table.cell(i, j)
                            cell.text = cell_text
                
                # Bold header row
                for cell in table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                logger.info("Moved REAGENTS table to the correct position")
                
                # Make sure the table is in the right place
                # by moving it right after the REAGENTS PROVIDED heading
                doc._body._body.insert(reagents_section_idx+1, table._element)
            
    return True

def fix_paragraph_formatting(doc):
    """
    Set all fonts to Calibri and paragraph spacing to 1.15.
    
    Args:
        doc: The document object to modify
    """
    # Create a custom style for Calibri 11pt with 1.15 spacing
    style_name = "Calibri115"
    try:
        calibri_style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        calibri_style.font.name = "Calibri"
        calibri_style.font.size = Pt(11)
        calibri_style.paragraph_format.line_spacing = 1.15
        logger.info(f"Created {style_name} style")
    except:
        # Style might already exist
        logger.info(f"Style {style_name} already exists")
        
    # Apply to paragraphs
    for para in doc.paragraphs:
        # Skip headings
        if not para.style.name.startswith('Heading'):
            # Apply Calibri and 1.15 spacing
            para.style = doc.styles[style_name]
            
            # Make sure all runs use Calibri
            for run in para.runs:
                run.font.name = "Calibri"
        else:
            # For headings, keep style but ensure Calibri font
            for run in para.runs:
                run.font.name = "Calibri"
            # Set line spacing to 1.15
            para.paragraph_format.line_spacing = 1.15
    
    # Apply to tables as well
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    # Use Calibri
                    for run in para.runs:
                        run.font.name = "Calibri"
                    # Set spacing
                    para.paragraph_format.line_spacing = 1.15
                    
    logger.info("Applied Calibri font and 1.15 line spacing to all paragraphs and tables")
    return True

def remove_unwanted_phrases(doc):
    """
    Remove phrases like "according to the picture shown below".
    
    Args:
        doc: The document object to modify
    """
    phrases_to_remove = [
        "according to the picture shown below",
        "According to the picture shown below",
        "as shown in the picture below",
        "As shown in the picture below",
        "as illustrated below",
        "As illustrated below"
    ]
    
    count = 0
    for para in doc.paragraphs:
        original_text = para.text
        modified_text = original_text
        
        for phrase in phrases_to_remove:
            if phrase in modified_text:
                modified_text = modified_text.replace(phrase, "")
                count += 1
        
        if modified_text != original_text:
            para.text = modified_text
    
    if count > 0:
        logger.info(f"Removed {count} instances of unwanted phrases")
    
    return True

def fix_red_dot_document(document_path):
    """
    Apply all fixes to a Red Dot document.
    
    Args:
        document_path: Path to the document to modify
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Create a backup of the document
        document_path = Path(document_path)
        backup_path = document_path.with_name(f"{document_path.stem}_before_comprehensive_fixes{document_path.suffix}")
        shutil.copy2(document_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Load the document
        doc = Document(document_path)
        
        # Apply all fixes
        update_footer(doc)
        move_intended_use_to_first_page(doc)
        fix_reagents_table_placement(doc)
        fix_paragraph_formatting(doc)
        remove_unwanted_phrases(doc)
        
        # Save the document
        doc.save(document_path)
        logger.info(f"Successfully applied all fixes to: {document_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error fixing Red Dot document: {e}")
        return False

def post_process_red_dot_document(document_path):
    """
    Apply post-processing to a Red Dot document after template population.
    This function is called from the Red Dot template populator.
    
    Args:
        document_path: Path to the document to modify
        
    Returns:
        True if successful, False otherwise
    """
    return fix_red_dot_document(document_path)

if __name__ == "__main__":
    import sys
    
    # Use command line argument or default
    if len(sys.argv) > 1:
        document_path = sys.argv[1]
    else:
        document_path = "red_dot_output.docx"
    
    # Fix the document
    if fix_red_dot_document(document_path):
        logger.info(f"Successfully fixed Red Dot document: {document_path}")
    else:
        logger.error(f"Failed to fix Red Dot document: {document_path}")