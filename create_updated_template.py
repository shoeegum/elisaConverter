#!/usr/bin/env python3
"""
Create an updated enhanced template with proper table formats.

Focuses on:
1. Creating a new template with the correct table formats
2. Adding proper variability tables and reproducibility section
"""

import logging
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import shutil

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def create_updated_template():
    """
    Create an updated template with proper table formats.
    """
    # Make a backup of the current template
    template_path = Path("templates_docx/enhanced_template.docx")
    backup_path = Path("templates_docx/enhanced_template_backup.docx")
    output_path = Path("templates_docx/enhanced_template_updated.docx")
    
    # Backup the original template
    shutil.copy(template_path, backup_path)
    logger.info(f"Created backup at {backup_path}")
    
    # Create a new document
    doc = Document()
    logger.info("Creating new template document")
    
    # Set document properties (margins, etc.)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Create the main title
    title = doc.add_paragraph("{{ kit_name|default('ELISA Kit') }}")
    title.style = 'Title'
    for run in title.runs:
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.name = 'Calibri'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add catalog and lot numbers
    cat_lot = doc.add_paragraph()
    cat_lot.add_run("Catalog Number: ").bold = True
    cat_lot.add_run("{{ catalog_number|default('') }}")
    cat_lot.add_run("\nLot Number: ").bold = True
    cat_lot.add_run("{{ lot_number|default('') }}")
    cat_lot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add intended use section
    intended_use_title = doc.add_paragraph("INTENDED USE")
    intended_use_title.style = 'Heading 2'
    intended_use_title.runs[0].font.color.rgb = RGBColor(0, 70, 180)
    
    intended_use_text = doc.add_paragraph("{{ intended_use|default('') }}")
    
    # Add page break
    doc.add_page_break()
    
    # Add Technical Details section
    tech_details_title = doc.add_paragraph("TECHNICAL DETAILS")
    tech_details_title.style = 'Heading 2'
    tech_details_title.runs[0].font.color.rgb = RGBColor(0, 70, 180)
    
    # Add technical details table
    tech_table = doc.add_table(rows=4, cols=2)
    tech_table.style = 'Table Grid'
    
    # Fill the technical details table
    tech_headers = ["Capture/Detection Antibodies", "Specificity", "Standard Protein", "Cross-reactivity"]
    for i, header in enumerate(tech_headers):
        cell = tech_table.cell(i, 0)
        cell.text = header
        # Bold the header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add template placeholders
    for i in range(4):
        cell = tech_table.cell(i, 1)
        cell.text = f"{{{{ technical_details_table[{i}].value|default('') }}}}"
    
    # Add overview section
    overview_title = doc.add_paragraph("OVERVIEW")
    overview_title.style = 'Heading 2'
    overview_title.runs[0].font.color.rgb = RGBColor(0, 70, 180)
    
    # Add overview table
    overview_table = doc.add_table(rows=8, cols=2)
    overview_table.style = 'Table Grid'
    
    # Fill the overview table
    overview_headers = ["Product Name", "Reactive Species", "Range", "Sensitivity", 
                       "Sample Type", "Cross Reactivity", "Storage", "Expiration"]
    for i, header in enumerate(overview_headers):
        cell = overview_table.cell(i, 0)
        cell.text = header
        # Bold the header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add template placeholders
    for i in range(8):
        cell = overview_table.cell(i, 1)
        cell.text = f"{{{{ overview_specifications_table[{i}].value|default('') }}}}"
    
    # Add background section
    background_title = doc.add_paragraph("BACKGROUND")
    background_title.style = 'Heading 2'
    background_title.runs[0].font.color.rgb = RGBColor(0, 70, 180)
    
    background_text = doc.add_paragraph("{{ background_text|default('') }}")
    
    # Add kit components section
    components_title = doc.add_paragraph("KIT COMPONENTS")
    components_title.style = 'Heading 2'
    components_title.runs[0].font.color.rgb = RGBColor(0, 70, 180)
    
    # Add kit components table
    components_table = doc.add_table(rows=8, cols=4)
    components_table.style = 'Table Grid'
    
    # Fill the header row
    header_row = components_table.rows[0]
    header_cells = ["Description", "Quantity", "Volume", "Storage of opened/reconstituted material"]
    for i, text in enumerate(header_cells):
        cell = header_row.cells[i]
        cell.text = text
        # Make the text bold
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add template placeholders
    for i in range(1, 8):
        for j, field in enumerate(["name", "quantity", "volume", "storage"]):
            cell = components_table.cell(i, j)
            cell.text = f"{{{{ reagent_{i}_{field}|default('') }}}}"
    
    # Add required materials section
    materials_title = doc.add_paragraph("MATERIALS REQUIRED BUT NOT PROVIDED")
    materials_title.style = 'Heading 2'
    materials_title.runs[0].font.color.rgb = RGBColor(0, 70, 180)
    
    # Add placeholder for materials
    materials_para = doc.add_paragraph("{{ required_materials_with_bullets|default('') }}")
    
    # Add reagent preparation section
    reagent_title = doc.add_paragraph("REAGENT PREPARATION")
    reagent_title.style = 'Heading 2'
    reagent_title.runs[0].font.color.rgb = RGBColor(0, 70, 180)
    
    reagent_text = doc.add_paragraph("{{ reagent_preparation|default('') }}")
    
    # Add standard dilution section
    dilution_title = doc.add_paragraph("DILUTION OF STANDARD")
    dilution_title.style = 'Heading 2'
    dilution_title.runs[0].font.color.rgb = RGBColor(0, 70, 180)
    
    dilution_text = doc.add_paragraph("{{ dilution_of_standard|default('') }}")
    
    # Add preparations before assay section
    prep_title = doc.add_paragraph("PREPARATIONS BEFORE ASSAY")
    prep_title.style = 'Heading 2'
    prep_title.runs[0].font.color.rgb = RGBColor(0, 70, 180)
    
    # Add numbered steps for preparations
    for i in range(1, 6):
        if i == 1:
            prep_text = doc.add_paragraph("1. Prepare all reagents, samples, and standards according to the instructions.")
        elif i == 2:
            prep_text = doc.add_paragraph("2. Confirm that you have the appropriate non-supplied equipment available.")
        elif i == 3:
            prep_text = doc.add_paragraph("3. Spin down all components to the bottom of the tube before opening.")
        elif i == 4:
            prep_text = doc.add_paragraph("4. Don't let the 96-well plate dry out as this will inactivate active components.")
        elif i == 5:
            prep_text = doc.add_paragraph("5. Don't reuse tips and tubes to avoid cross-contamination. Avoid using reagents from different batches.")
    
    # Add assay protocol section
    protocol_title = doc.add_paragraph("ASSAY PROTOCOL")
    protocol_title.style = 'Heading 2'
    protocol_title.runs[0].font.color.rgb = RGBColor(0, 70, 180)
    
    # Add numbered steps for protocol
    protocol_text = doc.add_paragraph("{{ assay_protocol_numbered|default('') }}")
    
    # Add standard curve section
    curve_title = doc.add_paragraph("TYPICAL DATA / STANDARD CURVE")
    curve_title.style = 'Heading 2'
    curve_title.runs[0].font.color.rgb = RGBColor(0, 70, 180)
    
    # Add curve description
    curve_desc = doc.add_paragraph("This standard curve is for demonstration only. A standard curve must be run with each assay.")
    
    # Add Standard Curve table (2 rows by 9 columns)
    std_table = doc.add_table(rows=2, cols=9)
    std_table.style = 'Table Grid'
    std_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Fill the header row
    header_row = std_table.rows[0]
    header_cells = ["Concentration (pg/ml)", "0", "62.5", "125", "250", "500", "1000", "2000", "4000"]
    for i, text in enumerate(header_cells):
        cell = header_row.cells[i]
        cell.text = text
        # Make the text bold
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Fill the data row
    data_row = std_table.rows[1]
    data_cells = ["O.D.", "{{ std_od_1|default('') }}", "{{ std_od_2|default('') }}", 
                 "{{ std_od_3|default('') }}", "{{ std_od_4|default('') }}", 
                 "{{ std_od_5|default('') }}", "{{ std_od_6|default('') }}",
                 "{{ std_od_7|default('') }}", "{{ std_od_8|default('') }}"]
    for i, text in enumerate(data_cells):
        cell = data_row.cells[i]
        cell.text = text
        # Set the text alignment
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add variability section
    var_title = doc.add_paragraph("INTRA/INTER-ASSAY VARIABILITY")
    var_title.style = 'Heading 2'
    var_title.runs[0].font.color.rgb = RGBColor(0, 70, 180)
    
    # Add intra-assay description
    intra_desc = doc.add_paragraph("Three samples of known concentration were tested on one plate to assess intra-assay precision.")
    
    # Add Intra-Assay Precision table
    intra_table = doc.add_table(rows=4, cols=5)
    intra_table.style = 'Table Grid'
    intra_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Fill the header row
    header_row = intra_table.rows[0]
    header_cells = ["Sample", "n", "Mean (pg/ml)", "Standard Deviation", "CV (%)"]
    for i, text in enumerate(header_cells):
        cell = header_row.cells[i]
        cell.text = text
        # Make the text bold and centered
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Fill data rows
    data_rows = [
        ["1", "16", "150", "9.15", "6.1%"],
        ["2", "16", "602", "43.94", "7.3%"],
        ["3", "16", "1476", "116.6", "7.9%"]
    ]
    
    for i, row_data in enumerate(data_rows):
        row = intra_table.rows[i+1]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            # Center the text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add spacing
    doc.add_paragraph()
    
    # Add inter-assay description
    inter_desc = doc.add_paragraph("Three samples of known concentration were tested in separate assays to assess inter-assay precision.")
    
    # Add Inter-Assay Precision table
    inter_table = doc.add_table(rows=4, cols=5)
    inter_table.style = 'Table Grid'
    inter_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Fill the header row
    header_row = inter_table.rows[0]
    header_cells = ["Sample", "n", "Mean (pg/ml)", "Standard Deviation", "CV (%)"]
    for i, text in enumerate(header_cells):
        cell = header_row.cells[i]
        cell.text = text
        # Make the text bold and centered
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Fill data rows
    data_rows = [
        ["1", "24", "145", "10.15", "7.0%"],
        ["2", "24", "618", "49.44", "8.0%"],
        ["3", "24", "1426", "128.34", "9.0%"]
    ]
    
    for i, row_data in enumerate(data_rows):
        row = inter_table.rows[i+1]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            # Center the text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add reproducibility section
    reprod_title = doc.add_paragraph("REPRODUCIBILITY")
    reprod_title.style = 'Heading 2'
    reprod_title.runs[0].font.color.rgb = RGBColor(0, 70, 180)
    
    # Add description paragraph
    reprod_desc = doc.add_paragraph("Samples were tested in four different assay lots to assess reproducibility.")
    
    # Add Reproducibility table
    reprod_table = doc.add_table(rows=4, cols=7)
    reprod_table.style = 'Table Grid'
    reprod_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Fill the header row
    header_row = reprod_table.rows[0]
    header_cells = ["", "Lot 1", "Lot 2", "Lot 3", "Lot 4", "Mean", "CV (%)"]
    for i, text in enumerate(header_cells):
        cell = header_row.cells[i]
        cell.text = text
        # Make the text bold and centered
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Fill data rows
    data_rows = [
        ["Sample 1", "150", "154", "170", "150", "156", "5.2%"],
        ["Sample 2", "602", "649", "645", "637", "633", "2.9%"],
        ["Sample 3", "1476", "1672", "1722", "1744", "1654", "7.2%"]
    ]
    
    for i, row_data in enumerate(data_rows):
        row = reprod_table.rows[i+1]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            # Set first column bold
            if j == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            # Center the text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save the document
    doc.save(output_path)
    logger.info(f"Created updated template at {output_path}")
    
    # Copy the updated template to the main template location
    shutil.copy(output_path, template_path)
    logger.info(f"Updated main template at {template_path}")
    
    return output_path

if __name__ == "__main__":
    create_updated_template()