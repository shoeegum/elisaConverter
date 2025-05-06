#!/usr/bin/env python3
"""
Check the structure of the Innovative Research output document.
"""

import logging
from docx import Document
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def check_document_structure(document_path="red_dot_output.docx"):
    """
    Check the structure of the document and print a detailed layout of sections,
    paragraphs, and tables.
    
    Args:
        document_path: Path to the document to check
    """
    try:
        # Convert to Path if string
        if isinstance(document_path, str):
            document_path = Path(document_path)
            
        # Load the document
        doc = Document(document_path)
        
        logger.info(f"=== Document Structure of {document_path} ===")
        
        # Check document title
        if len(doc.paragraphs) > 0:
            title = doc.paragraphs[0].text
            logger.info(f"Document Title: {title}")
        
        # Check for correct company name
        incorrect_name_count = 0
        for para in doc.paragraphs:
            if "Reddot Biotech" in para.text:
                incorrect_name_count += 1
                logger.warning(f"Found incorrect company name in paragraph: '{para.text[:50]}...'")
                
        if incorrect_name_count > 0:
            logger.warning(f"Found {incorrect_name_count} instances of incorrect company name 'Reddot Biotech'")
        else:
            logger.info("Company name appears to be correct (no 'Reddot Biotech' instances found)")
            
        # Find all section headings
        sections = []
        for i, para in enumerate(doc.paragraphs):
            # If paragraph has a style that starts with 'Heading' or contains uppercase text that could be a heading
            if (para.style.name.startswith('Heading') or 
                (para.text.isupper() and len(para.text.strip()) > 0 and len(para.text.strip()) < 50)):
                sections.append((i, para.text))
                logger.info(f"Section at P{i}: {para.text}")
                
                # Check next paragraph for placeholders
                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1]
                    if "{{" in next_para.text and "}}" in next_para.text:
                        logger.warning(f"  - Found unprocessed placeholder: {next_para.text}")
                    else:
                        # Show a snippet of the next paragraph
                        content = next_para.text[:50] + "..." if len(next_para.text) > 50 else next_para.text
                        if content.strip():
                            logger.info(f"  - Content starts with: {content}")
                            
        # Check tables
        logger.info("\n=== Tables ===")
        for i, table in enumerate(doc.tables):
            rows = len(table.rows)
            cols = len(table.columns) if rows > 0 else 0
            
            # Get table title (from preceding paragraph if possible)
            table_title = "Unknown"
            table_xml = table._element
            prev_paragraph = table_xml.getprevious()
            
            # Try to extract the closest preceding paragraph
            if prev_paragraph is not None:
                try:
                    import re
                    from docx.oxml.text.paragraph import CT_P
                    if isinstance(prev_paragraph, CT_P):
                        text = "".join([t.text for t in prev_paragraph.xpath(".//w:t")])
                        table_title = text
                except Exception:
                    # If we can't extract it, just use a generic title
                    pass
            
            logger.info(f"Table {i}: {rows}x{cols} (Title: {table_title})")
            
            # Check if this appears to be the reagents table
            reagents_table = False
            if rows > 0:
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                if 'Reagents' in header_cells or 'Component' in header_cells:
                    reagents_table = True
                    
            if reagents_table:
                # Check where this table is positioned
                table_index = i
                correct_position = False
                
                # Try to determine if this table is in the right place (after REAGENTS PROVIDED section)
                for sec_idx, (para_idx, section_title) in enumerate(sections):
                    if "REAGENTS PROVIDED" in section_title:
                        # Find the table closest to this section
                        # This is a bit of a heuristic since we can't directly know which table
                        # is associated with which paragraph in python-docx
                        if i == 0 or (i > 0 and para_idx > sections[sec_idx-1][0]):
                            correct_position = True
                            break
                            
                logger.info(f"Reagents Table Found at index {table_index}")
                if correct_position:
                    logger.info("  - Table appears to be in the correct position")
                else:
                    logger.warning("  - Table may not be in the correct position")
                    
                # Show some table contents
                if rows > 0:
                    for j, cell in enumerate(table.rows[0].cells):
                        logger.info(f"  - Column {j}: {cell.text}")
                        
                    # Show a sample of rows
                    max_sample = min(5, rows)
                    for j in range(1, max_sample):
                        try:
                            row_text = " | ".join([cell.text for cell in table.rows[j].cells])
                            logger.info(f"  - Row {j}: {row_text[:50]}..." if len(row_text) > 50 else f"  - Row {j}: {row_text}")
                        except:
                            pass
                        
        # Check footer
        logger.info("\n=== Footer ===")
        for i, section in enumerate(doc.sections):
            footer = section.footer
            for para in footer.paragraphs:
                logger.info(f"Footer text in section {i+1}: '{para.text}'")
                if "innov-research.com" in para.text.lower():
                    logger.info("Footer appears to be correct (contains 'innov-research.com')")
                else:
                    logger.warning("Footer does not contain expected text 'innov-research.com'")
                    
        logger.info("\n=== Check Complete ===")
        return True
        
    except Exception as e:
        logger.error(f"Error checking document structure: {e}")
        return False
        
if __name__ == "__main__":
    import sys
    
    # Use command line argument or default
    if len(sys.argv) > 1:
        document_path = sys.argv[1]
    else:
        document_path = "red_dot_output.docx"
    
    check_document_structure(document_path)