#!/usr/bin/env python3
"""
Improved checker for validating that numbered lists are properly preserved
in the Preparations Before Assay section of ELISA kit datasheets.
"""

import sys
from docx import Document
from pathlib import Path

def check_document_sections(document_path):
    """Check specific sections in a document for formatting and content."""
    doc = Document(document_path)
    filename = Path(document_path).name
    
    print(f"Analyzing document: {filename}")
    print("-" * 80)
    
    # First, get all section headings to understand document structure
    section_headings = []
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            section_headings.append(para.text.strip())
    
    print("Document sections found:")
    for heading in section_headings:
        print(f"  - {heading}")
    print()
    
    # Now specifically check the Preparations Before Assay section
    print("CHECKING 'PREPARATIONS BEFORE ASSAY' SECTION:")
    print("-" * 40)
    
    in_preparations_section = False
    found_preparations_section = False
    section_paragraphs = []
    has_numbered_lists = False
    
    # First, get the exact title and index of the Preparations Before Assay section
    prep_section_idx = -1
    for i, para in enumerate(doc.paragraphs):
        if "PREPARATIONS BEFORE ASSAY" in para.text.upper() and (
            para.style.name.startswith("Heading") or 
            len(para.text) < 100  # Short paragraph likely to be a heading
        ):
            prep_section_idx = i
            found_preparations_section = True
            print(f"Found section at paragraph {i}: {para.text}")
            break
            
    if not found_preparations_section:
        print("Could not find Preparations Before Assay section heading!")
        return False
        
    # Now find paragraphs between this section and the next section heading
    next_section_idx = len(doc.paragraphs)  # Default to end of document
    for i in range(prep_section_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if (para.style.name.startswith("Heading") or 
            (len(para.text) < 100 and any(keyword in para.text.upper() for keyword in 
             ["KIT COMPONENTS", "MATERIALS PROVIDED", "REAGENT", "PROTOCOL"]))):
            next_section_idx = i
            print(f"Section ends at paragraph {i}: {para.text}")
            break
    
    # Collect all paragraphs in the section
    for i in range(prep_section_idx + 1, next_section_idx):
        para = doc.paragraphs[i]
        
        # Include this paragraph if it has content
        if para.text.strip():  # Only include non-empty paragraphs
            section_paragraphs.append({
                "index": i,
                "style": para.style.name,
                "text": para.text.strip(),
                "is_numbered": bool(para.style.name == "List Number" or 
                                    (para.text.strip() and para.text.strip()[0].isdigit() and 
                                    ". " in para.text.strip()[:5]))
            })
            
            # Check if this is a numbered list paragraph
            if section_paragraphs[-1]["is_numbered"]:
                has_numbered_lists = True
    
    if not found_preparations_section:
        print("PREPARATIONS BEFORE ASSAY section not found in document!")
        return False
        
    # Print section details
    print(f"\nSection 'PREPARATIONS BEFORE ASSAY' contains {len(section_paragraphs)} paragraphs:")
    for para in section_paragraphs:
        print(f"  [{para['style']}] {para['text'][:60]}..." + 
              (" (NUMBERED)" if para["is_numbered"] else ""))
    
    # Summary
    print("\n" + "=" * 40)
    print(f"SUMMARY FOR {filename}:")
    print(f"  - 'PREPARATIONS BEFORE ASSAY' section found: {found_preparations_section}")
    print(f"  - Contains numbered lists: {has_numbered_lists}")
    print(f"  - Total paragraphs in section: {len(section_paragraphs)}")
    if has_numbered_lists:
        numbered_paras = [p for p in section_paragraphs if p["is_numbered"]]
        print(f"  - Number of numbered paragraphs: {len(numbered_paras)}")
        print(f"  - List items: {', '.join([p['text'][:10] + '...' for p in numbered_paras])}")
    
    return has_numbered_lists

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python improved_checker.py <docx_file> [<docx_file2> ...]")
        sys.exit(1)
    
    # Check all provided files
    for doc_file in sys.argv[1:]:
        check_document_sections(doc_file)
        if len(sys.argv) > 2:  # If multiple files, add a separator
            print("\n" + "=" * 80 + "\n")