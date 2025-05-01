#!/usr/bin/env python3
"""
Check if the footer text has been properly changed
"""

import logging
from pathlib import Path
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def check_footer_text(document_path):
    """
    Check if the footer text has been properly changed.
    
    Args:
        document_path: Path to the document to check
    """
    try:
        # Load the document
        doc = Document(document_path)
        
        # Check all sections for footers
        footer_texts = []
        
        # Get sections and check if we have any sections at all
        sections = list(doc.sections)
        logger.info(f"Found {len(sections)} sections in the document")
        
        if not sections:
            logger.warning("No sections found in the document!")
            
        # Process each section
        for i, section in enumerate(sections):
            logger.info(f"Checking section {i+1}")
            
            # Check if section has a footer
            if not hasattr(section, 'footer'):
                logger.warning(f"Section {i+1} has no footer attribute")
                continue
                
            # Check if footer has paragraphs
            if not hasattr(section.footer, 'paragraphs'):
                logger.warning(f"Section {i+1} footer has no paragraphs attribute")
                continue
                
            # Log section information
            logger.info(f"Section {i+1}: Footer linked to previous: {section.footer.is_linked_to_previous}")
            logger.info(f"Section {i+1}: Footer paragraph count: {len(section.footer.paragraphs)}")
            
            # Skip if linked to previous (except the first section)
            if i > 0 and section.footer.is_linked_to_previous:
                logger.info(f"Section {i+1} footer is linked to previous, skipping")
                continue
                
            # Get the text from each paragraph in the footer
            for j, paragraph in enumerate(section.footer.paragraphs):
                footer_texts.append(paragraph.text)
                logger.info(f"Section {i+1}, Footer paragraph {j+1}: '{paragraph.text}'")
                
                # Also check individual runs
                for k, run in enumerate(paragraph.runs):
                    if run.text.strip():
                        logger.info(f"Section {i+1}, Footer paragraph {j+1}, Run {k+1}: '{run.text.strip()}'")
        
        # Check header text as well (sometimes footers are stored in headers)
        header_texts = []
        for i, section in enumerate(sections):
            if hasattr(section, 'header') and hasattr(section.header, 'paragraphs'):
                for j, paragraph in enumerate(section.header.paragraphs):
                    header_texts.append(paragraph.text)
                    logger.info(f"Section {i+1}, Header paragraph {j+1}: '{paragraph.text}'")
        
        # Print all found footer texts
        logger.info(f"Found {len(footer_texts)} footer paragraphs")
        for i, text in enumerate(footer_texts):
            logger.info(f"Footer {i+1}: '{text}'")
            
        # Check if any footer contains "All rights reserved"
        if any("All rights reserved" in text for text in footer_texts + header_texts):
            logger.warning("Document still contains 'All rights reserved'")
        else:
            logger.info("Document does not contain 'All rights reserved'")
        
        # Check if any footer contains "Made by Sophie Gall"
        if any("Made by Sophie Gall" in text for text in footer_texts + header_texts):
            logger.info("Document successfully updated to include 'Made by Sophie Gall'")
        else:
            logger.warning("Document does not contain 'Made by Sophie Gall'")
        
        # Check if any footer contains the copyright symbol
        if any("©" in text for text in footer_texts + header_texts):
            logger.warning("Document still contains the copyright symbol (©)")
        else:
            logger.info("Document does not contain the copyright symbol (©)")
            
        return footer_texts
        
    except Exception as e:
        logger.error(f"Error checking footer: {e}")
        return []

if __name__ == "__main__":
    # Check the current output document
    check_footer_text("output_populated_template.docx")