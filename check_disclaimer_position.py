"""
Check the position of the DISCLAIMER section relative to DATA ANALYSIS.
"""

import docx
from docx.enum.text import WD_BREAK

def check_disclaimer_position(document_path="output_populated_template.docx"):
    """
    Check if there's a page break between DATA ANALYSIS and DISCLAIMER.
    
    Args:
        document_path: Path to the document to check
    """
    # Load the document
    doc = docx.Document(document_path)
    
    # Find the DATA ANALYSIS section and DISCLAIMER section
    data_analysis_idx = -1
    disclaimer_idx = -1
    
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().upper() == "DATA ANALYSIS":
            data_analysis_idx = i
        elif paragraph.text.strip().upper() == "DISCLAIMER":
            disclaimer_idx = i
    
    if data_analysis_idx == -1:
        print("DATA ANALYSIS section not found")
        return
        
    if disclaimer_idx == -1:
        print("DISCLAIMER section not found")
        return
    
    print(f"DATA ANALYSIS section found at paragraph {data_analysis_idx}")
    print(f"DISCLAIMER section found at paragraph {disclaimer_idx}")
    
    # Check if there are page breaks between DATA ANALYSIS and DISCLAIMER
    page_breaks_found = False
    for i in range(data_analysis_idx, disclaimer_idx):
        paragraph = doc.paragraphs[i]
        for run in paragraph.runs:
            for break_type in run._element.xpath(".//w:br"):
                if break_type.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "page":
                    print(f"Page break found in paragraph {i}")
                    page_breaks_found = True
    
    if not page_breaks_found:
        print("No page breaks found between DATA ANALYSIS and DISCLAIMER")
    
    # Check the number of paragraphs between DATA ANALYSIS and DISCLAIMER
    paragraph_count = disclaimer_idx - data_analysis_idx - 1
    print(f"Number of paragraphs between DATA ANALYSIS and DISCLAIMER: {paragraph_count}")
    
    # Show the content of paragraphs between DATA ANALYSIS and DISCLAIMER
    print("\nContent between DATA ANALYSIS and DISCLAIMER:")
    for i in range(data_analysis_idx + 1, disclaimer_idx):
        print(f"Para {i}: {doc.paragraphs[i].text[:100]}")

if __name__ == "__main__":
    check_disclaimer_position()