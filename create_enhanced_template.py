#!/usr/bin/env python3
"""
Create Enhanced Template for ELISA Kit Datasheets

This script creates an enhanced template with proper styles, formatting, and placeholders
for ELISA kit datasheets based on the provided requirements:
1. All section titles should use Heading 2 Style (blue, all caps)
2. Proper paragraph spacing (1.15) and font (Calibri)
3. Bold Calibri 36pt for main title
4. Company name in footer should be Calibri 24pt bold
5. Contact info in footer should be Open Sans Light 12pt
6. Required materials should use a bullet list
7. First page should only contain title, catalog/lot numbers, and intended use
"""

import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

# Set up logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def create_enhanced_template():
    """
    Create an enhanced template for ELISA kit datasheets with proper styling.
    """
    doc = Document()
    
    # Set document styles
    styles = doc.styles
    
    # Set default font for the entire document
    font = styles['Normal'].font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Create a hidden text style for template variables
    hidden_style = styles.add_style('Hidden Text', WD_STYLE_TYPE.PARAGRAPH)
    hidden_style.base_style = styles['Normal']
    hidden_style.hidden = True
    hidden_style.priority = 99
    hidden_font = hidden_style.font
    hidden_font.color.rgb = RGBColor(200, 200, 200)
    hidden_font.size = Pt(8)
    
    # Set narrow margins for the entire document
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Create heading styles
    # Heading 1 for kit name
    h1_style = styles['Heading 1']
    h1_style.font.name = 'Calibri'
    h1_style.font.size = Pt(36)
    h1_style.font.bold = True
    h1_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Heading 2 for section titles
    h2_style = styles['Heading 2']
    h2_style.font.name = 'Calibri'
    h2_style.font.size = Pt(12)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 70, 180)  # Bright blue
    h2_style.font.all_caps = True  # ALL CAPS formatting
    
    # Set paragraph spacing and line spacing
    paragraph_format = styles['Normal'].paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(8)
    paragraph_format.line_spacing = 1.15  # Match requested 1.15 spacing
    
    # Create bullet list style with proper formatting
    if 'List Bullet' not in styles:
        bullet_style = styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
    else:
        bullet_style = styles['List Bullet']
    
    bullet_style.base_style = styles['Normal']
    bullet_style.font.name = 'Calibri'
    bullet_style.font.size = Pt(11)
    bullet_style.paragraph_format.left_indent = Inches(0.25)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.25)
    
    # Create numbered list style with proper formatting
    if 'List Number' not in styles:
        number_style = styles.add_style('List Number', WD_STYLE_TYPE.PARAGRAPH)
    else:
        number_style = styles['List Number']
    
    number_style.base_style = styles['Normal']
    number_style.font.name = 'Calibri'
    number_style.font.size = Pt(11)
    number_style.paragraph_format.left_indent = Inches(0.25)
    number_style.paragraph_format.first_line_indent = Inches(-0.25)
    
    # Create footer styles
    footer_company_style = styles.add_style('Footer Company Style', WD_STYLE_TYPE.PARAGRAPH)
    footer_company_style.font.name = 'Calibri'
    footer_company_style.font.size = Pt(24)
    footer_company_style.font.bold = True
    
    footer_info_style = styles.add_style('Footer Info Style', WD_STYLE_TYPE.PARAGRAPH)
    footer_info_style.font.name = 'Open Sans Light'
    footer_info_style.font.size = Pt(12)
    
    # TABLE STYLES
    # Create a clean table style
    table_style = styles.add_style('Clean Table Style', WD_STYLE_TYPE.TABLE)
    table_style.base_style = styles['Table Grid']
    
    # PAGE 1 - Simple with only title, catalog/lot number, and intended use
    # Add the kit name title
    title_para = doc.add_paragraph("{{ kit_name }}", style='Heading 1')
    
    # First approach: Two separate paragraphs with different alignments
    
    # Add catalog number (left-aligned)
    catalog_para = doc.add_paragraph(style='Heading 2')
    catalog_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    catalog_run = catalog_para.add_run("CATALOG NO: ")
    catalog_run.font.color.rgb = RGBColor(0, 70, 180)  # Ensure blue color
    catalog_para.add_run("{{ catalog_number }}")
    
    # Add lot number (right-aligned)
    lot_para = doc.add_paragraph(style='Heading 2')
    lot_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    lot_run = lot_para.add_run("LOT NO: ")
    lot_run.font.color.rgb = RGBColor(0, 70, 180)  # Ensure blue color
    lot_para.add_run("{{ lot_number }}")
    
    # Add INTENDED USE section
    intended_use_header = doc.add_paragraph("INTENDED USE", style='Heading 2')
    intended_use_para = doc.add_paragraph("{{ intended_use }}")
    
    # Add page break after first page content
    doc.add_page_break()
    
    # PAGE 2 - Start with Background
    # BACKGROUND
    background_header = doc.add_paragraph("BACKGROUND", style='Heading 2')
    background_para = doc.add_paragraph("{{ background }}")
    
    # PRINCIPLE OF THE ASSAY
    principle_header = doc.add_paragraph("PRINCIPLE OF THE ASSAY", style='Heading 2')
    principle_para = doc.add_paragraph("{{ assay_principle }}")
    
    # OVERVIEW
    overview_header = doc.add_paragraph("OVERVIEW", style='Heading 2')
    overview_para = doc.add_paragraph("{{ overview }}")
    
    # Create a table for the specifications - no header row, just the 8 specification rows
    spec_table = doc.add_table(rows=8, cols=2)  # 8 standard specs only
    spec_table.style = 'Table Grid'
    
    # Set column widths for better readability
    for cell in spec_table.columns[0].cells:
        cell.width = Inches(2.5)
    for cell in spec_table.columns[1].cells:
        cell.width = Inches(3.5)
    
    # Add the standard specifications rows with Jinja2 variables
    standard_specs = [
        "Product Name", 
        "Reactive Species", 
        "Size", 
        "Description", 
        "Sensitivity", 
        "Detection Range", 
        "Storage Instructions", 
        "Uniprot ID"
    ]
    
    for i, prop in enumerate(standard_specs):
        row = spec_table.rows[i].cells
        row[0].text = prop
        row[1].text = "{{ overview_specifications_table[" + str(i) + "].value }}"
    
    # TECHNICAL DETAILS
    tech_details_header = doc.add_paragraph("TECHNICAL DETAILS", style='Heading 2')
    
    # Create a table for the technical details
    tech_table = doc.add_table(rows=4, cols=2)  # 4 standard fields
    tech_table.style = 'Table Grid'
    
    # Set column widths for better readability
    for cell in tech_table.columns[0].cells:
        cell.width = Inches(3.0)
    for cell in tech_table.columns[1].cells:
        cell.width = Inches(3.0)
    
    # Add the technical details rows with Jinja2 variables
    technical_properties = [
        "Capture/Detection Antibodies", 
        "Specificity", 
        "Standard Protein", 
        "Cross-reactivity"
    ]
    
    for i, prop in enumerate(technical_properties):
        row = tech_table.rows[i].cells
        row[0].text = prop
        row[1].text = "{{ technical_details_table[" + str(i) + "].value }}"
    
    # Add paragraph for any additional technical details text
    doc.add_paragraph("{{ technical_details }}")
    
    # PREPARATIONS BEFORE ASSAY
    prep_header = doc.add_paragraph("PREPARATIONS BEFORE ASSAY", style='Heading 2')
    
    # Add a placeholder for the preparation text (non-numbered portion)
    doc.add_paragraph("{{ preparations_text }}")
    
    # Add numbered preparation steps using Jinja2 loop - proper nesting
    doc.add_paragraph("{% if preparations_steps %}", style="Hidden Text")
    doc.add_paragraph("{% for step in preparations_steps %}", style="Hidden Text")
    
    # Placeholder for a numbered step - must be inside the for loop
    num_para = doc.add_paragraph()
    num_para.style = 'List Number'
    num_para.add_run("{{ step.text }}")
    
    # Close the loops in the correct order
    doc.add_paragraph("{% endfor %}", style="Hidden Text")
    doc.add_paragraph("{% endif %}", style="Hidden Text")
    
    # KIT COMPONENTS/MATERIALS PROVIDED
    kit_components_header = doc.add_paragraph("KIT COMPONENTS/MATERIALS PROVIDED", style='Heading 2')
    
    # Create a table with fixed number of rows for components
    # We'll create a static table with 7 rows, which should be sufficient for most kits
    reagents_table = doc.add_table(rows=8, cols=4)  # 1 header row + 7 data rows
    reagents_table.style = 'Table Grid'
    
    # Set column widths
    col_widths = [3.0, 1.0, 1.5, 2.5]  # inches for each column
    for i, width in enumerate(col_widths):
        for cell in reagents_table.columns[i].cells:
            cell.width = Inches(width)
    
    # Add headers to reagents table
    header_row = reagents_table.rows[0].cells
    header_row[0].text = "Description"
    header_row[1].text = "Quantity"
    header_row[2].text = "Volume"
    header_row[3].text = "Storage of opened/reconstituted material"
    
    # Make headers bold and center-aligned
    for cell in header_row:
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            
    # Add placeholders for up to 7 reagents (static approach)
    for i in range(1, 8):
        row = reagents_table.rows[i].cells
        row[0].text = f"{{{{ reagent_{i}_name|default('') }}}}"
        row[1].text = f"{{{{ reagent_{i}_quantity|default('') }}}}"
        row[2].text = f"{{{{ reagent_{i}_volume|default('') }}}}"
        row[3].text = f"{{{{ reagent_{i}_storage|default('') }}}}"
    
    # MATERIALS REQUIRED BUT NOT PROVIDED
    materials_header = doc.add_paragraph("MATERIALS REQUIRED BUT NOT PROVIDED", style='Heading 2')
    
    # Create static bullet points for required materials (typical items, up to 10)
    for i in range(1, 11):
        bullet_para = doc.add_paragraph(f"{{{{ req_material_{i}|default('') }}}}", style='List Bullet')
        # Only show bullet points with actual content
        doc.add_paragraph(f"{{% if not req_material_{i} %}}{{{{ '' }}}}{{% endif %}}", style="Hidden Text")
    
    # REAGENT PREPARATION
    reagent_prep_header = doc.add_paragraph("REAGENT PREPARATION", style='Heading 2')
    reagent_prep_para = doc.add_paragraph("{{ reagent_preparation }}")
    
    # SAMPLE PREPARATION
    sample_prep_header = doc.add_paragraph("SAMPLE PREPARATION", style='Heading 2')
    sample_prep_para = doc.add_paragraph("{{ sample_preparation }}")
    
    # DILUTION OF STANDARD
    dilution_header = doc.add_paragraph("DILUTION OF STANDARD", style='Heading 2')
    dilution_para = doc.add_paragraph("{{ dilution_of_standard }}")
    
    # TYPICAL DATA / STANDARD CURVE
    std_curve_header = doc.add_paragraph("TYPICAL DATA / STANDARD CURVE", style='Heading 2')
    std_curve_para = doc.add_paragraph("This standard curve is provided for demonstration only. A standard curve should be generated for each set of samples assayed.")
    
    # Add standard curve table with fixed number of rows
    # Typically there are 7-8 standard dilutions in an ELISA kit
    curve_table = doc.add_table(rows=9, cols=2)  # 1 header row + 8 data rows
    curve_table.style = 'Table Grid'
    
    # Add headers to curve table
    curve_headers = curve_table.rows[0].cells
    curve_headers[0].text = "Concentration (pg/ml)"
    curve_headers[1].text = "O.D."
    
    # Make headers bold and centered
    for cell in curve_headers:
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    
    # Add placeholders for standard curve values
    for i in range(1, 9):
        row = curve_table.rows[i].cells
        row[0].text = f"{{{{ std_conc_{i}|default('') }}}}"
        row[1].text = f"{{{{ std_od_{i}|default('') }}}}"
    
    # INTRA/INTER-ASSAY VARIABILITY
    variability_header = doc.add_paragraph("INTRA/INTER-ASSAY VARIABILITY", style='Heading 2')
    intra_para = doc.add_paragraph("- Intra-Assay Precision: Three samples of known concentration were tested on one plate to assess intra-assay precision.")
    inter_para = doc.add_paragraph("- Inter-Assay Precision: Three samples of known concentration were tested in separate assays to assess inter-assay precision.")
    precision_values = doc.add_paragraph("{{ variability_data }}")
    
    # ASSAY PROTOCOL
    protocol_header = doc.add_paragraph("ASSAY PROTOCOL", style='Heading 2')
    
    # Create static numbered steps for assay protocol (typical protocol has ~15-20 steps)
    for i in range(1, 21):
        num_para = doc.add_paragraph(f"{{{{ protocol_step_{i}|default('') }}}}", style='List Number')
        # Only show steps with actual content
        doc.add_paragraph(f"{{% if not protocol_step_{i} %}}{{{{ '' }}}}{{% endif %}}", style="Hidden Text")
    
    # DATA ANALYSIS
    analysis_header = doc.add_paragraph("DATA ANALYSIS", style='Heading 2')
    analysis_para = doc.add_paragraph("{{ data_analysis }}")
    
    # DISCLAIMER
    disclaimer_header = doc.add_paragraph("DISCLAIMER", style='Heading 2')
    disclaimer_para = doc.add_paragraph("This material is sold for in-vitro use only in manufacturing and research. This material is not suitable for human use. It is the responsibility of the user to undertake sufficient verification and testing to determine the suitability of each product's application. The statements herein are offered for informational purposes only and are intended to be used solely for your consideration, investigation and verification.")
    
    # Add footer
    footer_section = doc.sections[0]
    footer = footer_section.footer
    
    # Add footer paragraphs
    footer_website = footer.paragraphs[0]
    footer_website.text = "www.innov-research.com"
    footer_website.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    footer_website.style = 'Footer Info Style'
    
    # Create a new paragraph for contact info
    footer_contact = footer.add_paragraph("Ph: 248.896.0145 | Fx: 248.896.0149")
    footer_contact.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    footer_contact.style = 'Footer Info Style'
    
    # Add company name to the right side
    footer_company = footer.add_paragraph("Innovative Research, Inc.")
    footer_company.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    footer_company.style = 'Footer Company Style'
    
    # Save the template
    template_path = Path('templates_docx/enhanced_template.docx')
    doc.save(template_path)
    
    logger.info(f"Created enhanced template at {template_path}")
    return template_path

if __name__ == "__main__":
    create_enhanced_template()