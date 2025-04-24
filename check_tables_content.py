#!/usr/bin/env python3
"""
Check the content of the TECHNICAL DETAILS, OVERVIEW, and REPRODUCIBILITY tables
in the output document to identify missing values.
"""

import logging
from pathlib import Path
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def check_tables_content(document_path="output_populated_template.docx"):
    """
    Check the content of the tables in the document to identify missing values.
    
    Args:
        document_path: Path to the document to check
    """
    doc = Document(document_path)
    
    print(f"\n=== Table Population Status for {document_path} ===")
    
    # Dictionary to track table population status
    table_status = {
        "TECHNICAL DETAILS": {"found": False, "populated": False, "empty_cells": 0, "total_cells": 0},
        "OVERVIEW": {"found": False, "populated": False, "empty_cells": 0, "total_cells": 0},
        "REPRODUCIBILITY": {"found": False, "populated": False, "empty_cells": 0, "total_cells": 0}
    }
    
    # Find the Technical Details table (typically first table)
    for i, table in enumerate(doc.tables):
        if not table.rows:
            continue
            
        # Extract text from first row to identify table type
        table_content = ""
        for row in table.rows:
            for cell in row.cells:
                table_content += cell.text.lower() + " "
        
        # Technical Details Table
        if "capture" in table_content and "antibod" in table_content:
            table_status["TECHNICAL DETAILS"]["found"] = True
            
            # Check cells
            for row in table.rows:
                if len(row.cells) >= 2:
                    value_cell = row.cells[1].text.strip()
                    table_status["TECHNICAL DETAILS"]["total_cells"] += 1
                    if not value_cell or value_cell == "N/A":
                        table_status["TECHNICAL DETAILS"]["empty_cells"] += 1
        
        # Overview Table            
        elif "range" in table_content and "sensitivity" in table_content:
            table_status["OVERVIEW"]["found"] = True
            
            # Check cells
            for row in table.rows:
                if len(row.cells) >= 2:
                    value_cell = row.cells[1].text.strip()
                    table_status["OVERVIEW"]["total_cells"] += 1
                    if not value_cell or value_cell == "N/A":
                        table_status["OVERVIEW"]["empty_cells"] += 1
        
        # Reproducibility Tables (combine intra-assay, inter-assay, and lot-to-lot)
        elif (("intra-assay" in table_content or "inter-assay" in table_content) and 
              "sample" in table_content and "mean" in table_content) or "lot" in table_content:
            table_status["REPRODUCIBILITY"]["found"] = True
            
            # Check cells - skip header row
            for row_idx, row in enumerate(table.rows):
                if row_idx == 0:  # Skip header
                    continue
                    
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    table_status["REPRODUCIBILITY"]["total_cells"] += 1
                    if not cell_text:
                        table_status["REPRODUCIBILITY"]["empty_cells"] += 1
    
    # Determine population status
    for table_name, status in table_status.items():
        if not status["found"]:
            print(f"{table_name} Table: Not Found")
            continue
            
        # Calculate empty percentage if there are cells
        if status["total_cells"] > 0:
            empty_percentage = (status["empty_cells"] / status["total_cells"]) * 100
            status["populated"] = empty_percentage < 5  # Less than 5% empty is considered populated
            
            print(f"{table_name} Table: Found, {'Populated' if status['populated'] else 'Missing Values'}")
            print(f"  {empty_percentage:.1f}% empty cells ({status['empty_cells']}/{status['total_cells']})")
        else:
            print(f"{table_name} Table: Found, No Content")
    
    # Overall status
    print("\n=== Overall Status ===")
    all_found = all(status["found"] for status in table_status.values())
    all_populated = all(status["populated"] for status in table_status.values() if status["found"])
    
    if all_found and all_populated:
        print("✅ All tables are found and properly populated!")
    elif all_found:
        print("⚠️ All tables are found but some contain empty cells!")
    else:
        print("❌ Some tables are missing from the document!")

def is_table_after_paragraph(doc, table, paragraph_idx):
    """
    Check if a table appears after a given paragraph.
    
    Args:
        doc: The Document object
        table: The Table object to check
        paragraph_idx: The index of the paragraph to check against
        
    Returns:
        True if the table appears after the paragraph, False otherwise
    """
    # Get the parent element of the table
    table_element = table._element
    
    # Check if the table element is positioned after the paragraph
    paragraph_element = doc.paragraphs[paragraph_idx]._element
    
    # Traverse the document structure to find their relative positions
    for child in doc.element.body:
        if child == paragraph_element:
            # Found the paragraph
            for sibling in doc.element.body:
                if sibling == table_element:
                    # Found the table after the paragraph
                    return True
        elif child == table_element:
            # Found the table before the paragraph
            return False
    
    return False

if __name__ == "__main__":
    check_tables_content()