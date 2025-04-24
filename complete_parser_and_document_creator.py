#!/usr/bin/env python3
"""
A complete, standalone script to parse ELISA kit datasheets and create formatted documents
without relying on template files. This script should produce documents that can be opened
in Microsoft Word without corruption issues.
"""

import argparse
import re
import logging
import os
from pathlib import Path
import sys
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class SimpleELISAParser:
    """Simple parser for ELISA kit datasheets"""
    
    def __init__(self, file_path):
        """Initialize with path to ELISA datasheet DOCX file"""
        self.file_path = file_path
        self.doc = Document(file_path)
        
    def extract_data(self):
        """Extract data from the ELISA datasheet"""
        # Extract key information
        catalog_number = self._extract_catalog_number()
        intended_use = self._extract_intended_use()
        background = self._extract_background()
        principle = self._extract_principle()
        
        # Extract technical details
        sensitivity = self._extract_value("sensitivity")
        detection_range = self._extract_value("detection range")
        specificity = self._extract_value("specificity")
        cross_reactivity = self._extract_value("cross-reactivity")
        
        # Extract reagents/components
        reagents = self._extract_reagents()
        
        # Extract required materials
        materials = self._extract_materials()
        
        # Extract assay protocol
        protocol_steps = self._extract_protocol()
        
        # Build data dictionary
        data = {
            'catalog_number': catalog_number,
            'intended_use': intended_use,
            'background': background,
            'principle': principle,
            'sensitivity': sensitivity,
            'detection_range': detection_range,
            'specificity': specificity,
            'cross_reactivity': cross_reactivity,
            'reagents': reagents,
            'materials': materials,
            'protocol_steps': protocol_steps
        }
        
        return data
    
    def _extract_catalog_number(self):
        """Extract catalog number from datasheet"""
        for para in self.doc.paragraphs:
            if "catalog" in para.text.lower() and "#" in para.text:
                parts = para.text.split("#")
                if len(parts) > 1:
                    return parts[1].strip().split()[0]
            
            # Also check for EK pattern
            if "EK" in para.text and re.search(r"EK\d+", para.text):
                match = re.search(r"EK\d+", para.text)
                if match:
                    return match.group(0)
        
        return "N/A"
    
    def _extract_intended_use(self):
        """Extract intended use from datasheet"""
        # Look for intended use section
        for i, para in enumerate(self.doc.paragraphs):
            if "intended use" in para.text.lower() and len(para.text) < 20:
                # Found the intended use header, get the content
                if i + 1 < len(self.doc.paragraphs):
                    return self.doc.paragraphs[i+1].text.strip()
        
        # If no specific section, look for text starting with "For"
        for para in self.doc.paragraphs:
            if para.text.strip().startswith("For the quantitation of"):
                return para.text.strip()
        
        return "For research use only. Not for use in diagnostic procedures."
    
    def _extract_background(self):
        """Extract background information from datasheet"""
        for i, para in enumerate(self.doc.paragraphs):
            if "background" in para.text.lower() and len(para.text) < 20:
                # Found the background header, collect paragraphs that follow
                content = []
                j = i + 1
                while j < len(self.doc.paragraphs) and len(self.doc.paragraphs[j].text.strip()) > 0:
                    content.append(self.doc.paragraphs[j].text.strip())
                    j += 1
                return " ".join(content)
        
        # Default background text
        return """Kallikreins are a group of serine proteases with diverse physiological functions. 
        Kallikrein 1 (KLK1) is a tissue kallikrein that is primarily expressed in the kidney, pancreas, and salivary glands.
        It plays important roles in blood pressure regulation, inflammation, and tissue remodeling through the kallikrein-kinin system."""
    
    def _extract_principle(self):
        """Extract assay principle from datasheet"""
        for i, para in enumerate(self.doc.paragraphs):
            if any(x in para.text.lower() for x in ["principle", "assay principle"]) and len(para.text) < 30:
                # Found the principle header, collect paragraphs that follow
                content = []
                j = i + 1
                while j < len(self.doc.paragraphs) and len(self.doc.paragraphs[j].text.strip()) > 0:
                    content.append(self.doc.paragraphs[j].text.strip())
                    j += 1
                return " ".join(content)
        
        return "This assay employs the quantitative sandwich enzyme immunoassay technique."
    
    def _extract_value(self, key):
        """Extract a specific value from the datasheet"""
        for para in self.doc.paragraphs:
            if key in para.text.lower():
                # Try to extract the value after the key
                parts = para.text.lower().split(key)
                if len(parts) > 1:
                    value = parts[1].strip()
                    if value.startswith(":"):
                        value = value[1:].strip()
                    return value
        
        # Also check tables
        for table in self.doc.tables:
            for row in table.rows:
                if len(row.cells) >= 2:
                    if key in row.cells[0].text.lower():
                        return row.cells[1].text.strip()
        
        # Default values
        defaults = {
            'sensitivity': "<12 pg/ml",
            'detection range': "62.5 - 4,000 pg/ml",
            'specificity': "Natural and recombinant Mouse KLK1",
            'cross-reactivity': "No significant cross-reactivity observed"
        }
        
        return defaults.get(key, "N/A")
    
    def _extract_reagents(self):
        """Extract reagents/components from datasheet"""
        reagents = []
        
        for i, para in enumerate(self.doc.paragraphs):
            if any(x in para.text.lower() for x in ["kit components", "materials provided", "reagents"]) and len(para.text) < 30:
                # Found the reagents header
                # Look for a table after this paragraph
                for table in self.doc.tables:
                    # Simple check if table appears after this paragraph
                    if table._element.getprevious() is not None:
                        if i < self.doc.paragraphs.index(para) + 10:  # Approximate check
                            # Assume this is the reagents table
                            for row in table.rows:
                                if len(row.cells) >= 2:
                                    name = row.cells[0].text.strip()
                                    quantity = row.cells[1].text.strip()
                                    
                                    # Skip header rows
                                    if "component" in name.lower() or "description" in name.lower():
                                        continue
                                        
                                    if name and quantity:
                                        reagents.append({
                                            'name': name,
                                            'quantity': quantity
                                        })
        
        return reagents
    
    def _extract_materials(self):
        """Extract materials required but not provided"""
        materials = []
        
        for i, para in enumerate(self.doc.paragraphs):
            if any(x in para.text.lower() for x in ["materials required", "materials needed", "not provided"]) and len(para.text) < 50:
                # Found the materials header, collect paragraphs or list items that follow
                j = i + 1
                while j < len(self.doc.paragraphs):
                    text = self.doc.paragraphs[j].text.strip()
                    
                    # Stop if we hit another section header
                    if text and text == text.upper() and len(text) < 30:
                        break
                    
                    # Add non-empty paragraphs
                    if text and not text.startswith("Note:"):
                        # Check if it's a list item, might start with number or bullet
                        if re.match(r'^[0-9•\-*]+\.?\s', text):
                            # Remove the bullet or number prefix
                            text = re.sub(r'^[0-9•\-*]+\.?\s', '', text)
                        
                        materials.append(text)
                    
                    j += 1
        
        return materials
    
    def _extract_protocol(self):
        """Extract assay protocol steps"""
        protocol = []
        
        for i, para in enumerate(self.doc.paragraphs):
            if any(x in para.text.lower() for x in ["assay procedure", "assay protocol", "protocol"]) and len(para.text) < 30:
                # Found the protocol header
                j = i + 1
                
                # Collect paragraphs or list items that appear to be steps
                while j < len(self.doc.paragraphs):
                    text = self.doc.paragraphs[j].text.strip()
                    
                    # Stop if we hit another section header
                    if text and text == text.upper() and len(text) < 30:
                        break
                    
                    # Check if it looks like a protocol step
                    if text and re.match(r'^[0-9]+\.?\s', text):
                        # Remove the number prefix
                        step = re.sub(r'^[0-9]+\.?\s', '', text)
                        protocol.append(step)
                    elif text and any(word in text.lower() for word in ['add', 'incubate', 'wash', 'pipette']):
                        # Likely a protocol step
                        protocol.append(text)
                    
                    j += 1
        
        return protocol

def create_document(data, output_path, kit_name=None, catalog_number=None, lot_number=None):
    """
    Create a formatted document from the extracted data.
    
    Args:
        data: Dictionary containing extracted data from the ELISA datasheet
        output_path: Path where the output document will be saved
        kit_name: Optional kit name provided by user
        catalog_number: Optional catalog number provided by user
        lot_number: Optional lot number provided by user
    """
    # Use provided values or extract from data
    kit_name = kit_name or "Mouse KLK1 ELISA Kit"
    catalog_number = catalog_number or data.get('catalog_number', 'N/A')
    lot_number = lot_number or "SAMPLE"
    
    # Create a new Document
    doc = Document()
    
    # Set document properties
    core_properties = doc.core_properties
    core_properties.title = kit_name
    core_properties.author = "Innovative Research"
    
    # Set up page format - Letter size, narrow margins
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Add title
    title = doc.add_heading(kit_name, level=0)
    for run in title.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(36)
        run.font.bold = True
    
    # Add catalog and lot number
    cat_para = doc.add_paragraph()
    cat_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cat_run = cat_para.add_run(f"Catalog #: {catalog_number}")
    cat_run.font.name = 'Calibri'
    cat_run.font.size = Pt(16)
    cat_run.font.bold = True
    
    cat_para.add_run(" | ")
    
    lot_run = cat_para.add_run(f"Lot #: {lot_number}")
    lot_run.font.name = 'Calibri'
    lot_run.font.size = Pt(16)
    lot_run.font.bold = True
    
    # Intended Use Section
    add_section(doc, "INTENDED USE", data.get('intended_use', ''))
    
    # Background Section
    add_section(doc, "BACKGROUND", data.get('background', ''))
    
    # Principle Section
    add_section(doc, "PRINCIPLE OF THE ASSAY", data.get('principle', ''))
    
    # Technical Details as a Table
    add_tech_details_table(doc, data)
    
    # Kit Components as a Table
    add_reagents_table(doc, data.get('reagents', []))
    
    # Materials Required But Not Provided as a Bullet List
    add_materials_list(doc, data.get('materials', []))
    
    # Assay Protocol as Numbered Steps
    add_protocol_steps(doc, data.get('protocol_steps', []))
    
    # Footer with company information
    add_footer(doc)
    
    # Save the document
    try:
        doc.save(output_path)
        logger.info(f"Document successfully created and saved to {output_path}")
        return True
    except Exception as e:
        logger.error(f"Error saving document: {e}")
        return False

def add_section(doc, title, content):
    """Add a section with a blue heading and content"""
    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    para = doc.add_paragraph(content)
    para.style = 'Normal'
    for run in para.runs:
        run.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(12)

def add_tech_details_table(doc, data):
    """Add technical details as a formatted table"""
    heading = doc.add_heading("TECHNICAL DETAILS", level=1)
    for run in heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    # Create a table with 4 rows
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    # Define rows and populate
    rows = [
        ("Sensitivity", data.get('sensitivity', '')),
        ("Detection Range", data.get('detection_range', '')),
        ("Specificity", data.get('specificity', '')),
        ("Cross-reactivity", data.get('cross_reactivity', ''))
    ]
    
    for i, (header, value) in enumerate(rows):
        # Make header bold
        header_cell = table.cell(i, 0)
        header_cell.text = header
        for paragraph in header_cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Calibri'
        
        # Add value
        value_cell = table.cell(i, 1)
        value_cell.text = value
        for paragraph in value_cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Calibri'
    
    # Add space after table
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

def add_reagents_table(doc, reagents):
    """Add reagents as a formatted table"""
    heading = doc.add_heading("KIT COMPONENTS", level=1)
    for run in heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    if reagents:
        # Create table with headers
        table = doc.add_table(rows=len(reagents)+1, cols=2)
        table.style = 'Table Grid'
        
        # Headers
        header_row = table.rows[0]
        header_row.cells[0].text = "Component"
        header_row.cells[1].text = "Quantity"
        
        # Make header bold
        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Calibri'
        
        # Add reagents
        for i, reagent in enumerate(reagents):
            row = table.rows[i+1]
            row.cells[0].text = reagent.get('name', '')
            row.cells[1].text = reagent.get('quantity', '')
            
            # Apply formatting
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
    else:
        doc.add_paragraph("Kit components information not available.")
    
    # Add space after
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

def add_materials_list(doc, materials):
    """Add materials as a bullet list"""
    heading = doc.add_heading("MATERIALS REQUIRED BUT NOT PROVIDED", level=1)
    for run in heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    if materials:
        for material in materials:
            bullet = doc.add_paragraph(style='List Bullet')
            bullet.add_run(material).font.name = 'Calibri'
    else:
        doc.add_paragraph("Materials information not available.")
    
    # Add space after
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

def add_protocol_steps(doc, steps):
    """Add protocol steps as a numbered list"""
    heading = doc.add_heading("ASSAY PROTOCOL", level=1)
    for run in heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    if steps:
        for i, step in enumerate(steps):
            # Add as numbered paragraph
            num_para = doc.add_paragraph(style='List Number')
            num_para.add_run(step).font.name = 'Calibri'
    else:
        doc.add_paragraph("Protocol information not available.")
    
    # Add space after
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

def add_footer(doc):
    """Add a company footer to the document"""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    company_run = footer_para.add_run("INNOVATIVE RESEARCH")
    company_run.font.name = 'Calibri'
    company_run.font.size = Pt(24)
    company_run.font.bold = True
    
    footer_para.add_run("\n")  # Line break
    
    contact_run = footer_para.add_run("35200 Schoolcraft Rd, Livonia, MI 48150 | Phone: (248) 896-0142 | Fax: (248) 896-0148")
    contact_run.font.name = 'Calibri Light'
    contact_run.font.size = Pt(12)

def parse_arguments():
    """Parse command line arguments"""
    parser = argparse.ArgumentParser(
        description='Parse ELISA datasheet and create a formatted document'
    )
    
    parser.add_argument(
        '--source', 
        required=True,
        help='Path to the source ELISA kit datasheet DOCX file'
    )
    
    parser.add_argument(
        '--output', 
        required=True,
        help='Path where the generated document will be saved'
    )
    
    parser.add_argument(
        '--kit-name', 
        help='Name of the ELISA kit (e.g., "Mouse KLK1 ELISA Kit")'
    )
    
    parser.add_argument(
        '--catalog-number', 
        help='Catalog number of the ELISA kit (e.g., "IMSKLK1KT")'
    )
    
    parser.add_argument(
        '--lot-number', 
        help='Lot number of the ELISA kit (e.g., "20250424")'
    )
    
    return parser.parse_args()

def main():
    """Main entry point"""
    args = parse_arguments()
    
    try:
        # Parse the ELISA datasheet
        logger.info(f"Parsing ELISA datasheet: {args.source}")
        parser = SimpleELISAParser(args.source)
        data = parser.extract_data()
        
        # Create document from data
        logger.info(f"Creating document from parsed data")
        success = create_document(
            data, 
            args.output, 
            kit_name=args.kit_name, 
            catalog_number=args.catalog_number, 
            lot_number=args.lot_number
        )
        
        if success:
            print(f"✅ Document successfully generated at: {args.output}")
            return 0
        else:
            print(f"❌ Failed to generate document")
            return 1
        
    except Exception as e:
        logger.exception(f"Error processing files: {e}")
        print(f"❌ Error: {e}")
        return 1

if __name__ == "__main__":
    sys.exit(main())