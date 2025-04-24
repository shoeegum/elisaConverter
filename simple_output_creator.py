#!/usr/bin/env python3
"""
Extremely simple document creator with no complex templates or dependencies.
Creates a simple DOCX file with basic formatting that should be reliably openable.
"""

import sys
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Create a new Document
doc = Document()

# Set document properties
core_properties = doc.core_properties
core_properties.title = "Mouse KLK1 ELISA Kit"
core_properties.author = "Innovative Research"

# Set up page format - Letter size, standard margins
section = doc.sections[0]
section.page_height = Inches(11)
section.page_width = Inches(8.5)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# Add title
title = doc.add_heading("Mouse KLK1 ELISA Kit", level=0)
for run in title.runs:
    run.font.name = 'Calibri'
    run.font.size = Pt(32)
    run.font.bold = True

# Add catalog and lot numbers
cat_para = doc.add_paragraph()
cat_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
cat_para.add_run("Catalog #: IMSKLK1KT | Lot #: 20250424").bold = True

# Add a page break
doc.add_page_break()

# Add some sections
sections = [
    ("INTENDED USE", "This Mouse KLK1 ELISA Kit is intended for the quantitative determination of mouse KLK1 in serum, plasma, cell culture supernatants, and other biological fluids."),
    ("BACKGROUND", "Kallikreins are a group of serine proteases with diverse physiological functions. Kallikrein 1 (KLK1) is a tissue kallikrein primarily expressed in the kidney, pancreas, and salivary glands. It plays important roles in blood pressure regulation, inflammation, and tissue remodeling through the kallikrein-kinin system."),
    ("PRINCIPLE OF THE ASSAY", "This assay employs the quantitative sandwich enzyme immunoassay technique. A monoclonal antibody specific for Mouse KLK1 has been pre-coated onto a microplate. Standards and samples are pipetted into the wells, and any KLK1 present is bound by the immobilized antibody. After washing away unbound substances, an enzyme-linked polyclonal antibody specific for Mouse KLK1 is added to the wells. Following a wash to remove unbound antibody-enzyme reagent, a substrate solution is added, and color develops in proportion to the amount of KLK1 bound in the initial step. The color development is stopped, and the intensity of the color is measured."),
    ("TECHNICAL DETAILS", "Sensitivity: <12 pg/ml\nDetection Range: 62.5 - 4,000 pg/ml\nSpecificity: Natural and recombinant Mouse KLK1\nCross-reactivity: No significant cross-reactivity with other kallikrein family members")
]

# Add each section with proper formatting
for title, content in sections:
    # Section title in blue
    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 70, 180)
        run.font.bold = True
    
    # Section content
    para = doc.add_paragraph(content)
    para.style = 'Normal'
    para.paragraph_format.space_after = Pt(12)

# Add a simple footer
section = doc.sections[0]
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_para.add_run("INNOVATIVE RESEARCH").bold = True
footer_para.add_run("\n35200 Schoolcraft Rd, Livonia, MI 48150 | (248) 896-0142")

# Save the document
try:
    output_file = "simple_output.docx"
    doc.save(output_file)
    print(f"Document successfully created and saved to {output_file}")
except Exception as e:
    print(f"Error saving document: {e}")
    sys.exit(1)