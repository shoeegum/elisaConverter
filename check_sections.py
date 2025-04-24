#!/usr/bin/env python3
"""
Check specific sections in the output document to identify missing content.
"""

from docx import Document
import sys

def check_document_sections(document_path):
    doc = Document(document_path)
    print(f"Analyzing document: {document_path}")
    print("-" * 80)
    
    # Check the Materials Required section
    print("\nMATERIALS REQUIRED BUT NOT PROVIDED SECTION:")
    print("-" * 40)
    in_materials_section = False
    materials_list_items = []
    
    for i, para in enumerate(doc.paragraphs):
        if "MATERIALS REQUIRED BUT NOT PROVIDED" in para.text:
            in_materials_section = True
            print(f"Found section at paragraph {i}: {para.text}")
        elif in_materials_section and any(s in para.text.upper() for s in ["REAGENT PREPARATION", "KIT COMPONENTS"]):
            in_materials_section = False
            print(f"Section ends at paragraph {i}: {para.text}")
        elif in_materials_section and para.text.strip() and para.style.name == "List Bullet":
            materials_list_items.append(para.text.strip())
            print(f"  - Found bullet point: {para.text.strip()}")
    
    print(f"Total bullet points found: {len(materials_list_items)}")
    
    # Check the Standard Curve Example table
    print("\nSTANDARD CURVE EXAMPLE SECTION:")
    print("-" * 40)
    found_standard_curve_table = False
    
    for i, table in enumerate(doc.tables):
        # Check if this looks like a standard curve table
        if table.rows[0].cells[0].text.strip().lower() in ["concentration", "concentration (pg/ml)"]:
            found_standard_curve_table = True
            print(f"Found Standard Curve table at index {i}")
            print(f"  Rows: {len(table.rows)}, Columns: {len(table.columns)}")
            print("  Header row:")
            for j, cell in enumerate(table.rows[0].cells):
                print(f"    Cell {j}: {cell.text}")
            if len(table.rows) > 1:
                print("  Data row:")
                for j, cell in enumerate(table.rows[1].cells):
                    print(f"    Cell {j}: {cell.text}")
            break
    
    if not found_standard_curve_table:
        print("Standard Curve table not found!")
    
    # Check the Intra/Inter-Assay Variability section
    print("\nINTRA/INTER-ASSAY VARIABILITY SECTION:")
    print("-" * 40)
    found_variability_section = False
    
    for i, para in enumerate(doc.paragraphs):
        if "INTRA/INTER-ASSAY VARIABILITY" in para.text:
            found_variability_section = True
            print(f"Found section at paragraph {i}: {para.text}")
            # Look for tables near this section
            for j, table in enumerate(doc.tables):
                if j > i and j < i + 10:  # Look for tables within 10 paragraphs
                    print(f"  Found table at index {j} with {len(table.rows)} rows and {len(table.columns)} columns")
            break
    
    if not found_variability_section:
        print("Intra/Inter-Assay Variability section not found!")
    
    # Check for Reproducibility section
    print("\nREPRODUCIBILITY SECTION:")
    print("-" * 40)
    found_reproducibility_section = False
    
    for i, para in enumerate(doc.paragraphs):
        if "REPRODUCIBILITY" in para.text:
            found_reproducibility_section = True
            print(f"Found section at paragraph {i}: {para.text}")
            # Check if next paragraph has content
            if i + 1 < len(doc.paragraphs):
                print(f"  Next paragraph: {doc.paragraphs[i+1].text[:100]}...")
            break
    
    if not found_reproducibility_section:
        print("Reproducibility section not found!")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python check_sections.py <docx_file>")
        sys.exit(1)
    
    check_document_sections(sys.argv[1])