#!/usr/bin/env python3
"""
Format Document

This script applies consistent formatting to a DOCX document:
1. Calibri font throughout
2. 1.15 line spacing for all paragraphs
3. Ensures Title formatting is correct (36pt, bold)
"""

import logging
from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.shared import Pt, RGBColor

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def apply_document_formatting(document_path):
    """
    Apply Calibri font and 1.15 line spacing to all paragraphs in the document.
    Also ensures Title formatting is correct.
    
    Args:
        document_path: Path to the document to modify
    """
    try:
        # Make a backup of the document
        document_path = Path(document_path)
        backup_path = document_path.with_name(f"{document_path.stem}_before_formatting{document_path.suffix}")
        import shutil
        shutil.copy2(document_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Load the document
        doc = Document(document_path)
        
        # First set the default style
        if 'Normal' in doc.styles:
            style = doc.styles['Normal']
            style.font.name = "Calibri"
            style.font.size = Pt(11)  # 11pt for body text
            style.paragraph_format.line_spacing = 1.15
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        
        # Ensure Title style is correct
        if 'Title' in doc.styles:
            title_style = doc.styles['Title']
            title_style.font.name = "Calibri"
            title_style.font.size = Pt(36)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
        # First check and fix the title paragraph specifically
        if len(doc.paragraphs) > 0:
            title_para = doc.paragraphs[0]
            if title_para.style.name == 'Title':
                # Make sure title paragraphs have correct formatting
                title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # Fix title paragraph formatting
                for run in title_para.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(36)
                    run.font.bold = True
                    
                # If there are no runs, add them with proper formatting
                if len(title_para.runs) == 0:
                    title_text = title_para.text
                    title_para.clear()
                    new_run = title_para.add_run(title_text)
                    new_run.font.name = "Calibri"
                    new_run.font.size = Pt(36)
                    new_run.font.bold = True
            
        # Apply to all paragraphs
        for para in doc.paragraphs:
            # Skip title paragraph which we've already handled
            if para.style.name == 'Title':
                continue
                
            # Apply paragraph formatting
            para.paragraph_format.line_spacing = 1.15
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            
            # Apply font to all runs
            for run in para.runs:
                run.font.name = "Calibri"
                # Ensure 11pt font size for body text (unless it's a heading)
                if para.style.name not in ['Heading 1', 'Heading 2', 'Heading 3', 'Title']:
                    run.font.size = Pt(11)
        
        # Apply to all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        # Apply paragraph formatting
                        para.paragraph_format.line_spacing = 1.15
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                        
                        # Apply font to all runs
                        for run in para.runs:
                            run.font.name = "Calibri"
                            
        # Make one final pass for any styled paragraphs
        for style_id in ['Heading 1', 'Heading 3', 'List Bullet', 'List Number']:
            if style_id in doc.styles:
                style = doc.styles[style_id]
                style.font.name = "Calibri"
                # Keep line spacing consistent
                style.paragraph_format.line_spacing = 1.15
                style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        
        # Specific settings for Heading 2 style (section headings)
        if 'Heading 2' in doc.styles:
            style = doc.styles['Heading 2']
            style.font.name = "Calibri"
            style.font.size = Pt(12)  # 12pt for section headings
            style.font.color.rgb = RGBColor(0, 70, 180)  # Blue color
            style.font.bold = True
            # Keep line spacing consistent
            style.paragraph_format.line_spacing = 1.15
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                
        # Save the document
        doc.save(document_path)
        logger.info(f"Successfully formatted document: {document_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error formatting document: {e}")
        return False

if __name__ == "__main__":
    # Format the current output document
    apply_document_formatting("output_populated_template.docx")