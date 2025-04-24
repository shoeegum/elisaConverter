#!/usr/bin/env python3
"""
Check the Sample Preparation, Storage, and Dilution Guideline sections
"""

from docx import Document

def check_sample_sections(document_path="output_populated_template.docx"):
    """
    Check the sample preparation, storage, and dilution guideline sections
    in both the source and output documents.
    
    Args:
        document_path: Path to the document to check
    """
    doc = Document(document_path)
    
    # Find sections by keyword
    sample_prep_idx = None
    sample_dilution_idx = None
    assay_procedure_idx = None
    
    print("=== OUTPUT DOCUMENT ANALYSIS ===")
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "SAMPLE PREPARATION AND STORAGE" in text:
            sample_prep_idx = i
            print(f"Found SAMPLE PREPARATION AND STORAGE at paragraph {i}")
        elif "SAMPLE DILUTION GUIDELINE" in text:
            sample_dilution_idx = i
            print(f"Found SAMPLE DILUTION GUIDELINE at paragraph {i}")
        elif "ASSAY PROCEDURE" in text or "ASSAY PROTOCOL" in text:
            assay_procedure_idx = i
            print(f"Found ASSAY PROCEDURE at paragraph {i}")
    
    # Check Sample Preparation section
    if sample_prep_idx is not None:
        print("\n=== SAMPLE PREPARATION AND STORAGE SECTION ===")
        # Print a few paragraphs after the section heading
        start_idx = sample_prep_idx + 1
        end_idx = sample_dilution_idx if sample_dilution_idx else min(sample_prep_idx + 10, len(doc.paragraphs))
        
        for i in range(start_idx, end_idx):
            print(f"Para {i}: {doc.paragraphs[i].text[:100]}...")
        
        # Check for tables between the headings
        tables_between = []
        for i, table in enumerate(doc.tables):
            tables_between.append(f"Table {i}: {len(table.rows)}x{len(table.columns)} - First cell: {table.cell(0, 0).text[:50]}...")
        
        if tables_between:
            print("\nTables in document:")
            for table_info in tables_between:
                print(f"  {table_info}")
    
    # Check Sample Dilution section
    if sample_dilution_idx is not None:
        print("\n=== SAMPLE DILUTION GUIDELINE SECTION ===")
        # Print content between dilution and assay procedure
        start_idx = sample_dilution_idx + 1
        end_idx = assay_procedure_idx if assay_procedure_idx else min(sample_dilution_idx + 10, len(doc.paragraphs))
        
        for i in range(start_idx, end_idx):
            print(f"Para {i}: {doc.paragraphs[i].text[:100]}...")
    
    # Check the source document as well
    print("\n=== SOURCE DOCUMENT ANALYSIS ===")
    source_doc = Document("attached_assets/EK1586_Mouse_KLK1Kallikrein_1_ELISA_Kit_PicoKine_Datasheet.docx")
    
    sample_prep_idx = None
    sample_dilution_idx = None
    sample_collection_idx = None
    assay_procedure_idx = None
    
    for i, para in enumerate(source_doc.paragraphs):
        text = para.text.strip()
        if "Sample Collection Notes" in text:
            sample_collection_idx = i
            print(f"Found Sample Collection Notes at paragraph {i}")
        elif "Specimen Collection" in text or "Sample Preparation" in text:
            sample_prep_idx = i
            print(f"Found Sample Preparation at paragraph {i}: {text}")
        elif "Sample Dilution Guideline" in text:
            sample_dilution_idx = i
            print(f"Found Sample Dilution Guideline at paragraph {i}: {text}")
        elif "Assay Procedure" in text:
            assay_procedure_idx = i
            print(f"Found Assay Procedure at paragraph {i}: {text}")
    
    # Check for tables in the source document
    print("\nTables in source document:")
    for i, table in enumerate(source_doc.tables):
        if len(table.rows) > 0 and len(table.columns) > 0:
            cols = len(table.columns)
            rows = len(table.rows)
            first_cell = table.cell(0, 0).text.strip()
            print(f"  Table {i}: {rows}x{cols} - First cell: {first_cell[:50]}...")
            
            # If this might be the sample preparation table, print more details
            if i >= 3 and i <= 5:  # Tables around the sample sections
                print("  Possible sample preparation table - Headers:")
                for j, cell in enumerate(table.row_cells(0)):
                    print(f"    Cell {j}: {cell.text}")
                if rows > 1:
                    print("  First data row:")
                    for j, cell in enumerate(table.row_cells(1)):
                        print(f"    Cell {j}: {cell.text}")
    
    # Sample Preparation content
    if sample_prep_idx is not None:
        print("\nSource Sample Preparation content:")
        start_idx = sample_prep_idx
        end_idx = sample_dilution_idx if sample_dilution_idx else min(sample_prep_idx + 10, len(source_doc.paragraphs))
        
        for i in range(start_idx, end_idx):
            print(f"Para {i}: {source_doc.paragraphs[i].text[:100]}...")
    
    # Sample Dilution content
    if sample_dilution_idx is not None:
        print("\nSource Sample Dilution Guideline content:")
        start_idx = sample_dilution_idx
        end_idx = assay_procedure_idx if assay_procedure_idx else min(sample_dilution_idx + 10, len(source_doc.paragraphs))
        
        for i in range(start_idx, end_idx):
            print(f"Para {i}: {source_doc.paragraphs[i].text[:100]}...")

if __name__ == "__main__":
    check_sample_sections()