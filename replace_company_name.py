#!/usr/bin/env python3
"""
Replace Company Name and Brand References

This script replaces:
1. All instances of "Boster" with "Innovative Research, Inc." 
2. All instances of "PicoKine®" and its variations with empty string (removing it)

It processes the entire document and maintains all formatting.
"""

import logging
import re
from pathlib import Path
import shutil
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def replace_company_references(document_path):
    """
    Replace all company and brand references in the document.
    
    Replaces:
    - "Boster" with "Innovative Research, Inc."
    - "PicoKine®" (and variations) with ""
    
    Args:
        document_path: Path to the document to modify
    """
    try:
        # Make a backup of the document
        document_path = Path(document_path)
        backup_path = document_path.with_name(f"{document_path.stem}_before_company_replace{document_path.suffix}")
        shutil.copy2(document_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Load the document
        doc = Document(document_path)
        
        # Process all paragraphs in the document
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if "Boster" in run.text:
                    run.text = run.text.replace("Boster", "Innovative Research, Inc.")
                    logger.info(f"Replaced 'Boster' with 'Innovative Research, Inc.' in paragraph")
                
                if "PicoKine®" in run.text or "PicoKine" in run.text:
                    run.text = re.sub(r'PicoKine®?', '', run.text)
                    logger.info(f"Removed 'PicoKine®' from paragraph")
        
        # Process all table cells in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if "Boster" in run.text:
                                run.text = run.text.replace("Boster", "Innovative Research, Inc.")
                                logger.info(f"Replaced 'Boster' with 'Innovative Research, Inc.' in table cell")
                            
                            if "PicoKine®" in run.text or "PicoKine" in run.text:
                                run.text = re.sub(r'PicoKine®?', '', run.text)
                                logger.info(f"Removed 'PicoKine®' from table cell")
        
        # Save the document
        doc.save(document_path)
        logger.info(f"Successfully replaced company references in: {document_path}")
        return True
    
    except Exception as e:
        logger.error(f"Error replacing company references: {e}")
        return False

if __name__ == "__main__":
    # Replace company references in the current output document
    replace_company_references("output_populated_template.docx")