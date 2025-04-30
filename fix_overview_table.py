#!/usr/bin/env python3
"""
Fix Overview Table

This script fixes the OVERVIEW table in the document to ensure it's properly populated
with Mouse KLK1/Kallikrein 1 data.
"""

import logging
from pathlib import Path
import shutil
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def fix_overview_table(document_path):
    """
    Fix the OVERVIEW table in the document to ensure it's properly populated.
    
    Args:
        document_path: Path to the document to modify
    """
    try:
        # Make a backup of the document
        document_path = Path(document_path)
        backup_path = document_path.with_name(f"{document_path.stem}_before_fixing_table{document_path.suffix}")
        shutil.copy2(document_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Load the document
        doc = Document(document_path)
        
        # Find the OVERVIEW section and table
        overview_idx = None
        overview_table = None
        
        # Find the OVERVIEW section
        for i, para in enumerate(doc.paragraphs):
            if "OVERVIEW" in para.text.upper() and "TECHNICAL DETAILS" not in para.text.upper():
                overview_idx = i
                logger.info(f"Found OVERVIEW section at paragraph {i}")
                break
        
        if overview_idx is None:
            logger.warning("Could not find OVERVIEW section")
            return False
        
        # Find the table that follows the OVERVIEW heading
        table_count = 0
        for i, table in enumerate(doc.tables):
            # Try to check if this is the overview table
            if (len(table.rows) >= 8 and len(table.rows[0].cells) >= 2 and 
                "Capture Ab" in table.rows[0].cells[0].text):
                overview_table = table
                logger.info(f"Found OVERVIEW table at index {i}")
                break
                
            # Alternative: find by checking if any row contains "Sample Type"
            for row in table.rows:
                if len(row.cells) >= 2 and "Sample Type" in row.cells[0].text:
                    overview_table = table
                    logger.info(f"Found OVERVIEW table at index {i} by 'Sample Type' keyword")
                    break
                    
            if overview_table is not None:
                break
            
            table_count += 1
        
        if overview_table is None:
            logger.warning("Could not find OVERVIEW table")
            return False
            
        # Map of expected rows and their values for Mouse KLK1
        overview_data = {
            "Capture Ab": "Rabbit polyclonal antibody",
            "Detection Ab": "Biotinylated rabbit polyclonal antibody",
            "Standard": "Recombinant mouse KLK1/Kallikrein 1 protein",
            "Sample Type": "Cell culture supernatants, cell lysates, serum, plasma",
            "Detection Method": "Colorimetric",
            "Sensitivity": "≤ 93.75 pg/mL",
            "Range": "156.25-10000 pg/mL",
            "Recovery": "80-100%"
        }
        
        # Update each row in the table
        rows_updated = 0
        for row in overview_table.rows:
            if len(row.cells) >= 2:
                cell_text = row.cells[0].text.strip()
                if cell_text in overview_data:
                    # Save the original style
                    original_style = row.cells[1].paragraphs[0].style if row.cells[1].paragraphs else None
                    
                    # Update the cell content
                    row.cells[1].text = overview_data[cell_text]
                    
                    # Restore the original style if possible
                    if original_style and row.cells[1].paragraphs:
                        row.cells[1].paragraphs[0].style = original_style
                    
                    rows_updated += 1
        
        logger.info(f"Updated {rows_updated} rows in the OVERVIEW table")
        
        # Save the document
        doc.save(document_path)
        logger.info(f"Successfully fixed OVERVIEW table: {document_path}")
        
        # Apply consistent formatting to make sure table text has proper font
        from format_document import apply_document_formatting
        apply_document_formatting(document_path)
        
        return True
        
    except Exception as e:
        logger.error(f"Error fixing OVERVIEW table: {e}")
        return False

if __name__ == "__main__":
    # Fix the OVERVIEW table in the current output document
    fix_overview_table("output_populated_template.docx")