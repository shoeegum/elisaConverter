#!/usr/bin/env python3
"""
Modify an existing ELISA kit datasheet to add numbered lists to the
Preparations Before Assay section for testing purposes.
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path

def modify_preparations_section(input_path, output_path):
    """
    Modify the Preparations Before Assay section in an existing document
    to include numbered lists.
    """
    doc = Document(input_path)
    
    # Find the Preparations Before Assay section
    prep_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "PREPARATIONS BEFORE ASSAY" in para.text.upper():
            prep_idx = i
            break
    
    if prep_idx is None:
        print("Preparations Before Assay section not found!")
        return False
    
    # Find the end of the section (next heading)
    end_idx = None
    for i in range(prep_idx + 1, len(doc.paragraphs)):
        if "KIT COMPONENTS" in doc.paragraphs[i].text.upper():
            end_idx = i
            break
    
    if end_idx is None:
        print("Could not find the end of the Preparations Before Assay section!")
        return False
    
    print(f"Found Preparations section from paragraph {prep_idx} to {end_idx}")
    
    # Collect the original paragraph texts
    original_paragraphs = []
    for i in range(prep_idx + 1, end_idx):
        text = doc.paragraphs[i].text.strip()
        if text:
            original_paragraphs.append(text)
    
    print(f"Found {len(original_paragraphs)} paragraphs in the section")
    
    # Delete all paragraphs in the section
    for _ in range(end_idx - (prep_idx + 1)):
        doc.paragraphs[prep_idx + 1]._element.getparent().remove(doc.paragraphs[prep_idx + 1]._element)
    
    # Add intro paragraphs and numbered lists after the section heading
    current_idx = prep_idx + 1
    
    # Intro paragraphs (keep the first two as regular paragraphs)
    for i in range(min(2, len(original_paragraphs))):
        p = doc.add_paragraph(original_paragraphs[i])
        # Move it to the right position
        p._element.addnext(doc.paragraphs[current_idx]._element)
        current_idx += 1
    
    # Convert some paragraphs to numbered lists
    for i in range(2, min(7, len(original_paragraphs))):
        p = doc.add_paragraph(f"{i-1}. {original_paragraphs[i]}")
        # Set numbering style - note that this might not actually apply proper numbering in python-docx
        p.style = 'List Number'
        # Move it to the right position
        p._element.addnext(doc.paragraphs[current_idx]._element)
        current_idx += 1
    
    # Add remaining paragraphs as regular text
    for i in range(7, len(original_paragraphs)):
        p = doc.add_paragraph(original_paragraphs[i])
        # Move it to the right position
        p._element.addnext(doc.paragraphs[current_idx]._element)
        current_idx += 1
    
    # Save the modified document
    doc.save(output_path)
    print(f"Modified document saved to: {output_path}")
    return True

if __name__ == "__main__":
    source_file = "attached_assets/EK1586_Mouse_KLK1Kallikrein_1_ELISA_Kit_PicoKine_Datasheet.docx"
    output_file = "attached_assets/modified_with_numbered_lists.docx"
    modify_preparations_section(source_file, output_file)