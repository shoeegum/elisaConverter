#!/usr/bin/env python3
"""
Update the enhanced template to add placeholders for intra/inter variability tables
and reproducibility section.
"""

import logging
import sys
from docx import Document
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches
from pathlib import Path

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def update_enhanced_template():
    """
    Update the enhanced template to include the variability and reproducibility tables.
    """
    # Load the existing template
    template_path = Path("templates_docx/enhanced_template.docx")
    backup_path = Path("templates_docx/enhanced_template_backup.docx")
    output_path = Path("templates_docx/enhanced_template_updated.docx")
    
    # Make a backup
    import shutil
    shutil.copy(template_path, backup_path)
    logger.info(f"Created backup at {backup_path}")
    
    # Load the document
    doc = Document(template_path)
    logger.info(f"Loaded template from {template_path}")
    
    # First check if the document already has tables for intra/inter variability
    section_updated = False
    
    # Find the variability section
    variability_section = None
    reproducibility_section = None
    for i, para in enumerate(doc.paragraphs):
        if "INTRA/INTER-ASSAY VARIABILITY" in para.text.upper():
            variability_section = i
            logger.info(f"Found variability section at paragraph {i}")
        elif "REPRODUCIBILITY" in para.text.upper():
            reproducibility_section = i
            logger.info(f"Found reproducibility section at paragraph {i}")
    
    if variability_section:
        # Look for tables near this section
        has_intra_table = False
        has_inter_table = False
        
        for table in doc.tables:
            # Check if this is an intra-assay table
            first_cell_text = table.cell(0, 0).text if table.rows and table.columns else ""
            if "Sample" in first_cell_text and "Mean" in table.cell(0, 2).text if table.rows and len(table.columns) > 2 else "":
                has_intra_table = True
                logger.info("Found intra-assay table")
                # Update the table with placeholder values
                if len(table.rows) >= 4 and len(table.columns) >= 5:
                    # Sample 1 row
                    table.cell(1, 0).text = "1"
                    table.cell(1, 1).text = "{{ intra_var_sample1_n }}"
                    table.cell(1, 2).text = "{{ intra_var_sample1_mean }}"
                    table.cell(1, 3).text = "{{ intra_var_sample1_sd }}"
                    table.cell(1, 4).text = "{{ intra_var_sample1_cv }}"
                    
                    # Sample 2 row
                    table.cell(2, 0).text = "2"
                    table.cell(2, 1).text = "{{ intra_var_sample2_n }}"
                    table.cell(2, 2).text = "{{ intra_var_sample2_mean }}"
                    table.cell(2, 3).text = "{{ intra_var_sample2_sd }}"
                    table.cell(2, 4).text = "{{ intra_var_sample2_cv }}"
                    
                    # Sample 3 row
                    table.cell(3, 0).text = "3"
                    table.cell(3, 1).text = "{{ intra_var_sample3_n }}"
                    table.cell(3, 2).text = "{{ intra_var_sample3_mean }}"
                    table.cell(3, 3).text = "{{ intra_var_sample3_sd }}"
                    table.cell(3, 4).text = "{{ intra_var_sample3_cv }}"
            
            # Check if this is an inter-assay table            
            if has_intra_table and "Sample" in first_cell_text and "Mean" in table.cell(0, 2).text if table.rows and len(table.columns) > 2 else "":
                has_inter_table = True
                logger.info("Found inter-assay table")
                # Update the table with placeholder values
                if len(table.rows) >= 4 and len(table.columns) >= 5:
                    # Sample 1 row
                    table.cell(1, 0).text = "1"
                    table.cell(1, 1).text = "{{ inter_var_sample1_n }}"
                    table.cell(1, 2).text = "{{ inter_var_sample1_mean }}"
                    table.cell(1, 3).text = "{{ inter_var_sample1_sd }}"
                    table.cell(1, 4).text = "{{ inter_var_sample1_cv }}"
                    
                    # Sample 2 row
                    table.cell(2, 0).text = "2"
                    table.cell(2, 1).text = "{{ inter_var_sample2_n }}"
                    table.cell(2, 2).text = "{{ inter_var_sample2_mean }}"
                    table.cell(2, 3).text = "{{ inter_var_sample2_sd }}"
                    table.cell(2, 4).text = "{{ inter_var_sample2_cv }}"
                    
                    # Sample 3 row
                    table.cell(3, 0).text = "3"
                    table.cell(3, 1).text = "{{ inter_var_sample3_n }}"
                    table.cell(3, 2).text = "{{ inter_var_sample3_mean }}"
                    table.cell(3, 3).text = "{{ inter_var_sample3_sd }}"
                    table.cell(3, 4).text = "{{ inter_var_sample3_cv }}"
    
    # Add intra-assay variability content if missing
    if not has_intra_table:
        # Find the paragraph index to add content
        if variability_section:
            # Add intra-assay precision description
            para_idx = variability_section + 1
            intra_desc = doc.paragraphs[para_idx]
            intra_desc.text = "Three samples of known concentration were tested on one plate to assess intra-assay precision."
            
            # Create intra-assay precision table
            intra_table = doc.add_table(rows=4, cols=5)
            intra_table.style = 'Table Grid'
            intra_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Fill header row
            header_cells = ["Sample", "n", "Mean (pg/ml)", "Standard Deviation", "CV (%)"]
            for i, text in enumerate(header_cells):
                intra_table.cell(0, i).text = text
                for paragraph in intra_table.cell(0, i).paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add placeholder values for data rows
            intra_table.cell(1, 0).text = "1"
            intra_table.cell(1, 1).text = "{{ intra_var_sample1_n }}"
            intra_table.cell(1, 2).text = "{{ intra_var_sample1_mean }}"
            intra_table.cell(1, 3).text = "{{ intra_var_sample1_sd }}"
            intra_table.cell(1, 4).text = "{{ intra_var_sample1_cv }}"
            
            intra_table.cell(2, 0).text = "2"
            intra_table.cell(2, 1).text = "{{ intra_var_sample2_n }}"
            intra_table.cell(2, 2).text = "{{ intra_var_sample2_mean }}"
            intra_table.cell(2, 3).text = "{{ intra_var_sample2_sd }}"
            intra_table.cell(2, 4).text = "{{ intra_var_sample2_cv }}"
            
            intra_table.cell(3, 0).text = "3"
            intra_table.cell(3, 1).text = "{{ intra_var_sample3_n }}"
            intra_table.cell(3, 2).text = "{{ intra_var_sample3_mean }}"
            intra_table.cell(3, 3).text = "{{ intra_var_sample3_sd }}"
            intra_table.cell(3, 4).text = "{{ intra_var_sample3_cv }}"
            
            # Center all cell contents
            for row in intra_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
            # Add spacing paragraph
            doc.add_paragraph()
            
            # Add inter-assay variability content
            inter_desc = doc.add_paragraph()
            inter_desc.text = "Three samples of known concentration were tested in separate assays to assess inter-assay precision."
            
            # Create inter-assay precision table
            inter_table = doc.add_table(rows=4, cols=5)
            inter_table.style = 'Table Grid'
            inter_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Fill header row
            for i, text in enumerate(header_cells):
                inter_table.cell(0, i).text = text
                for paragraph in inter_table.cell(0, i).paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add placeholder values for data rows
            inter_table.cell(1, 0).text = "1"
            inter_table.cell(1, 1).text = "{{ inter_var_sample1_n }}"
            inter_table.cell(1, 2).text = "{{ inter_var_sample1_mean }}"
            inter_table.cell(1, 3).text = "{{ inter_var_sample1_sd }}"
            inter_table.cell(1, 4).text = "{{ inter_var_sample1_cv }}"
            
            inter_table.cell(2, 0).text = "2"
            inter_table.cell(2, 1).text = "{{ inter_var_sample2_n }}"
            inter_table.cell(2, 2).text = "{{ inter_var_sample2_mean }}"
            inter_table.cell(2, 3).text = "{{ inter_var_sample2_sd }}"
            inter_table.cell(2, 4).text = "{{ inter_var_sample2_cv }}"
            
            inter_table.cell(3, 0).text = "3"
            inter_table.cell(3, 1).text = "{{ inter_var_sample3_n }}"
            inter_table.cell(3, 2).text = "{{ inter_var_sample3_mean }}"
            inter_table.cell(3, 3).text = "{{ inter_var_sample3_sd }}"
            inter_table.cell(3, 4).text = "{{ inter_var_sample3_cv }}"
            
            # Center all cell contents
            for row in inter_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            section_updated = True
    
    # Add reproducibility section if missing
    if not reproducibility_section:
        # Add spacing paragraph
        doc.add_paragraph()
        
        # Add reproducibility section header
        reprod_title = doc.add_paragraph("REPRODUCIBILITY")
        reprod_title.style = 'Heading 2'
        # Set blue color
        for run in reprod_title.runs:
            run.font.color.rgb = RGBColor(0, 70, 180)
        
        # Add description
        reprod_desc = doc.add_paragraph()
        reprod_desc.text = "Samples were tested in four different assay lots to assess reproducibility."
        
        # Create reproducibility table
        reprod_table = doc.add_table(rows=4, cols=7)
        reprod_table.style = 'Table Grid'
        reprod_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Fill header row
        header_cells = ["", "Lot 1", "Lot 2", "Lot 3", "Lot 4", "Mean", "CV (%)"]
        for i, text in enumerate(header_cells):
            reprod_table.cell(0, i).text = text
            for paragraph in reprod_table.cell(0, i).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add placeholder values for data rows
        reprod_table.cell(1, 0).text = "Sample 1"
        reprod_table.cell(1, 1).text = "{{ repro_sample1_lot1 }}"
        reprod_table.cell(1, 2).text = "{{ repro_sample1_lot2 }}"
        reprod_table.cell(1, 3).text = "{{ repro_sample1_lot3 }}"
        reprod_table.cell(1, 4).text = "{{ repro_sample1_lot4 }}"
        reprod_table.cell(1, 5).text = "{{ repro_sample1_mean }}"
        reprod_table.cell(1, 6).text = "{{ repro_sample1_cv }}"
        
        reprod_table.cell(2, 0).text = "Sample 2"
        reprod_table.cell(2, 1).text = "{{ repro_sample2_lot1 }}"
        reprod_table.cell(2, 2).text = "{{ repro_sample2_lot2 }}"
        reprod_table.cell(2, 3).text = "{{ repro_sample2_lot3 }}"
        reprod_table.cell(2, 4).text = "{{ repro_sample2_lot4 }}"
        reprod_table.cell(2, 5).text = "{{ repro_sample2_mean }}"
        reprod_table.cell(2, 6).text = "{{ repro_sample2_cv }}"
        
        reprod_table.cell(3, 0).text = "Sample 3"
        reprod_table.cell(3, 1).text = "{{ repro_sample3_lot1 }}"
        reprod_table.cell(3, 2).text = "{{ repro_sample3_lot2 }}"
        reprod_table.cell(3, 3).text = "{{ repro_sample3_lot3 }}"
        reprod_table.cell(3, 4).text = "{{ repro_sample3_lot4 }}"
        reprod_table.cell(3, 5).text = "{{ repro_sample3_mean }}"
        reprod_table.cell(3, 6).text = "{{ repro_sample3_cv }}"
        
        # Bold the first column
        for i in range(1, 4):
            for paragraph in reprod_table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Center all cell contents
        for row in reprod_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        section_updated = True
    
    # Check if standard curve table is correct
    curve_table_updated = False
    for i, para in enumerate(doc.paragraphs):
        if "STANDARD CURVE" in para.text.upper() or "TYPICAL DATA" in para.text.upper():
            logger.info(f"Found standard curve section at paragraph {i}")
            
            # Look for tables near this section
            for j, table in enumerate(doc.tables):
                # This is a heuristic, but should work for most cases
                if j > 1 and j < len(doc.tables) - 1:  # Skip the first technical details table and last table
                    # Check if it's the standard curve table
                    if table.rows and table.columns and "Concentration" in table.cell(0, 0).text:
                        logger.info(f"Found standard curve table at index {j}")
                        
                        # Check table dimensions
                        if len(table.rows) == 9 and len(table.columns) == 2:
                            logger.info("Standard curve table has wrong dimensions (9x2), needs to be (2x9)")
                            
                            # Get all the data from the current table
                            conc_values = []
                            od_values = []
                            for row_idx in range(1, len(table.rows)):
                                conc_cell = table.cell(row_idx, 0).text.strip()
                                od_cell = table.cell(row_idx, 1).text.strip()
                                
                                # Only include if it has template variables or valid numbers
                                if "std_conc" in conc_cell or re.search(r'\d', conc_cell):
                                    conc_values.append(conc_cell)
                                if "std_od" in od_cell or re.search(r'\d', od_cell):
                                    od_values.append(od_cell)
                            
                            # If we don't have enough values, use placeholders
                            if len(conc_values) < 8:
                                for i in range(len(conc_values), 8):
                                    conc_values.append(f"{{{{ std_conc_{i+1}|default('') }}}}")
                            
                            if len(od_values) < 8:
                                for i in range(len(od_values), 8):
                                    od_values.append(f"{{{{ std_od_{i+1}|default('') }}}}")
                            
                            # Create a new 2x9 table
                            new_table = doc.add_table(rows=2, cols=9)
                            new_table.style = 'Table Grid'
                            new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                            
                            # Fill header row
                            header_cells = ["Concentration (pg/ml)"]
                            header_cells.extend(conc_values)
                            for i, text in enumerate(header_cells):
                                if i < len(new_table.columns):
                                    new_table.cell(0, i).text = text
                                    for paragraph in new_table.cell(0, i).paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Fill data row
                            data_cells = ["O.D."]
                            data_cells.extend(od_values)
                            for i, text in enumerate(data_cells):
                                if i < len(new_table.columns):
                                    new_table.cell(1, i).text = text
                                    for paragraph in new_table.cell(1, i).paragraphs:
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Delete the old table
                            table._element.getparent().remove(table._element)
                            curve_table_updated = True
                            
                            # Move the new table to the correct position (TODO)
                            
                            break
    
    # Save the updated document
    doc.save(output_path)
    logger.info(f"Updated template saved to {output_path}")
    
    # Overwrite the original template
    import shutil
    shutil.copy(output_path, template_path)
    logger.info(f"Updated original template at {template_path}")
    
    return section_updated or curve_table_updated

if __name__ == "__main__":
    updated = update_enhanced_template()
    if updated:
        print("Template updated successfully.")
    else:
        print("No updates needed to the template.")