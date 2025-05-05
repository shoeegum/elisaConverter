#!/usr/bin/env python3
"""
Update Red Dot Footer

This script modifies the footer in Red Dot documents to:
1. Use Calibri 26pt font
2. Place 'Innovative Research, Inc.' on the right side
"""

import logging
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configure logging
logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def update_red_dot_footer(document_path):
    """
    Update the footer in Red Dot documents.
    
    Args:
        document_path: Path to the document to modify
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Create a backup of the document
        document_path = Path(document_path)
        backup_path = document_path.with_name(f"{document_path.stem}_before_footer_update{document_path.suffix}")
        shutil.copy2(document_path, backup_path)
        logger.info(f"Created backup at {backup_path}")
        
        # Load the document
        doc = Document(document_path)
        
        # Process sections
        section_count = len(doc.sections)
        logger.info(f"Found {section_count} sections in the document")
        
        for i, section in enumerate(doc.sections):
            logger.info(f"Processing section {i+1} footer")
            
            # Access the footer
            footer = section.footer
            
            # Clear existing content
            for paragraph in footer.paragraphs:
                paragraph.clear()
            
            # If there are no paragraphs, add one
            if not footer.paragraphs:
                footer.add_paragraph()
                
            # Get the first paragraph
            paragraph = footer.paragraphs[0]
            
            # Set the alignment to right
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            # Add the company name
            run = paragraph.add_run("Innovative Research, Inc.")
            
            # Set font to Calibri 26pt
            run.font.name = "Calibri"
            run.font.size = Pt(26)
            
            logger.info(f"Set footer text in section {i+1} to 'Innovative Research, Inc.' (Calibri 26pt, right-aligned)")
        
        # Save the document
        doc.save(document_path)
        logger.info(f"Successfully updated footer in: {document_path}")
        return True
        
    except Exception as e:
        logger.error(f"Error updating footer: {e}")
        return False

if __name__ == "__main__":
    import sys
    
    # Use command line argument or default
    if len(sys.argv) > 1:
        document_path = sys.argv[1]
    else:
        document_path = "red_dot_output.docx"
    
    # Update the footer
    if update_red_dot_footer(document_path):
        logger.info(f"Successfully updated footer in: {document_path}")
    else:
        logger.error(f"Failed to update footer in: {document_path}")